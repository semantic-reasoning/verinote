# SPDX-License-Identifier: MPL-2.0
import builtins
from contextlib import contextmanager
from dataclasses import replace
import functools
import importlib.resources
import importlib.util
import inspect
import json
import logging
from pathlib import Path
import re
import sqlite3
import sys
import threading
import time
import unicodedata
from html import unescape

import pytest

pytest.importorskip("fastapi")
from fastapi.testclient import TestClient  # noqa: E402

import verinote.config  # noqa: E402
import verinote.llm.openrouter_adapter  # noqa: E402
import verinote.web.app as webapp  # noqa: E402
from verinote.config import (  # noqa: E402
    Config,
    ConfigCorruptError,
    CredentialsCorruptError,
    app_config_path,
    read_app_config,
    save_settings,
)
from verinote.engine import CheckReport, DEFAULT_POLICY, FindingDetail  # noqa: E402
from verinote.engine.terms import Atom, Compound, StringLit  # noqa: E402
from verinote.kb_location import KBRootSafetyError  # noqa: E402
from verinote.llm.anthropic_adapter import AnthropicAdapter  # noqa: E402
from verinote.llm.base import ExtractedFact, LLMError, ModelListing  # noqa: E402
from verinote.llm.claude_cli_adapter import ClaudeCliAdapter  # noqa: E402
from verinote.llm.ollama_adapter import OllamaAdapter  # noqa: E402
from verinote.llm.openai_adapter import OpenAIAdapter  # noqa: E402
from verinote.pipeline import ChunkedExtractionResult, ExtractionJobBusyError  # noqa: E402
from verinote.pipeline.verify import _with_unrecorded_policy_warning  # noqa: E402
from verinote.pipeline.policy_state import (  # noqa: E402
    POLICY_RELPATH,
    PolicyMissingError,
    policy_sha256,
)
from verinote.pipeline.query import query_path  # noqa: E402
from verinote.pipeline.query_intent import parse_query_intent  # noqa: E402
from verinote.store import Store  # noqa: E402
from verinote.store import db as store_db  # noqa: E402
from verinote.store.duckdb_fact_terms import (  # noqa: E402
    DuckDBFactTermStoreLockedError,
    fact_terms_path,
)
from verinote.store.fact_input import structural_term  # noqa: E402
from verinote.web import create_app  # noqa: E402


def _client(tmp_path) -> TestClient:
    cfg = Config(
        root=tmp_path, db_path=tmp_path / "kb.sqlite",
        provider="anthropic", model="m", api_key=None, base_url=None,
    )
    app = create_app(cfg)
    client = TestClient(app)
    store = app.state.store
    client.fact_id = store.add_fact("A", "is_a", "B", status="needs_review", confidence=0.9)
    return client


def _target(kind: str, value: str | None) -> dict | None:
    return None if value is None else {"kind": kind, "value": value}


def _intent(kind: str, *, subject: str | None = None) -> dict:
    return {
        "kind": kind,
        "subject": _target("entity", subject),
        "relation": None,
        "object": None,
        "relation_candidates": [],
        "operator": None,
        "value_type": None,
        "value": None,
        "reason": None,
    }


class IntentOnlyClient:
    name = "intent-only"

    def __init__(self, intent):
        self.intent = intent
        self.intent_calls = 0
        self.direct_datalog_calls = 0

    def extract_query_intent(self, *, question: str, schema_hint: str = ""):
        self.intent_calls += 1
        raw = self.intent(question) if callable(self.intent) else self.intent
        return parse_query_intent(raw)

    def translate_query(self, *, question: str, qid: int, schema_hint: str = "") -> str:
        self.direct_datalog_calls += 1
        raise AssertionError("supported planner path must not call direct Datalog")


def _wait_for(assertion, *, timeout: float = 2.0) -> None:
    deadline = time.monotonic() + timeout
    last_error = None
    while time.monotonic() < deadline:
        try:
            assertion()
            return
        except AssertionError as e:
            last_error = e
            time.sleep(0.01)
    if last_error is not None:
        raise last_error


def test_dashboard_renders(tmp_path):
    c = _client(tmp_path)
    r = c.get("/")
    assert r.status_code == 200
    assert "verinote" in r.text


def test_report_falls_back_to_legacy_finding_strings(tmp_path, monkeypatch):
    report = CheckReport(
        ok=True,
        errors=0,
        warnings=1,
        text="legacy report",
        findings=["WARN legacy: rendered without structured metadata"],
    )
    monkeypatch.setattr(webapp, "verify", lambda _store: report)

    body = _client(tmp_path).get("/report").text

    assert "finding-list-legacy" in body
    assert "WARN legacy: rendered without structured metadata" in body


def test_report_renders_structured_finding_metadata(tmp_path, monkeypatch):
    finding = "ERROR synthetic_rule: Example"
    report = CheckReport(
        ok=False,
        errors=1,
        warnings=0,
        text="structured report",
        findings=[finding],
        finding_details=[FindingDetail(finding, "error", "synthetic_rule")],
    )
    monkeypatch.setattr(webapp, "verify", lambda _store: report)

    body = _client(tmp_path).get("/report").text

    assert 'class="finding finding-error" data-code="synthetic_rule"' in body
    assert finding in body


def test_report_renders_all_legacy_findings_after_policy_warning(tmp_path, monkeypatch):
    legacy_findings = [
        "ERROR legacy_rule: first finding",
        "WARN legacy_rule: second finding",
    ]
    report = _with_unrecorded_policy_warning(
        CheckReport(
            ok=False,
            errors=1,
            warnings=0,
            text="\n".join(legacy_findings),
            findings=legacy_findings,
        )
    )
    monkeypatch.setattr(webapp, "verify", lambda _store: report)

    body = _client(tmp_path).get("/report").text

    assert [detail.text for detail in report.finding_details] == report.findings
    assert all(finding in body for finding in report.findings)
    assert [body.index(finding) for finding in report.findings] == sorted(
        body.index(finding) for finding in report.findings
    )


def test_dashboard_shows_factlog_borrowed_source_signals(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    a = store.add_source("sources/a.md")
    b = store.add_source("sources/b.md")
    csrc = store.add_source("sources/c.md")
    store.add_fact("Acme", "uses", "FastAPI", status="confirmed", source_id=a)
    store.add_fact("Acme", "uses", "FastAPI", status="accepted", source_id=b)
    store.add_fact("Acme", "uses", "FastAPI", status="candidate", source_id=csrc)
    store.add_fact("Org", "established_on", "2020", status="confirmed", source_id=a)
    store.add_fact("Org", "established_on", "2021", status="confirmed", source_id=b)

    body = unescape(c.get("/").text)

    assert "Source corroboration" in body
    assert "Acme" in body
    assert "FastAPI" in body
    assert ">2</td>" in body
    assert "Single-valued conflicts" in body
    assert "Org" in body
    assert "2020" in body
    assert "2021" in body
    assert "(1 source)" in body


def _engine_input_card(body: str) -> int:
    """Read the number the dashboard's "engine input" card actually renders."""
    match = re.search(
        r'<div class="num">(\d+)</div><div class="lbl">engine input</div>', body
    )
    assert match is not None, "engine input card not found in dashboard"
    return int(match.group(1))


def _review_card(body: str) -> int:
    match = re.search(
        r'<div class="num">(\d+)</div><div class="lbl">needs review</div>', body
    )
    assert match is not None, "needs review card not found in dashboard"
    return int(match.group(1))


def test_dashboard_engine_input_card_follows_engine_statuses(tmp_path, monkeypatch):
    """The dashboard's headline card must answer from ENGINE_STATUSES too.

    It used to sum `confirmed + accepted` in the template — a third answer to
    "what is engine input". Widening the tier must move this card in lockstep
    with coverage and the Sources badge, or the top of the UI asserts a number
    the engine disagrees with.
    """
    c = _client(tmp_path)  # seeds one needs_review fact
    store = c.app.state.store
    sid = store.add_source("sources/a.md")
    store.add_fact("A", "is_a", "B", status="confirmed", source_id=sid)
    store.add_fact("C", "is_a", "D", status="accepted", source_id=sid)
    store.add_fact("E", "is_a", "F", status="superseded", source_id=sid)

    assert _engine_input_card(c.get("/").text) == 2

    monkeypatch.setattr(
        store_db, "ENGINE_STATUSES", store_db.ENGINE_STATUSES | {"superseded"}
    )

    body = c.get("/").text
    assert _engine_input_card(body) == 3
    # ...and it still agrees with what the engine would actually read.
    assert _engine_input_card(body) == len(
        store.facts(statuses=store_db.ENGINE_STATUSES)
    )


def test_dashboard_uses_call_time_review_and_all_status_accessors(tmp_path, monkeypatch):
    c = _client(tmp_path)
    store = c.app.state.store
    store.add_fact("C", "is_a", "D", status="candidate")
    store.add_fact("E", "is_a", "F", status="confirmed")

    assert _review_card(c.get("/").text) == 2

    monkeypatch.setattr(store_db, "REVIEW_STATUSES", frozenset({"candidate"}))
    monkeypatch.setattr(
        store_db,
        "ALL_FACT_STATUSES",
        frozenset({"candidate", "confirmed", "superseded"}),
    )

    body = c.get("/").text
    assert _review_card(body) == 1
    assert '<span class="badge badge-needs_review">needs_review</span>' not in body
    assert '<span class="badge badge-candidate">candidate</span>' in body
    assert '<span class="badge badge-confirmed">confirmed</span>' in body


def test_dashboard_shows_action_queue_counts_and_links(tmp_path):
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="anthropic",
        model="m",
        api_key=None,
        base_url=None,
    )
    c = TestClient(create_app(cfg))
    store = c.app.state.store
    store.add_fact("Unsupported Sample", "uses", "Sample Service", status="candidate")
    review_source = store.add_source("sources/review.txt")
    support_a = store.add_source("sources/support-a.txt")
    support_b = store.add_source("sources/support-b.txt")
    store.add_fact(
        "Reviewed Sample",
        "uses",
        "Sample Service",
        status="candidate",
        source_id=review_source,
    )
    store.add_fact(
        "Reviewed Sample",
        "uses",
        "Sample Service",
        status="confirmed",
        source_id=support_a,
    )
    store.add_fact(
        "Reviewed Sample",
        "uses",
        "Sample Service",
        status="accepted",
        source_id=support_b,
    )
    conflict_a = store.add_source("sources/conflict-a.txt")
    conflict_b = store.add_source("sources/conflict-b.txt")
    store.add_fact("Org", "established_on", "2020", status="confirmed", source_id=conflict_a)
    store.add_fact("Org", "established_on", "2021", status="accepted", source_id=conflict_b)
    failed_source = store.add_source("sources/failed.txt")
    job_id = store.create_extraction_job(
        source_id=failed_source,
        provider="fake",
        model="sample-model",
        total_chunks=1,
    )
    chunk_id = store.add_source_chunks(
        job_id=job_id,
        source_id=failed_source,
        chunks=["Sample body"],
    )[0]
    store.mark_extraction_job_running(job_id)
    store.mark_chunk_running(chunk_id)
    store.mark_chunk_failed(chunk_id, "provider down")
    changed = store.add_fact("Changed Sample", "uses", "Service", status="confirmed")
    store.amend_fact(changed, subject="Changed Sample", relation="uses", obj="Service v2")

    body = unescape(c.get("/").text)

    assert "Action queues" in body
    assert "Unsupported review items" in body
    assert "Corroborated review targets" in body
    assert "Single-valued conflicts" in body
    assert "Failed source analyses" in body
    assert "Recent lifecycle changes" in body
    assert "Source-backed engine facts" in body
    assert 'href="/review?filter=unsupported"' in body
    assert 'href="/review?filter=corroborated"' in body
    assert 'href="/workbench"' in body
    assert 'href="/sources"' in body
    assert ">1</td>" in body
    assert ">3</td>" in body


def test_no_active_kb_shows_selector(tmp_path, monkeypatch):
    monkeypatch.delenv("VERINOTE_ROOT", raising=False)
    monkeypatch.setenv("HOME", str(tmp_path / "home"))
    monkeypatch.setenv("XDG_CONFIG_HOME", str(tmp_path / "xdg"))
    monkeypatch.setenv("APPDATA", str(tmp_path / "appdata"))
    monkeypatch.chdir(tmp_path)

    c = TestClient(create_app())
    r = c.get("/")

    assert r.status_code == 200
    assert "Select a knowledge base" in r.text


def test_select_kb_activates_app(tmp_path, monkeypatch):
    monkeypatch.delenv("VERINOTE_ROOT", raising=False)
    monkeypatch.setenv("HOME", str(tmp_path / "home"))
    monkeypatch.setenv("XDG_CONFIG_HOME", str(tmp_path / "xdg"))
    monkeypatch.setenv("APPDATA", str(tmp_path / "appdata"))
    monkeypatch.chdir(tmp_path)
    kb = tmp_path / "chosen"

    c = TestClient(create_app())
    r = c.post("/kb/select", data={"root": str(kb)}, follow_redirects=False)

    assert r.status_code == 303
    assert (kb / "kb.sqlite").is_file()
    assert (kb / "policy" / "logic-policy.dl").is_file()
    assert c.app.state.cfg.root == kb.resolve()
    assert not app_config_path().exists()
    assert "Knowledge base" in c.get("/").text


def _unsafe_ui_root(tmp_path, kind):
    worktree = tmp_path / "synthetic-worktree"
    worktree.mkdir()
    if kind == "linked-worktree":
        (worktree / ".git").write_text("gitdir: synthetic\n", encoding="utf-8")
    else:
        (worktree / ".git").mkdir()
    expected = worktree / "nested" / "kb"
    if kind == "symlink":
        alias = tmp_path / "synthetic-worktree-alias"
        alias.symlink_to(worktree, target_is_directory=True)
        return alias / "nested" / "kb", expected
    return expected, expected


@pytest.mark.parametrize("kind", ["normal-worktree", "linked-worktree", "symlink"])
def test_ui_startup_refuses_worktree_descendant_before_initializing(tmp_path, kind):
    root, expected = _unsafe_ui_root(tmp_path, kind)

    with pytest.raises(KBRootSafetyError, match="inside Git worktree"):
        create_app(Config.for_root(root))

    assert not expected.exists()


@pytest.mark.parametrize("path", ["/kb/select", "/settings/root"])
@pytest.mark.parametrize("kind", ["normal-worktree", "linked-worktree", "symlink"])
def test_ui_root_selection_refuses_worktree_descendant_before_initializing(
    tmp_path, path, kind
):
    root, expected = _unsafe_ui_root(tmp_path, kind)
    client = (
        _client(tmp_path)
        if path == "/settings/root"
        else TestClient(create_app())
    )

    response = client.post(path, data={"root": str(root)}, follow_redirects=False)

    assert response.status_code == 400
    assert "inside Git worktree" in response.text
    assert not expected.exists()


@pytest.mark.parametrize("path", ["/kb/select", "/settings/root"])
def test_kb_root_selection_reports_unopenable_db_instead_of_500(tmp_path, path, monkeypatch):
    """A KB root whose kb.sqlite cannot be opened (e.g. a directory owned by a
    different account, or a locked file) surfaces as a 400 with a clear message,
    not an unhandled 500 — sqlite3.OperationalError must be caught alongside
    KBLocationError/OSError in _open_root's callers (#Windows admin-owned KB dir)."""
    target = tmp_path / "kb"
    target.mkdir()
    client = _client(tmp_path) if path == "/settings/root" else TestClient(create_app())

    def boom(self, db_path):
        raise sqlite3.OperationalError("unable to open database file")

    monkeypatch.setattr(store_db.Store, "__init__", boom)

    response = client.post(path, data={"root": str(target)}, follow_redirects=False)

    assert response.status_code == 400
    assert "could not open KB" in response.text


def test_dashboard_shows_coverage_gap(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    sid = store.add_source("sources/a.txt")  # no file on disk
    store.add_fact("X", "is_a", "Y", status="needs_review", source_id=sid)
    r = c.get("/")
    assert r.status_code == 200
    assert "Coverage" in r.text
    assert "sources/a.txt" in r.text
    assert "gap" in r.text


def test_review_shows_queue(tmp_path):
    c = _client(tmp_path)
    r = c.get("/review")
    assert r.status_code == 200
    assert "is_a" in r.text
    assert "Inspect" in r.text
    assert "Conf." not in r.text


def test_review_paginates_large_queue_and_preserves_controls(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    for idx in range(1200):
        store.add_fact(f"Bulk {idx:03d}", "uses", "Sample", status="candidate")

    body = unescape(c.get("/review").text)

    assert body.count('<tr id="fact-') == 50
    assert "Showing 1-50 of 1201 review facts" in body
    assert "Bulk 1199" in body
    assert "Bulk 1149" not in body
    assert 'href="/review?filter=needs-human-decision&sort=newest&page_size=50&page=2"' in body

    page_two = unescape(c.get("/review?page=2").text)
    assert page_two.count('<tr id="fact-') == 50
    assert "Showing 51-100 of 1201 review facts" in page_two
    assert "Bulk 1149" in page_two
    assert "Bulk 1099" not in page_two

    controls = unescape(c.get("/review?sort=oldest&page_size=25&page=2").text)
    assert 'href="/review?filter=unsupported&sort=oldest&page_size=25&page=1"' in controls
    assert '<option value="oldest" selected>Oldest</option>' in controls
    assert '<option value="25" selected>25</option>' in controls


def test_review_clamps_invalid_pagination_params(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    for idx in range(60):
        store.add_fact(f"Clamp {idx:03d}", "uses", "Sample", status="candidate")

    body = unescape(c.get("/review?page=bad&page_size=1000&sort=subject").text)

    assert body.count('<tr id="fact-') == 50
    assert "Showing 1-50 of 61 review facts" in body
    assert '<option value="newest" selected>Newest</option>' in body
    assert '<option value="50" selected>50</option>' in body


def test_review_trust_filter_applies_before_pagination(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    for idx in range(60):
        source_id = store.add_source(f"sources/support-{idx:03d}.txt")
        store.add_fact(f"Supported {idx:03d}", "uses", "Sample", status="candidate")
        store.add_fact(
            f"Supported {idx:03d}",
            "uses",
            "Sample",
            status="confirmed",
            source_id=source_id,
        )

    unfiltered = unescape(c.get("/review").text)
    assert "is_a" not in unfiltered

    body = unescape(c.get("/review?filter=unsupported").text)

    assert body.count('<tr id="fact-') == 1
    assert "Showing 1-1 of 1 review facts" in body
    assert "is_a" in body
    assert "Supported 059" not in body
    assert "Trust filter is applied to this page" not in body


def test_review_renders_structural_terms_from_duckdb_and_distinguishes_strings(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    store.add_fact(
        'person("Ada")',
        "has_role",
        'role(person("Ada"), "PI")',
        status="candidate",
    )
    store.add_fact(
        structural_term('person("Ada")'),
        structural_term("has_role"),
        structural_term('role(person("Ada"), "PI")'),
        status="candidate",
    )

    body = unescape(c.get("/review").text)

    assert 'class="subj term-string">"person(\\"Ada\\")"' in body
    assert 'class="subj term-term">person("Ada")' in body
    assert 'class="obj term-term">role(person("Ada"), "PI")' in body


def test_review_rows_show_trust_signals_evidence_and_inspect_link(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    policy = tmp_path / "policy"
    policy.mkdir()
    (policy / "logic-policy.dl").write_text(
        '.decl functional(rel: symbol)\nfunctional("published_year").\n',
        encoding="utf-8",
    )
    (policy / "relation-aliases.md").write_text(
        "- `publication_year` -> `published_year`\n",
        encoding="utf-8",
    )
    (policy / "typed-relations.md").write_text(
        "- published_year : number as year_number\n",
        encoding="utf-8",
    )
    source_id = store.add_source("sources/sample-candidate.txt")
    artifact_id = store.add_source_artifact(
        source_id=source_id,
        kind="original_text",
        path="sources/sample-candidate.txt",
    )
    job_id = store.create_extraction_job(
        source_id=source_id,
        artifact_id=artifact_id,
        provider="fake",
        model="sample-model",
        total_chunks=1,
    )
    chunk_id = store.add_source_chunks(
        job_id=job_id,
        source_id=source_id,
        chunks=["Sample Report was published in 2024."],
    )[0]
    fact_id = store.add_fact(
        "Sample Report",
        "publication_year",
        "2024",
        status="candidate",
        source_id=source_id,
        job_id=job_id,
    )
    store.add_fact_evidence(
        fact_id=fact_id,
        source_id=source_id,
        artifact_id=artifact_id,
        job_id=job_id,
        chunk_id=chunk_id,
        snippet="Sample Report was published in 2024.",
    )
    support_a = store.add_source("sources/sample-support-a.txt")
    support_b = store.add_source("sources/sample-support-b.txt")
    conflict = store.add_source("sources/sample-conflict.txt")
    store.add_fact(
        "Sample Report",
        "published_year",
        "2024",
        status="accepted",
        source_id=support_a,
    )
    store.add_fact(
        "Sample Report",
        "published_year",
        "2024",
        status="confirmed",
        source_id=support_b,
    )
    store.add_fact(
        "Sample Report",
        "published_year",
        "2025",
        status="accepted",
        source_id=conflict,
    )

    body = unescape(c.get("/review").text)

    assert "source backed" in body
    assert "corroborated" in body
    assert "conflicted" in body
    assert "canonical" in body
    assert "published_year" in body
    assert "year_number" in body
    assert "Sample Report was published in 2024." in body
    assert f'href="/facts/{fact_id}/provenance"' in body
    assert "Inspect" in body


def test_review_filters_by_deterministic_trust_signals(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    unsupported = store.add_fact(
        "Unsupported Sample",
        "uses",
        "Sample Service",
        status="candidate",
    )
    source_a = store.add_source("sources/sample-a.txt")
    source_b = store.add_source("sources/sample-b.txt")
    reviewed = store.add_source("sources/sample-reviewed.txt")
    corroborated = store.add_fact(
        "Reviewed Sample",
        "uses",
        "Sample Service",
        status="candidate",
        source_id=reviewed,
    )
    store.add_fact(
        "Reviewed Sample",
        "uses",
        "Sample Service",
        status="accepted",
        source_id=source_a,
    )
    store.add_fact(
        "Reviewed Sample",
        "uses",
        "Sample Service",
        status="confirmed",
        source_id=source_b,
    )

    unsupported_body = unescape(c.get("/review?filter=unsupported").text)
    assert "Unsupported Sample" in unsupported_body
    assert "Reviewed Sample" not in unsupported_body
    assert f'href="/facts/{unsupported}/provenance"' in unsupported_body

    corroborated_body = unescape(c.get("/review?filter=corroborated").text)
    assert "Reviewed Sample" in corroborated_body
    assert "Unsupported Sample" not in corroborated_body
    assert f'href="/facts/{corroborated}/provenance"' in corroborated_body


def test_workbench_renders_corroboration_conflicts_and_normalization(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    policy = tmp_path / "policy"
    policy.mkdir()
    (policy / "logic-policy.dl").write_text(
        '.decl functional(rel: symbol)\n'
        'functional("published_year").\n'
        'functional("revenue").\n',
        encoding="utf-8",
    )
    (policy / "relation-aliases.md").write_text(
        "- `pub_year` -> `published_year`\n",
        encoding="utf-8",
    )
    (policy / "typed-relations.md").write_text(
        "- revenue : amount as revenue_scalar\n",
        encoding="utf-8",
    )
    source_a = store.add_source("sources/a.txt")
    source_b = store.add_source("sources/b.txt")
    source_c = store.add_source("sources/c.txt")
    candidate_source = store.add_source("sources/candidate.txt")
    fact_id = store.add_fact(
        "Sample Report",
        "pub_year",
        "2024",
        status="confirmed",
        source_id=source_a,
    )
    store.add_fact(
        "Sample Report",
        "published_year",
        "2024",
        status="accepted",
        source_id=source_b,
    )
    candidate_id = store.add_fact(
        "Sample Report",
        "published_year",
        "2024",
        status="candidate",
        source_id=candidate_source,
    )
    store.add_fact(
        "Sample Company",
        "revenue",
        'amount(5000,"억")',
        status="confirmed",
        source_id=source_a,
    )
    store.add_fact(
        "Sample Company",
        "revenue",
        'amount(0.54,"조")',
        status="accepted",
        source_id=source_b,
    )
    store.add_fact(
        "Sample Company",
        "revenue",
        'amount(5400,"억")',
        status="confirmed",
        source_id=source_c,
    )

    body = unescape(c.get("/workbench").text)

    assert "Trust workbench" in body
    assert "Corroborated facts" in body
    assert "Sample Report" in body
    assert "pub_year -> published_year" in body
    assert "sources/a.txt" in body
    assert "sources/b.txt" in body
    assert "Related candidates" in body
    assert f'href="/facts/{candidate_id}/provenance"' in body
    assert "Single-valued conflicts" in body
    assert "Sample Company" in body
    assert 'amount(5000,"억")' in body
    assert 'amount(0.54,"조")' in body
    assert "revenue_scalar=540000000000" in body
    assert f'href="/facts/{fact_id}/provenance"' in body


def test_review_shows_accept_recommendation(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    policy = tmp_path / "policy"
    policy.mkdir()
    (policy / "logic-policy.dl").write_text(
        '.decl functional(rel: symbol)\nfunctional("published_year").\n',
        encoding="utf-8",
    )
    source_a = store.add_source("sources/a.txt")
    source_b = store.add_source("sources/b.txt")
    job_a = store.create_extraction_job(
        source_id=source_a,
        provider="fake",
        model="sample-model",
        total_chunks=1,
    )
    job_b = store.create_extraction_job(
        source_id=source_b,
        provider="fake",
        model="sample-model",
        total_chunks=1,
    )
    chunk_a = store.add_source_chunks(job_id=job_a, source_id=source_a, chunks=["a"])[0]
    chunk_b = store.add_source_chunks(job_id=job_b, source_id=source_b, chunks=["b"])[0]
    store.mark_extraction_job_running(job_a)
    store.mark_chunk_running(chunk_a)
    store.mark_chunk_done(chunk_a)
    store.finish_extraction_job(job_a)
    store.mark_extraction_job_running(job_b)
    store.mark_chunk_running(chunk_b)
    store.mark_chunk_done(chunk_b)
    store.finish_extraction_job(job_b)
    store.add_fact(
        "Sample Report",
        "published_year",
        "2024",
        status="candidate",
        source_id=source_a,
        job_id=job_a,
    )
    store.add_fact(
        "Sample Report",
        "published_year",
        "2024",
        status="confirmed",
        source_id=source_b,
        job_id=job_b,
    )

    body = unescape(c.get("/review").text)

    assert "accept recommended" in body


def test_toggle_endpoint_swaps_row(tmp_path):
    c = _client(tmp_path)
    r = c.post(f"/facts/{c.fact_id}/toggle")
    assert r.status_code == 200
    assert "confirmed" in r.text
    # the only queued fact was promoted, so the review queue is now empty
    assert "Review queue is empty" in c.get("/review").text


def test_upload_extracts_and_redirects(tmp_path, monkeypatch, fake_client):
    monkeypatch.setattr(
        webapp, "get_client", lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)])
    )
    c = _client(tmp_path)
    r = c.post(
        "/sources",
        files={"file": ("note.txt", b"some text", "text/plain")},
        follow_redirects=False,
    )
    assert r.status_code == 303
    assert r.headers["location"] == "/sources"
    # the file is saved immediately; extraction finishes in a background job.
    assert (tmp_path / "sources" / "note.txt").read_text() == "some text"

    def extracted():
        assert "is_a" in c.get("/review").text
        body = c.get("/sources").text
        assert "Analysis complete: 1/1 chunk(s)" in body
        assert "1 awaiting review" in body

    _wait_for(extracted)


def test_upload_normalizes_nfc_source_identity_and_chunk_text(tmp_path, monkeypatch):
    nfd_filename = unicodedata.normalize("NFD", "\ubb38\uc11c.txt")
    nfc_filename = unicodedata.normalize("NFC", nfd_filename)
    nfd_text = unicodedata.normalize("NFD", "\uac00\uac01 \uc815\ubcf4")
    nfc_text = unicodedata.normalize("NFC", nfd_text)
    assert nfd_filename != nfc_filename
    assert nfd_text != nfc_text

    chunk_inputs = []
    real_create_chunked_job = webapp.create_chunked_extraction_job

    def capture_chunk_job(store, **kwargs):
        chunk_inputs.append(kwargs["source_text"])
        return real_create_chunked_job(store, **kwargs)

    monkeypatch.setattr(webapp, "create_chunked_extraction_job", capture_chunk_job)
    monkeypatch.setattr(webapp, "get_client", lambda _cfg: object())
    processed_jobs = set()
    processed_lock = threading.Lock()
    first_worker_finished = threading.Event()
    second_worker_finished = threading.Event()

    def complete_job(store, _client, *, job_id, **_kwargs):
        assert store.claim_pending_extraction_job(job_id)
        store.finish_extraction_job(job_id)
        with processed_lock:
            processed_jobs.add(job_id)
            if len(processed_jobs) == 1:
                first_worker_finished.set()
            elif len(processed_jobs) == 2:
                second_worker_finished.set()
        return ChunkedExtractionResult(job_id=job_id)

    monkeypatch.setattr(
        webapp,
        "process_extraction_job",
        complete_job,
    )
    c = _client(tmp_path)

    first = c.post(
        "/sources",
        files={"file": (nfd_filename, nfd_text.encode("utf-8"), "text/plain")},
        follow_redirects=False,
    )
    assert first.status_code == 303
    assert first_worker_finished.wait(timeout=2.0)
    assert (tmp_path / "sources" / nfc_filename).read_bytes() == nfd_text.encode("utf-8")
    source = c.app.state.store.get_source_by_path(f"sources/{nfc_filename}")
    assert source is not None
    artifact = c.app.state.store.latest_source_text_artifact(int(source["id"]))
    assert artifact is not None
    assert (tmp_path / artifact["path"]).read_text(encoding="utf-8") == nfc_text
    second = c.post(
        "/sources",
        files={"file": (nfc_filename, nfc_text.encode("utf-8"), "text/plain")},
        follow_redirects=False,
    )

    assert second.status_code == 303
    assert second_worker_finished.wait(timeout=2.0)
    sources = c.app.state.store.sources()
    assert len(sources) == 1
    assert sources[0]["path"] == f"sources/{nfc_filename}"
    assert chunk_inputs == [nfc_text, nfc_text]


def test_upload_nfd_korean_fact_and_query_match_nfc(
    tmp_path, monkeypatch, fake_client, intent_payload
):
    """#411: NFC/NFD-equivalent Korean input must remain query-equivalent end to end."""
    nfd_filename = unicodedata.normalize("NFD", "합성문서.txt")
    nfc_filename = unicodedata.normalize("NFC", nfd_filename)
    nfd_subject = unicodedata.normalize("NFD", "합성프로젝트")
    nfc_subject = unicodedata.normalize("NFC", nfd_subject)
    nfd_relation = unicodedata.normalize("NFD", "목적")
    nfc_relation = unicodedata.normalize("NFC", nfd_relation)
    nfd_object = unicodedata.normalize("NFD", "검증목표")
    nfc_object = unicodedata.normalize("NFC", nfd_object)
    nfd_text = f"{nfd_subject} {nfd_relation} {nfd_object}"
    nfc_text = unicodedata.normalize("NFC", nfd_text)
    assert nfd_text != nfc_text

    nfc_question = f"사실 확인 {nfc_subject}"
    nfd_question = unicodedata.normalize("NFD", nfc_question)
    assert nfd_question != nfc_question

    def intent_for(question: str):
        assert question in {nfc_question, nfd_question}
        return intent_payload(
            "lookup_object",
            subject=nfd_subject if question == nfd_question else nfc_subject,
            relation=nfd_relation if question == nfd_question else nfc_relation,
        )

    client = fake_client(
        [ExtractedFact(nfd_subject, nfd_relation, nfd_object, 0.9)], intent=intent_for
    )
    extraction_inputs = []
    extract_facts = client.extract_facts

    def capture_extract_facts(*, source_text: str, schema_hint: str = ""):
        extraction_inputs.append(source_text)
        return extract_facts(source_text=source_text, schema_hint=schema_hint)

    monkeypatch.setattr(client, "extract_facts", capture_extract_facts)
    monkeypatch.setattr(webapp, "get_client", lambda _cfg: client)
    c = _client(tmp_path)

    upload = c.post(
        "/sources",
        files={"file": (nfd_filename, nfd_text.encode("utf-8"), "text/plain")},
        follow_redirects=False,
    )
    assert upload.status_code == 303

    store = c.app.state.store

    def extracted_fact():
        facts = [fact for fact in store.review_queue() if fact["subject"] == nfc_subject]
        assert len(facts) == 1

    _wait_for(extracted_fact)
    source = store.get_source_by_path(f"sources/{nfc_filename}")
    assert source is not None
    artifact = store.latest_source_text_artifact(int(source["id"]))
    assert artifact is not None
    assert (tmp_path / artifact["path"]).read_text(encoding="utf-8") == nfc_text
    assert extraction_inputs == [nfc_text]

    fact = next(fact for fact in store.review_queue() if fact["subject"] == nfc_subject)
    assert (fact["subject"], fact["relation"], fact["object"]) == (
        nfc_subject,
        nfc_relation,
        nfc_object,
    )
    assert store.get_fact_terms(int(fact["id"])) == (
        StringLit(nfc_subject),
        StringLit(nfc_relation),
        StringLit(nfc_object),
    )

    accepted = c.post(f"/facts/{fact['id']}/accept", follow_redirects=False)
    assert accepted.status_code == 200
    assert store.get_fact(int(fact["id"]))["status"] == "confirmed"

    answers = []
    for question in (nfd_question, nfc_question):
        response = c.post("/ask", data={"question": question})
        assert response.status_code == 200
        body = unescape(response.text)
        assert "VERIFIED — engine" in body
        match = re.search(
            r'<div class="answer-box">\s*<pre>(.*?)</pre>', body, flags=re.DOTALL
        )
        assert match is not None
        answers.append(match.group(1).strip())

    assert answers[0] == answers[1]
    assert answers[0].startswith(f"{nfc_subject}, {nfc_relation}, {nfc_object}")


def test_upload_rejects_unsupported_type(tmp_path):
    c = _client(tmp_path)
    r = c.post("/sources", files={"file": ("blob.bin", b"\x00\x01", "application/octet-stream")})
    assert r.status_code == 400
    assert "unsupported source type" in r.text


def test_sources_page_lists_sources(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    sid = store.add_source("sources/a.txt", kind="text")
    long_artifact_path = (
        "artifacts/sources/1/"
        "very-long-extracted-text-artifact-name-that-should-not-stretch-the-sources-table.txt"
    )
    store.add_source_artifact(
        source_id=sid,
        kind="extracted_text",
        path=long_artifact_path,
    )
    store.add_fact("A", "is_a", "B", status="candidate", source_id=sid)
    job_id = store.create_extraction_job(
        source_id=sid,
        provider="ollama",
        model="qwen3.5:9b",
        total_chunks=2,
    )
    chunks = store.add_source_chunks(job_id=job_id, source_id=sid, chunks=["a", "b"])
    store.mark_extraction_job_running(job_id)
    store.mark_chunk_running(chunks[0])
    store.mark_chunk_done(chunks[0])
    store.mark_chunk_running(chunks[1])
    store.mark_chunk_failed(chunks[1], "provider down")

    r = c.get("/sources")

    assert r.status_code == 200
    assert "sources/a.txt" in r.text
    assert 'title="sources/a.txt"' in r.text
    assert "text" in r.text
    assert "failed" in r.text
    assert "1/2 chunk(s)" in r.text
    assert "provider down" in r.text
    assert "1 awaiting review" in r.text
    assert "1 unsupported" in r.text
    assert "ollama" in r.text
    assert "qwen3.5:9b" in r.text
    assert "Retry" in r.text
    assert "extracted_text" in r.text
    assert f'title="{long_artifact_path}"' in r.text
    assert 'class="truncate"' in r.text
    assert "Accept all" in r.text


def test_sources_page_labels_the_engine_count_by_what_it_means(tmp_path, monkeypatch):
    """`engine_count` is the engine-input tier, not the `confirmed` status.

    The badge used to read "N confirmed" while the number already counted
    `accepted` too — and, since the tier is read at call time, it counts whatever
    ENGINE_STATUSES says. Labelling it after one member status misreads the KB.
    """
    c = _client(tmp_path)
    store = c.app.state.store
    sid = store.add_source("sources/a.txt", kind="text")
    store.add_fact("A", "is_a", "B", status="accepted", source_id=sid)
    store.add_fact("C", "is_a", "D", status="superseded", source_id=sid)

    body = c.get("/sources").text
    assert "1 engine input" in body
    assert "confirmed" not in body

    monkeypatch.setattr(
        store_db, "ENGINE_STATUSES", store_db.ENGINE_STATUSES | {"superseded"}
    )
    assert "2 engine input" in c.get("/sources").text


def test_sources_accept_all_promotes_review_facts_for_that_source(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    source_id = store.add_source("sources/sample.txt", kind="text")
    other_source_id = store.add_source("sources/other.txt", kind="text")
    candidate_id = store.add_fact(
        "Sample Subject",
        "uses",
        "Sample Object",
        status="candidate",
        source_id=source_id,
    )
    needs_review_id = store.add_fact(
        "Sample Subject",
        "mentions",
        "Sample Note",
        status="needs_review",
        source_id=source_id,
    )
    confirmed_id = store.add_fact(
        "Already Confirmed",
        "is_a",
        "Sample",
        status="confirmed",
        source_id=source_id,
    )
    other_id = store.add_fact(
        "Other Subject",
        "uses",
        "Other Object",
        status="candidate",
        source_id=other_source_id,
    )

    r = c.post(f"/sources/{source_id}/accept-all", follow_redirects=False)

    assert r.status_code == 303
    assert r.headers["location"] == "/sources"
    assert store.get_fact(candidate_id)["status"] == "confirmed"
    assert store.get_fact(needs_review_id)["status"] == "confirmed"
    assert store.get_fact(confirmed_id)["status"] == "confirmed"
    assert store.get_fact(other_id)["status"] == "candidate"
    assert [event["action"] for event in store.fact_log(candidate_id)] == ["accepted"]
    assert [event["action"] for event in store.fact_log(needs_review_id)] == ["accepted"]
    assert store.fact_log(confirmed_id) == []
    assert store.fact_log(other_id) == []

    body = c.get("/sources").text
    assert "0 awaiting review" in body


def test_sources_page_shows_trust_counts_and_evidence_snippets(tmp_path, monkeypatch):
    c = _client(tmp_path)
    store = c.app.state.store
    policy = tmp_path / "policy"
    policy.mkdir()
    (policy / "logic-policy.dl").write_text(
        '.decl functional(rel: symbol)\nfunctional("published_year").\n',
        encoding="utf-8",
    )
    source_id = store.add_source("sources/sample-source.txt", kind="text")
    artifact_id = store.add_source_artifact(
        source_id=source_id,
        kind="original_text",
        path="sources/sample-source.txt",
    )
    job_id = store.create_extraction_job(
        source_id=source_id,
        artifact_id=artifact_id,
        provider="fake",
        model="sample-model",
        total_chunks=1,
    )
    chunk_id = store.add_source_chunks(
        job_id=job_id,
        source_id=source_id,
        chunks=["Sample Report was published in 2024."],
    )[0]
    store.mark_extraction_job_running(job_id)
    store.mark_chunk_running(chunk_id)
    store.mark_chunk_done(chunk_id)
    fact_id = store.add_fact(
        "Sample Report",
        "published_year",
        "2024",
        status="candidate",
        source_id=source_id,
        job_id=job_id,
    )
    store.add_fact_evidence(
        fact_id=fact_id,
        source_id=source_id,
        artifact_id=artifact_id,
        job_id=job_id,
        chunk_id=chunk_id,
        snippet="Sample Report was published in 2024.",
    )
    support_a = store.add_source("sources/sample-support-a.txt")
    support_b = store.add_source("sources/sample-support-b.txt")
    conflict = store.add_source("sources/sample-conflict.txt")
    store.add_fact(
        "Sample Report",
        "published_year",
        "2024",
        status="accepted",
        source_id=support_a,
    )
    store.add_fact(
        "Sample Report",
        "published_year",
        "2024",
        status="confirmed",
        source_id=support_b,
    )
    store.add_fact(
        "Sample Report",
        "published_year",
        "2025",
        status="accepted",
        source_id=conflict,
    )
    monkeypatch.setattr(
        webapp,
        "fact_trust_summary",
        lambda *args, **kwargs: pytest.fail("sources page should use source rollups"),
    )

    body = unescape(c.get("/sources").text)

    assert "sources/sample-source.txt" in body
    assert "original_text" in body
    assert "sample-model" in body
    assert "chunk size" in body
    assert "max facts" in body
    assert "0 unsupported" in body
    assert "1 conflicted" in body
    assert "1 corroborated" in body
    assert "Sample Report was published in 2024." in body


def test_delete_source_removes_file_and_extracted_facts(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    source_path = tmp_path / "sources" / "a.txt"
    source_path.parent.mkdir()
    source_path.write_text("source body", encoding="utf-8")
    sid = store.add_source("sources/a.txt", kind="text")
    artifact_path = tmp_path / "artifacts" / "sources" / str(sid) / "extracted.txt"
    artifact_path.parent.mkdir(parents=True)
    artifact_path.write_text("artifact body", encoding="utf-8")
    store.add_source_artifact(
        source_id=sid,
        kind="extracted_text",
        path=f"artifacts/sources/{sid}/extracted.txt",
    )
    source_fact = store.add_fact("A", "is_a", "B", status="candidate", source_id=sid)
    unrelated_fact = c.fact_id

    r = c.post(f"/sources/{sid}/delete", follow_redirects=False)

    assert r.status_code == 303
    assert r.headers["location"] == "/sources"
    assert not source_path.exists()
    assert not artifact_path.exists()
    assert store.sources() == []
    assert store.get_fact(source_fact) is None
    assert store.get_fact_terms(source_fact) is None
    assert store.get_fact(unrelated_fact) is not None
    assert store.source_extraction_jobs() == []


def test_reanalyze_source_reuses_artifact_and_replaces_extracted_facts(
    tmp_path, monkeypatch, fake_client
):
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("New", "is_a", "Fact", 0.9)]),
    )
    c = _client(tmp_path)
    store = c.app.state.store
    sid = store.add_source("sources/a.txt", kind="text")
    artifact_path = tmp_path / "sources" / "a.txt"
    artifact_path.parent.mkdir()
    artifact_path.write_text("source body", encoding="utf-8")
    artifact_id = store.add_source_artifact(
        source_id=sid,
        kind="original_text",
        path="sources/a.txt",
    )
    store.add_fact("Old", "is_a", "Fact", status="candidate", source_id=sid)
    old_job = store.create_extraction_job(
        source_id=sid,
        artifact_id=artifact_id,
        provider="fake",
        model="old",
        total_chunks=1,
    )
    old_chunk = store.add_source_chunks(
        job_id=old_job, source_id=sid, chunks=["old body"]
    )[0]
    store.mark_extraction_job_running(old_job)
    store.mark_chunk_running(old_chunk)
    store.mark_chunk_done(old_chunk)

    body = c.get("/sources").text
    assert "Re-analyze" in body

    r = c.post(f"/sources/{sid}/reanalyze", follow_redirects=False)

    assert r.status_code == 303
    assert r.headers["location"] == "/sources"

    def reanalyzed():
        assert any(row["subject"] == "New" for row in store.review_queue())
        assert "Analysis complete: 1/1 chunk(s)" in c.get("/sources").text

    _wait_for(reanalyzed)
    assert not any(row["subject"] == "Old" for row in store.review_queue())
    assert len(store.source_extraction_jobs()) == 1
    assert store.get_source(sid)["path"] == "sources/a.txt"
    assert artifact_path.read_text(encoding="utf-8") == "source body"
    fact = [row for row in store.review_queue() if row["subject"] == "New"][0]
    assert fact["source_id"] == sid
    assert store.get_extraction_job_detail(fact["job_id"])["artifact_id"] == artifact_id


def test_reanalyze_preserves_human_decisions_and_suppresses_resurrection(
    tmp_path, monkeypatch, fake_client
):
    # #339 through the real route: reanalyze must not delete a human's confirmed
    # fact nor resurrect a superseded one. The fake run re-emits the superseded
    # triple, so reconcile_fact finds the preserved row and suppresses the
    # re-insert (#160). The confirmed/superseded facts carry NO artifact-anchored
    # evidence, isolating this from the #329 staleness sweep (see the next test).
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client(
            [
                ExtractedFact("New", "is_a", "Fact", 0.9),
                ExtractedFact("Ghost", "is_a", "Spirit", 0.9),
            ]
        ),
    )
    c = _client(tmp_path)
    store = c.app.state.store
    sid = store.add_source("sources/a.txt", kind="text")
    artifact_path = tmp_path / "sources" / "a.txt"
    artifact_path.parent.mkdir()
    artifact_path.write_text("source body", encoding="utf-8")
    store.add_source_artifact(source_id=sid, kind="original_text", path="sources/a.txt")
    store.add_fact("Old", "is_a", "Thing", status="candidate", source_id=sid)
    confirmed = store.add_fact(
        "Kept", "is_a", "Confirmed", status="confirmed", source_id=sid
    )
    superseded = store.add_fact(
        "Ghost", "is_a", "Spirit", status="candidate", source_id=sid
    )
    store.reject_fact(superseded)
    assert store.get_fact(superseded)["status"] == "superseded"

    r = c.post(f"/sources/{sid}/reanalyze", follow_redirects=False)
    assert r.status_code == 303

    def reanalyzed():
        assert any(row["subject"] == "New" for row in store.review_queue())
        assert "Analysis complete: 1/1 chunk(s)" in c.get("/sources").text

    _wait_for(reanalyzed)

    # The review-tier candidate is gone; both human decisions are untouched.
    assert not any(row["subject"] == "Old" for row in store.review_queue())
    assert store.get_fact(confirmed)["status"] == "confirmed"
    assert store.get_fact(superseded)["status"] == "superseded"
    # Exactly one row for the superseded triple: the preserved rejection, never a
    # fresh candidate resurrected from the re-emitted triple.
    ghost_rows = [
        f
        for f in store.facts()
        if (f["subject"], f["relation"], f["object"]) == ("Ghost", "is_a", "Spirit")
    ]
    assert len(ghost_rows) == 1
    assert ghost_rows[0]["id"] == superseded
    assert ghost_rows[0]["status"] == "superseded"


def test_reanalyze_demotes_a_now_unsupported_confirmed_fact_via_the_sweep(
    tmp_path, monkeypatch, fake_client
):
    # #339 composition proof: preserving the confirmed row is all it takes for the
    # #329 post-extraction sweep to demote it when the current artifact no longer
    # supports it -- NO demotion logic lives in clear_source_analysis or the route.
    # The confirmed fact is anchored at an OLDER artifact; reanalyze runs over the
    # newer one and re-extracts a different triple, so the sweep finds no evidence
    # for the confirmed fact at the current artifact and returns it to review,
    # stale=1 (never deleted, never silently left confirmed).
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("Fresh", "is_a", "Fact", 0.9)]),
    )
    c = _client(tmp_path)
    store = c.app.state.store
    sid = store.add_source("sources/a.txt", kind="text")
    artifact_path = tmp_path / "sources" / "a.txt"
    artifact_path.parent.mkdir()
    artifact_path.write_text("current body", encoding="utf-8")
    old_artifact = store.add_source_artifact(
        source_id=sid, kind="original_text", path="sources/a-v1.txt", checksum="v1"
    )
    new_artifact = store.add_source_artifact(
        source_id=sid, kind="original_text", path="sources/a.txt", checksum="v2"
    )
    assert new_artifact != old_artifact
    confirmed = store.add_fact(
        "Ada", "is_a", "Analyst", status="confirmed", source_id=sid
    )
    store.add_fact_evidence(
        fact_id=confirmed, source_id=sid, artifact_id=old_artifact, snippet="Ada"
    )

    r = c.post(f"/sources/{sid}/reanalyze", follow_redirects=False)
    assert r.status_code == 303

    # The sweep runs AFTER the job is marked done, so wait on the demotion itself
    # rather than the "Analysis complete" message.
    def swept():
        assert any(row["subject"] == "Fresh" for row in store.review_queue())
        assert store.get_fact(confirmed)["status"] == "needs_review"

    _wait_for(swept)

    row = store.get_fact(confirmed)
    assert row is not None
    assert row["status"] == "needs_review"
    assert row["stale"] == 1


def test_upload_docx_converts_and_extracts(tmp_path, monkeypatch, fake_client):
    import io

    import docx

    monkeypatch.setattr(
        webapp, "get_client", lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)])
    )
    d = docx.Document()
    d.add_paragraph("converted text")
    buf = io.BytesIO()
    d.save(buf)

    c = _client(tmp_path)
    r = c.post(
        "/sources",
        files={
            "file": (
                "report.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        follow_redirects=False,
    )
    assert r.status_code == 303
    # the original file is preserved and converted text is tracked as an artifact.
    assert (tmp_path / "sources" / "report.docx").is_file()
    artifact_files = list((tmp_path / "artifacts" / "sources" / "1").glob("*.txt"))
    assert len(artifact_files) == 1
    assert artifact_files[0].read_text().strip() == "converted text"
    kinds = {s["kind"] for s in c.app.state.store.sources_with_counts()}
    assert "binary" in kinds
    sources_body = c.get("/sources").text
    assert "sources/report.docx" in sources_body
    assert "artifacts/sources/1/" in sources_body
    assert "extracted_text" in sources_body

    def extracted():
        assert any(row["subject"] == "X" for row in c.app.state.store.review_queue())
        assert "complete" in c.get("/sources").text

    _wait_for(extracted)
    fact = [row for row in c.app.state.store.review_queue() if row["subject"] == "X"][0]
    provenance = c.get(f"/facts/{fact['id']}/provenance").text
    assert "sources/report.docx" in provenance
    assert "artifacts/sources/1/" in provenance


def test_upload_surfaces_llm_error(tmp_path, monkeypatch, fake_client):
    monkeypatch.setattr(
        webapp, "get_client", lambda cfg: fake_client(error=LLMError("provider down"))
    )
    c = _client(tmp_path)
    r = c.post(
        "/sources",
        files={"file": ("note.txt", b"x", "text/plain")},
        follow_redirects=False,
    )
    assert r.status_code == 303
    assert r.headers["location"] == "/sources"

    def failed():
        body = c.get("/sources").text
        assert "Analysis failed: 1 chunk(s) failed" in body
        assert "provider down" in body

    _wait_for(failed)


def test_retry_failed_source_chunks(tmp_path, monkeypatch, fake_client):
    state = {"error": LLMError("provider down")}

    def client_factory(cfg):
        if state["error"] is not None:
            return fake_client(error=state["error"])
        return fake_client([ExtractedFact("X", "is_a", "Y", 0.9)])

    # Record how every worker is dispatched, so we can prove the retry button drives
    # the atomic claim in human-override retry mode rather than a standalone reset.
    real_process = webapp.process_extraction_job
    dispatched = []

    def recording_process(*args, retry=False, retry_max_attempts=None, **kwargs):
        dispatched.append({"retry": retry, "retry_max_attempts": retry_max_attempts})
        return real_process(
            *args, retry=retry, retry_max_attempts=retry_max_attempts, **kwargs
        )

    monkeypatch.setattr(webapp, "process_extraction_job", recording_process)
    monkeypatch.setattr(webapp, "get_client", client_factory)
    c = _client(tmp_path)
    upload = c.post(
        "/sources",
        files={"file": ("note.txt", b"some text", "text/plain")},
        follow_redirects=False,
    )
    assert upload.status_code == 303

    def failed():
        assert "provider down" in c.get("/sources").text

    _wait_for(failed)
    job_id = c.app.state.store.source_extraction_jobs()[0]["id"]
    state["error"] = None

    retry = c.post(f"/sources/jobs/{job_id}/retry", follow_redirects=False)

    assert retry.status_code == 303

    def retried():
        assert "is_a" in c.get("/review").text
        assert "Analysis complete: 1/1 chunk(s)" in c.get("/sources").text

    _wait_for(retried)

    # The button started the worker through the atomic claim in retry mode with the
    # uncapped human override (`retry_max_attempts=None`) — the upload's own worker
    # ran in the ordinary non-retry mode — and there is no standalone reset function
    # left in the store for the handler to have called instead.
    assert [call for call in dispatched if call["retry"]] == [
        {"retry": True, "retry_max_attempts": None}
    ]
    assert not hasattr(Store, "retry_failed_chunks")


def test_create_app_resumes_pending_source_jobs(tmp_path, monkeypatch, fake_client):
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="anthropic",
        model="m",
        api_key=None,
        base_url=None,
    )
    with Store(cfg.db_path) as store:
        store.init_schema()
        sid = store.add_source("sources/a.txt")
        job_id = store.create_extraction_job(
            source_id=sid, provider="anthropic", model="m", total_chunks=1
        )
        store.add_source_chunks(job_id=job_id, source_id=sid, chunks=["some text"])
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )

    c = TestClient(create_app(cfg))

    def resumed():
        assert "is_a" in c.get("/review").text
        assert "Analysis complete: 1/1 chunk(s)" in c.get("/sources").text

    _wait_for(resumed)
    # a `pending` job is resumed directly; only a genuinely-`running` one is rolled
    # back first, so no rollback event should appear for a pending resume (#242).
    assert "extraction_job_rolled_back" not in _job_event_types(cfg, job_id)


def _job_kb(tmp_path, *, with_policy: bool):
    """A KB with one pending extraction job — optionally with a *recorded* policy."""
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="anthropic",
        model="m",
        api_key=None,
        base_url=None,
    )
    policy = tmp_path / POLICY_RELPATH
    with Store(cfg.db_path) as store:
        store.init_schema()
        sid = store.add_source("sources/a.txt")
        job_id = store.create_extraction_job(
            source_id=sid, provider="anthropic", model="m", total_chunks=1
        )
        store.add_source_chunks(job_id=job_id, source_id=sid, chunks=["some text"])
        policy.parent.mkdir(parents=True, exist_ok=True)
        policy.write_text(DEFAULT_POLICY, encoding="utf-8")
        store.record_policy_marker(policy_sha256(DEFAULT_POLICY), origin="scaffold")
    if not with_policy:
        policy.unlink()  # the KB recorded a policy and the file is gone: halted
    return cfg, job_id, policy


def _job_row(cfg, job_id):
    with Store(cfg.db_path) as store:
        store.init_schema()
        return dict(store.get_extraction_job(job_id))


def _job_event_types(cfg, job_id):
    with Store(cfg.db_path) as store:
        store.init_schema()
        return [
            row["event_type"]
            for row in store._conn.execute(
                "SELECT event_type FROM fact_events WHERE job_id = ? ORDER BY id",
                (job_id,),
            )
        ]


def _empty_policy_web_kb(tmp_path, *, text=""):
    """A KB that recorded a real policy whose file has since been truncated to
    empty. Returns (cfg, fact_id): the fact is reviewable, so a POST acting on it
    is a genuine write the halt guard must refuse (#171)."""
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="anthropic",
        model="m",
        api_key=None,
        base_url=None,
    )
    policy = tmp_path / POLICY_RELPATH
    with Store(cfg.db_path) as store:
        store.init_schema()
        fact_id = store.add_fact("A", "is_a", "B", status="needs_review", confidence=0.9)
        policy.parent.mkdir(parents=True, exist_ok=True)
        policy.write_text(DEFAULT_POLICY, encoding="utf-8")
        store.record_policy_marker(policy_sha256(DEFAULT_POLICY), origin="scaffold")
    policy.write_text(text, encoding="utf-8")  # truncated to empty under the KB
    return cfg, fact_id


def test_web_mutating_route_returns_409_on_an_empty_policy_kb(tmp_path):
    """#171: an empty-policy KB is halted for writes on the web too. A mutating
    route (accepting a fact) is refused by the existing middleware guard with the
    existing 409 halted page — no new template — because PolicyEmptyError is a
    PolicyMissingError that guard already catches."""
    cfg, fact_id = _empty_policy_web_kb(tmp_path)
    c = TestClient(create_app(cfg))

    r = c.post(f"/facts/{fact_id}/accept")

    assert r.status_code == 409
    assert "empty" in r.text
    assert "policy reset --force" in r.text


def test_web_report_renders_the_empty_policy_diagnosis(tmp_path):
    """/report is exempt from the write guard and calls verify() directly, so it
    renders the new empty-policy diagnosis instead of the cryptic engine error."""
    cfg, _ = _empty_policy_web_kb(tmp_path)
    c = TestClient(create_app(cfg))

    r = c.get("/report")

    assert r.status_code == 200
    assert "empty" in r.text
    assert "must declare relation/3" not in r.text


def test_launching_the_ui_on_a_halted_kb_resumes_nothing(tmp_path, monkeypatch, fake_client):
    """Zero HTTP requests, and still the launcher used to write to a halted KB (#194).

    `_resume_source_extraction_jobs` runs inside `create_app()`, outside the
    request middleware that guards every route — so the middleware's halt never
    covered it. The worker it started raised, `except Exception` "helpfully"
    marked the job `failed`, and a KB whose rules were gone took a write from a
    process that had not served a single request.
    """
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=False)
    clients = []
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: clients.append(cfg) or fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )
    before = _job_row(cfg, job_id)

    create_app(cfg)

    time.sleep(0.2)  # a worker thread, had one been started, would have written by now
    assert clients == []  # no worker was started at all
    assert _job_row(cfg, job_id) == before  # job untouched: still pending, same message
    assert before["status"] == "pending"
    assert _job_event_types(cfg, job_id) == []  # no started / failed / rolled_back


def test_launching_the_ui_on_a_healthy_kb_still_resumes(tmp_path, monkeypatch, fake_client):
    """The gate above refuses a *halted* KB, not every KB."""
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)
    clients = []
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: clients.append(cfg) or fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )

    c = TestClient(create_app(cfg))

    def resumed():
        assert "is_a" in c.get("/review").text
        assert _job_row(cfg, job_id)["status"] == "done"

    _wait_for(resumed)
    assert clients != []  # the worker really was started


def test_resume_polls_to_done_not_just_to_fact_visibility(tmp_path, monkeypatch, fake_client):
    """Widen the fact-visible -> job-done gap; the resume guard must still reach `done`.

    Facts land in /review inside `_extract_chunk`; the job flips to `done` a step
    later in `mark_chunk_done`. Delay that flip and a guard that stopped at fact
    visibility would read `running`. This pins the guard to the real finish signal.
    """
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )
    real_mark = Store.mark_chunk_done

    def slow_mark(self, chunk_id):
        time.sleep(0.3)
        return real_mark(self, chunk_id)

    monkeypatch.setattr(Store, "mark_chunk_done", slow_mark)

    c = TestClient(create_app(cfg))

    def resumed():
        assert "is_a" in c.get("/review").text
        assert _job_row(cfg, job_id)["status"] == "done"

    _wait_for(resumed, timeout=3.0)  # > the injected 0.3s delay
    assert _job_row(cfg, job_id)["status"] == "done"


def test_worker_halt_does_not_mark_the_job_failed(tmp_path, monkeypatch, fake_client):
    """`except PolicyMissingError` must stay ABOVE `except Exception` in the worker.

    Below it, the generic handler "reports" the halt by calling
    `fail_extraction_job` — a write to the very KB the halt exists to protect, and
    one that buries the job in a `failed` state nothing resumes.
    `process_extraction_job` already rewound the job to `pending`; the right move
    is to touch the DB not at all.
    """
    cfg, job_id, policy = _job_kb(tmp_path, with_policy=True)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )
    called = threading.Event()

    def halt(*args, **kwargs):
        # the shape of a mid-job halt: the job has already been rolled back to
        # `pending` by `process_extraction_job` before the error reaches the worker
        policy.unlink()
        called.set()
        raise PolicyMissingError("policy file is missing; run `verinote policy reset --force`")

    monkeypatch.setattr(webapp, "process_extraction_job", halt)

    create_app(cfg)

    assert called.wait(timeout=2.0)
    time.sleep(0.2)  # let a (wrong) `fail_extraction_job` land, if the order regressed
    job = _job_row(cfg, job_id)
    assert job["status"] == "pending", "the halted job was buried in `failed` by a write"
    assert "analysis failed" not in job["message"]
    assert "extraction_job_failed" not in _job_event_types(cfg, job_id)


def test_worker_still_fails_the_job_on_an_ordinary_error(tmp_path, monkeypatch, fake_client):
    """The control for the test above: a non-halt crash *is* still reported."""
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )

    def boom(*args, **kwargs):
        raise RuntimeError("boom")

    monkeypatch.setattr(webapp, "process_extraction_job", boom)

    create_app(cfg)

    def failed():
        assert _job_row(cfg, job_id)["status"] == "failed"

    _wait_for(failed)
    # Exact equality, not `"boom" in message`: the weaker form does not pin the
    # no-op property of `_error_cause` (#551) — that a message-bearing exception
    # passes through unchanged rather than being type-qualified.
    assert _job_row(cfg, job_id)["message"] == "analysis failed: boom"


def test_worker_names_the_type_when_the_generic_error_has_no_message(
    tmp_path, monkeypatch, fake_client
):
    """#551: a bare `ValueError()` must not leave the job row reading "analysis
    failed: " with nothing after the colon. `_error_cause` names the exception's
    type when `str(exc)` is blank."""
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )

    def boom(*args, **kwargs):
        raise ValueError()

    monkeypatch.setattr(webapp, "process_extraction_job", boom)

    create_app(cfg)

    def failed():
        assert _job_row(cfg, job_id)["status"] == "failed"

    _wait_for(failed)
    assert _job_row(cfg, job_id)["message"] == "analysis failed: ValueError"


def test_worker_names_the_type_when_the_llm_error_has_no_message(
    tmp_path, monkeypatch, fake_client
):
    """#551: the `except LLMError` clause has the same blank-cause symptom as the
    generic one — a bare `LLMError("")` must not leave "extraction failed: "."""
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)

    def raise_blank_llm_error(cfg):
        raise LLMError("")

    monkeypatch.setattr(webapp, "get_client", raise_blank_llm_error)

    create_app(cfg)

    def failed():
        assert _job_row(cfg, job_id)["status"] == "failed"

    _wait_for(failed)
    assert _job_row(cfg, job_id)["message"] == "extraction failed: LLMError"


def test_worker_leaves_a_message_bearing_error_unbounded_on_the_job_row(
    tmp_path, monkeypatch, fake_client
):
    """Truncation stays out of `_error_cause`. #551's own scope declaration
    ("범위 밖" / out of scope in the issue) names message-length bounding as
    something neither of the two job-level clauses it discusses — the CLI's
    `fail_extraction_job` call and the web worker's generic clause — has,
    and something this issue does not add: it covers cause notation only.
    The extraction worker's two direct sites (S1/S2) are unbounded for that
    reason, unlike `_short_error` and S6's own inline normalisation, which
    both bound their reason to 240 characters. A 300-character message must
    survive on the job row whole."""
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )
    long_message = "x" * 300

    def boom(*args, **kwargs):
        raise RuntimeError(long_message)

    monkeypatch.setattr(webapp, "process_extraction_job", boom)

    create_app(cfg)

    def failed():
        assert _job_row(cfg, job_id)["status"] == "failed"

    _wait_for(failed)
    assert _job_row(cfg, job_id)["message"] == f"analysis failed: {long_message}"


def test_worker_busy_does_not_mark_the_job_failed(tmp_path, monkeypatch, fake_client):
    """A job another worker owns must not be buried in `failed` by this worker (#240).

    `except ExtractionJobBusyError` must sit ABOVE `except Exception`: below it, the
    generic handler calls `fail_extraction_job` — a write that corrupts a job this
    process does not own. The right move is to touch the DB not at all.
    """
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )
    called = threading.Event()

    def busy(*args, **kwargs):
        called.set()
        raise ExtractionJobBusyError(job_id)

    monkeypatch.setattr(webapp, "process_extraction_job", busy)

    create_app(cfg)

    assert called.wait(timeout=2.0)
    time.sleep(0.2)  # let a (wrong) `fail_extraction_job` land, if the branch were missing
    job = _job_row(cfg, job_id)
    assert job["status"] == "pending", "a job owned by another worker was buried in `failed`"
    assert "analysis failed" not in job["message"]
    assert "extraction failed" not in job["message"]
    assert "extraction_job_failed" not in _job_event_types(cfg, job_id)


# --- #525: the worker's terminal clauses re-read before they write ------------


def _join_worker(job_id, *, timeout: float = 5.0) -> None:
    """Wait for THE worker thread, not for a clock.

    `_wait_for(status == "done")` is useless as a settle point here: on the broken
    tree `mark_chunk_done` has ALREADY written `done` before `finish_extraction_job`
    runs, so it returns immediately and the whole verdict falls back onto whatever
    sleep follows it. `_start_source_extraction` names its thread, so the test can
    join the actual worker and know the handler has run — or find it already gone,
    which is the same guarantee.
    """
    name = f"verinote-source-extract-{job_id}"
    for thread in threading.enumerate():
        if thread.name == name:
            thread.join(timeout=timeout)
            assert not thread.is_alive(), "the worker thread did not finish"


def _fail_job_spy(monkeypatch):
    """Count `fail_extraction_job` calls, so "declined" is a POSITIVE assertion.

    `_join_worker` already rules out the "the write has not landed yet" reading of a
    still-`done` row, so this is not about timing. It distinguishes a guard that
    DECLINED from one that wrote a `failed` row some later step happened to restore:
    the row and the event tell you the end state, and only the call count tells you
    the write never happened.
    """
    calls = []
    real = store_db.Store.fail_extraction_job

    def spy(self, job_id, message):
        calls.append((job_id, message))
        return real(self, job_id, message)

    monkeypatch.setattr(store_db.Store, "fail_extraction_job", spy)
    return calls


def _raise_on_finish(monkeypatch, exc):
    """Make `finish_extraction_job` raise AFTER `mark_chunk_done` wrote `done`.

    This is the reachable shape of #525 and needs no invention: the chunk loop marks
    the last chunk `done` (which writes the JOB `done` too), and only then does
    `process_extraction_job` call `finish_extraction_job`. Anything raised there
    escapes with the job already complete.
    """

    def boom(self, job_id):
        raise exc

    monkeypatch.setattr(store_db.Store, "finish_extraction_job", boom)


def test_worker_leaves_a_done_job_done_when_finishing_it_raises(
    tmp_path, monkeypatch, fake_client, caplog
):
    """The regression #525 is about: a completed job buried by a later failure.

    Measured on the tree before the fix, driving the real worker: the job came to
    rest `failed` with 'analysis failed: post-done store error', its chunk `done`,
    and an `extraction_job_failed` event appended beside it. Every chunk had
    completed and its candidate facts were committed — the KB reporting a run it
    finished as one that failed (#194/#239).

    The event assertion is not redundant with the status one: the row could read
    `done` while `fail_extraction_job` had appended the event anyway.
    """
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )
    failures = _fail_job_spy(monkeypatch)
    _raise_on_finish(monkeypatch, RuntimeError("post-done store error"))

    with caplog.at_level(logging.WARNING, logger="verinote.web.app"):
        create_app(cfg)
        _join_worker(job_id)

    job = _job_row(cfg, job_id)
    assert job["status"] == "done", "a completed job was buried by a post-`done` error"
    # the real completion message, not merely the absence of a failure one: a guard
    # that declined the write but cleared the message would pass the weaker check.
    assert job["message"] == "Analysis complete: 1/1 chunk(s)"
    assert "extraction_job_failed" not in _job_event_types(cfg, job_id)
    assert failures == [], "the terminal clause wrote over a job it no longer owned"
    # DECLINING IS NOT DROPPING. Nothing else records this error — the job keeps the
    # `done` row `mark_chunk_done` wrote, and `finish_extraction_job` raised before
    # its own completion event and before the run summary — so if the guard did not
    # log, a real sqlite/WAL-class failure would leave no trace anywhere at all.
    assert "not recording on the job row" in caplog.text
    assert "post-done store error" in caplog.text
    # `exc_info` specifically, not just the message: the row would be type-qualified
    # only when the message is blank (`_error_cause`, #551), so a bare `ValueError()`
    # here would read "analysis failed: ValueError" and the traceback would still
    # be the only thing naming anything more specific than that. Asserting the
    # message text alone does not pin it — that string is in the message too.
    assert "RuntimeError" in caplog.text
    declined = next(
        r for r in caplog.records if "not recording on the job row" in r.getMessage()
    )
    assert declined.exc_info is not None
    # nothing else recorded this run either, which is why the log line above is the
    # whole of the record: `finish_extraction_job` raised before its own event.
    assert "extraction_job_completed" not in _job_event_types(cfg, job_id)
    with Store(cfg.db_path) as store:
        store.init_schema()
        assert [r["status"] for r in store.source_chunks(job_id)] == ["done"]


def test_worker_leaves_a_done_job_done_when_finishing_it_raises_an_llm_error(
    tmp_path, monkeypatch, fake_client, caplog
):
    """The `except LLMError` half of the same guard. CHARACTERISATION, not a scenario.

    No production path raises an `LLMError` with the job already `done`: the chunk
    loop swallows `LLMError` through `_release_claimed_chunk` and continues, the two
    calls after `process_extraction_job` returns carry their own guards, and
    `finish_extraction_job` raises no `LLMError` of its own. The exception is
    injected here. What the test pins is that the guard on that clause is the same
    guard — "does this call still own this job?" is not a question the exception
    TYPE answers — so the two clauses cannot drift apart.

    The clause ITSELF is live, and this test says nothing against that: an
    unreadable prompt override reaches it through `_extraction_schema_hint` (#539),
    and `process_extraction_job` raises `LLMError` for a missing job or source. On
    all of those the job is still `pending` and the guard correctly permits the
    write — which is what `test_a_broken_extraction_limit_hint_is_extraction_failed
    _not_analysis_failed` pins.
    """
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )
    failures = _fail_job_spy(monkeypatch)
    _raise_on_finish(monkeypatch, LLMError("post-done llm error"))

    with caplog.at_level(logging.WARNING, logger="verinote.web.app"):
        create_app(cfg)
        _join_worker(job_id)

    job = _job_row(cfg, job_id)
    assert job["status"] == "done"
    assert job["message"] == "Analysis complete: 1/1 chunk(s)"
    assert "extraction_job_failed" not in _job_event_types(cfg, job_id)
    assert failures == []
    assert "not recording on the job row" in caplog.text
    assert "post-done llm error" in caplog.text
    assert "LLMError" in caplog.text
    declined = next(
        r for r in caplog.records if "not recording on the job row" in r.getMessage()
    )
    assert declined.exc_info is not None
    assert "extraction_job_completed" not in _job_event_types(cfg, job_id)


def test_worker_writes_nothing_when_the_job_row_is_gone(tmp_path, monkeypatch, fake_client):
    """The `is not None` half of the guard, which nothing else reaches.

    The docstring's own bullet says a deleted source reads back as `None` and the
    write is then a no-op. Nothing measured that: with the conjunct removed the
    guard raises `TypeError` on `None["status"]` before reaching the write, so the
    decline and the no-op write are indistinguishable to every other test here.
    """
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )
    failures = _fail_job_spy(monkeypatch)

    def delete_the_source_then_raise(self, job_id):
        # `PRAGMA foreign_keys = ON` is set per connection (`store/db.py`) and
        # `extraction_jobs.source_id` is `ON DELETE CASCADE` (`schema.sql`), so the
        # job row goes with the source and the guard's re-read sees `None`.
        self._conn.execute(
            "DELETE FROM sources WHERE id = "
            "(SELECT source_id FROM extraction_jobs WHERE id = ?)",
            (job_id,),
        )
        raise RuntimeError("post-done store error")

    monkeypatch.setattr(
        store_db.Store, "finish_extraction_job", delete_the_source_then_raise
    )

    create_app(cfg)
    _join_worker(job_id)

    # FIRST, for the diagnosis rather than against a vacuous pass: this test cannot
    # go green with the cascade not firing — the job row would survive as `done`,
    # the guard would decline, and the write assertion below would fail anyway.
    # MEASURED, both `PRAGMA foreign_keys` sites flipped to OFF: without these three
    # lines it fails at the write assertion with a bare `[] != [(1, ...)]`, naming
    # nothing; with them it fails here instead, saying which half broke.
    with Store(cfg.db_path) as s:
        s.init_schema()
        assert s.get_extraction_job(job_id) is None
    # The guard fell THROUGH to the write rather than raising on `None["status"]`.
    # This is the assertion the conjunct owns, and it observes the CALL, nothing more.
    # What the write then did is a separate fact, read off the store rather than
    # measured here: `fail_extraction_job` is `UPDATE ... WHERE id = ?` (`store/db.py`)
    # over a row the assertion above already read back as gone — that read happens
    # after `_join_worker`, so after the write — and its event append sits behind
    # `if after is not None`. Do not "strengthen" this with `_job_event_types`: TWO
    # separate things empty it, and neither is about the guard. `fact_events.job_id`
    # is `ON DELETE SET NULL` (`schema.sql`), so the events this run really did write
    # drop out of a `WHERE job_id = ?` query; and `fail_extraction_job` appends
    # nothing anyway once the row is gone. Probed with the write allowed: the table
    # holds [('extraction_job_started', None), ('candidate_created', None)] and zero
    # `extraction_job_failed` rows counted table-wide.
    assert failures == [(job_id, "analysis failed: post-done store error")]


def test_worker_still_fails_a_claimed_job_whose_chunk_crashed(
    tmp_path, monkeypatch, fake_client
):
    """The other side of the predicate: `running` must still be written `failed`.

    THIS IS THE TEST THAT CONSTRAINS THE GUARD FROM BELOW, and the reason the
    predicate is "refuse `done`" rather than anything broader. Every other test
    around it reaches the terminal clauses with the job `pending` (stubbed before
    the claim), `done` (finished, then broken), or gone altogether
    (`test_worker_writes_nothing_when_the_job_row_is_gone`, whose source is deleted
    from inside the raising callable). None drives the ordinary case: a
    job this worker genuinely claimed, whose chunk then failed for a non-`LLMError`
    reason. `process_extraction_job` releases the claim and re-raises, and
    `_refresh_extraction_job` deliberately keeps an owned job `running` across that
    (#337), so the exception arrives with the job `running` and MUST be recorded.

    Widen the refusal set to `{"done", "running"}` and every other test in this file
    still passes while the failure the clause exists to report is silently swallowed.
    This one goes red — measured, and it is the only red in the whole suite.

    OF THE THREE WAYS OF GETTING THIS PREDICATE WRONG NAMED HERE, THE WIDENING IS
    THE ONE THIS TEST ISOLATES — the sentence used to say "or invert the predicate"
    as if that were a second one. Measured, it is not: swapping
    in `cmd_sync`'s `running` predicate leaves THIS test green (a `running` job is
    exactly what it writes) and reddens the pre-claim tests instead, while inverting
    to `!= "done"` reddens this test, both `done` tests and the pre-claim ones all
    together. Either would be caught by something; only the widening is caught by
    this test and nothing else.

    No stub stands in for `process_extraction_job`: the real one runs, so the status
    the clause sees is the one production would produce.
    """
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client(error=RuntimeError("chunk exploded")),
    )

    create_app(cfg)
    _join_worker(job_id)

    job = _job_row(cfg, job_id)
    assert job["status"] == "failed", "a claimed job whose chunk crashed was not recorded"
    assert "chunk exploded" in job["message"], "the cause was dropped from the job row"
    assert "extraction_job_failed" in _job_event_types(cfg, job_id)


def test_startup_revives_a_job_left_running_by_a_crash(tmp_path, monkeypatch, fake_client):
    """A job a crash left `running` (with a `running` chunk) is revived to `done` (#242).

    This is the ONE path that seeds a `running` job at startup — every other resume
    test seeds `pending`. DB state cannot tell this zombie from a live owner, so the
    resume loop rolls it back to `pending` first, which is what lets the new claim
    take it and actually re-process the stranded chunk (rather than walk past it).
    """
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)  # a pending job + recorded policy
    with Store(cfg.db_path) as store:
        store.init_schema()
        store.mark_extraction_job_running(job_id)
        chunk_id = store.source_chunks(job_id)[0]["id"]
        store.mark_chunk_running(chunk_id)
    assert _job_row(cfg, job_id)["status"] == "running"
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )

    c = TestClient(create_app(cfg))

    def revived():
        assert "is_a" in c.get("/review").text
        assert _job_row(cfg, job_id)["status"] == "done"

    _wait_for(revived)
    # the stranded chunk was actually processed, not skipped, and the revival went
    # through a rollback (only a genuinely-`running` job is rolled back first).
    with Store(cfg.db_path) as store:
        store.init_schema()
        assert store.source_chunks(job_id)[0]["status"] == "done"
    assert "extraction_job_rolled_back" in _job_event_types(cfg, job_id)


def _auto_accept_kb(tmp_path):
    """A KB one finished job away from auto-accepting a fact.

    `sources/b.txt` already contributed `X is_a Y` from a `done` job, so the moment
    the pending job on `sources/a.txt` extracts the same triple, both facts have two
    distinct corroborating sources and become auto-accept eligible.
    """
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="anthropic",
        model="m",
        api_key=None,
        base_url=None,
        auto_accept_recommendations=True,
    )
    policy = tmp_path / POLICY_RELPATH
    with Store(cfg.db_path) as store:
        store.init_schema()
        corroborating = store.add_source("sources/b.txt")
        done_job = store.create_extraction_job(
            source_id=corroborating, provider="anthropic", model="m", total_chunks=1
        )
        chunk = store.add_source_chunks(
            job_id=done_job, source_id=corroborating, chunks=["prior text"]
        )[0]
        store.mark_extraction_job_running(done_job)
        store.mark_chunk_running(chunk)
        store.mark_chunk_done(chunk)
        store.finish_extraction_job(done_job)
        store.add_fact(
            "X", "is_a", "Y", status="candidate", source_id=corroborating, job_id=done_job
        )

        sid = store.add_source("sources/a.txt")
        job_id = store.create_extraction_job(
            source_id=sid, provider="anthropic", model="m", total_chunks=1
        )
        store.add_source_chunks(job_id=job_id, source_id=sid, chunks=["some text"])
        policy.parent.mkdir(parents=True, exist_ok=True)
        policy.write_text(DEFAULT_POLICY, encoding="utf-8")
        store.record_policy_marker(policy_sha256(DEFAULT_POLICY), origin="scaffold")
    return cfg, job_id, policy


def _fact_statuses(cfg):
    with Store(cfg.db_path) as store:
        store.init_schema()
        return sorted(str(fact["status"]) for fact in store.facts())


def test_worker_does_not_auto_accept_on_a_kb_that_went_halted(
    tmp_path, monkeypatch, fake_client
):
    """End-to-end: a policy lost after the last chunk must not still auto-accept (#194).

    `apply_auto_accept_recommendations` fires in the same `with Store(...)` block the
    instant `process_extraction_job` returns, so a policy deleted after the final
    chunk's write boundary lands squarely on it — and `status='accepted'` on a KB
    with no rules is a review gate no rule was ever applied to.

    Scope, honestly: this path is double-guarded, so it does not pin
    `apply_auto_accept_recommendations`'s own `assert_writable` — the engine's
    `store_functional_relations` -> `load_policy` would also raise here.
    `test_auto_accept_refuses_a_halted_kb_by_its_own_guard` is the test that pins the
    guard; this one pins the worker's end-to-end behaviour: job `done`, nothing
    accepted, no `auto_accept_applied` in the history.
    """
    cfg, job_id, policy = _auto_accept_kb(tmp_path)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )
    real = webapp.process_extraction_job

    def vanish_after_the_job(*args, **kwargs):
        # the exact gap the guard covers: every chunk passed its write boundary, the
        # job is `done`, and the policy disappears before auto-accept runs
        result = real(*args, **kwargs)
        policy.unlink()
        return result

    monkeypatch.setattr(webapp, "process_extraction_job", vanish_after_the_job)

    create_app(cfg)

    def job_finished():
        assert _job_row(cfg, job_id)["status"] == "done"

    _wait_for(job_finished)
    time.sleep(0.2)  # let an unguarded auto-accept land, if the guard regressed
    # both facts were auto-accept *eligible*; on a halted KB neither may be accepted
    assert _fact_statuses(cfg) == ["candidate", "candidate"]
    with Store(cfg.db_path) as store:
        store.init_schema()
        events = [
            row["event_type"]
            for row in store._conn.execute("SELECT event_type FROM fact_events")
        ]
    assert "auto_accept_applied" not in events


def test_worker_still_auto_accepts_on_a_healthy_kb(tmp_path, monkeypatch, fake_client):
    """The control: the guard refuses a *halted* KB, it does not disable auto-accept.

    Without this, deleting the auto-accept call outright would pass the test above.
    """
    cfg, job_id, _ = _auto_accept_kb(tmp_path)  # policy left in place
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )

    create_app(cfg)

    def auto_accepted():
        assert _job_row(cfg, job_id)["status"] == "done"
        assert _fact_statuses(cfg) == ["accepted", "accepted"]

    _wait_for(auto_accepted)


def test_worker_leaves_a_done_job_done_when_auto_accept_raises(
    tmp_path, monkeypatch, fake_client
):
    # #340 exception-safety, the sibling of the #329 sweep guard directly above it:
    # auto-accept is a call inside the worker's try/except, so it USED TO BE that
    # without the local guard an auto-accept error reached the outer `except
    # Exception -> fail_extraction_job` and retroactively flipped an already-`done`
    # extraction to `failed` — the KB lying about its own run state (#194/#239).
    # Since #525 the outer clause re-reads and declines a `done` job, so this test
    # no longer distinguishes the local guard: MEASURED, removing that guard leaves
    # it green. What it still pins is that the job SURVIVES an auto-accept error,
    # whichever layer contains it — the extraction genuinely succeeded and its
    # facts are already committed. `test_worker_leaves_a_done_job_done_when_
    # finishing_it_raises` is what now pins the outer re-read itself.
    cfg, job_id, _ = _auto_accept_kb(tmp_path)  # auto-accept on, policy present
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )

    def boom(*args, **kwargs):
        raise RuntimeError("auto-accept exploded")

    # Patch the module-level function as bound in app.py's namespace (auto-accept is
    # a module-level function reference here, not a Store method like the sweep).
    monkeypatch.setattr(webapp, "apply_auto_accept_recommendations", boom)

    create_app(cfg)  # resumes the pending job -> _start_source_extraction

    def job_finished():
        assert _job_row(cfg, job_id)["status"] == "done"

    _wait_for(job_finished)
    time.sleep(0.2)  # let a late fail_extraction_job land, if the guard regressed
    row = _job_row(cfg, job_id)
    assert row["status"] == "done"  # the auto-accept error did NOT flip it to failed
    assert row["failed_chunks"] == 0
    assert "analysis failed" not in (row["message"] or "")  # never marked failed
    # extraction genuinely completed; only the (now contained) auto-accept failed.
    with Store(cfg.db_path) as store:
        store.init_schema()
        assert any(f["subject"] == "X" for f in store.facts())


def test_worker_lets_a_policy_missing_halt_from_auto_accept_reach_the_halt_handler(
    tmp_path, monkeypatch, fake_client, caplog
):
    # A DIFFERENT test from the one above: a `PolicyMissingError` out of auto-accept
    # is a #194 halt, not an ordinary failure. It must reach the worker's outer
    # `except PolicyMissingError` handler (which writes NOTHING), NOT be swallowed by
    # the local `except Exception` guard. Job status alone can't tell the two apart
    # (`done` either way — this test passes on today's unguarded code too), so it
    # pins the distinction on the LOG: the halt handler's line must appear and the
    # local guard's must not. That is what catches a guard missing the
    # `except PolicyMissingError: raise` clause, or ordering it below `except
    # Exception`.
    cfg, job_id, _ = _auto_accept_kb(tmp_path)  # auto-accept on, policy present
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )

    def halt(*args, **kwargs):
        raise PolicyMissingError("policy vanished after the job finished")

    monkeypatch.setattr(webapp, "apply_auto_accept_recommendations", halt)

    with caplog.at_level(logging.WARNING, logger="verinote.web.app"):
        create_app(cfg)  # resumes the pending job -> _start_source_extraction

        def halt_logged():
            assert "halted (KB policy missing)" in caplog.text

        _wait_for(halt_logged)  # the halt reached the outer PolicyMissingError handler
        # ...and the local `except Exception` guard did NOT swallow it as a failure.
        assert "auto-accept failed" not in caplog.text
    # the outer handler wrote nothing: the completed job is left exactly `done`.
    assert _job_row(cfg, job_id)["status"] == "done"


def test_report_ok_for_consistent_kb(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    store.add_fact("Org", "established_on", "2020", status="confirmed")
    r = c.get("/report")
    assert r.status_code == 200
    assert "errors: 0" in r.text
    assert "backend: DuckDB" in r.text


def test_report_shows_query_trace_links_and_candidate_exclusion(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    source_id = store.add_source("sources/sample.txt")
    fact_id = store.add_fact(
        "Sample Person",
        "born_in",
        "Sample City",
        status="confirmed",
        source_id=source_id,
    )
    store.add_fact_evidence(
        fact_id=fact_id,
        source_id=source_id,
        snippet="Sample Person was born in Sample City.",
    )
    store.add_fact(
        "Candidate Person",
        "born_in",
        "Draft City",
        status="candidate",
        source_id=source_id,
    )
    path = query_path(tmp_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        '.decl answer_q1(value: symbol)\n'
        'answer_q1(O) :- relation("Sample Person", "born_in", O).\n',
        encoding="utf-8",
    )

    body = unescape(c.get("/report").text)

    assert "Traceability" in body
    # one needs_review fact from _client, one candidate above: both review statuses
    # count, and the report must name which is which.
    assert "2 fact(s) awaiting review were excluded from engine input" in body
    assert "candidate 1, needs_review 1" in body
    assert f'href="/facts/{fact_id}/provenance"' in body
    assert "Sample Person was born in Sample City." in body
    assert "Draft City" not in body


def test_report_traceability_renders_one_status_then_none(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store

    # _client seeds a single needs_review fact: one status, so no separator.
    body = unescape(c.get("/report").text)
    assert "1 fact(s) awaiting review were excluded from engine input" in body
    assert "(needs_review 1)" in body

    store.accept_fact(c.fact_id)

    body = unescape(c.get("/report").text)
    assert "No facts were held back from engine input pending review." in body
    assert "awaiting review" not in body


def test_report_flags_contradiction(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    store.add_fact("Org", "established_on", "2020", status="confirmed")
    store.add_fact("Org", "established_on", "2021", status="confirmed")
    r = c.get("/report")
    assert r.status_code == 200
    assert "ERRORS" in r.text
    assert "functional_conflict" in r.text


def test_report_does_not_claim_a_promotion_gate(tmp_path):
    """The errors banner must not resurrect the false "promotion/query gated" claim.

    No gate exists (see #164): promotion writes never consult the report, and this
    page renders its own answers unconditionally. The banner may say the KB is
    inconsistent, but not that anything is blocked.
    """
    c = _client(tmp_path)
    store = c.app.state.store
    store.add_fact("Org", "established_on", "2020", status="confirmed")
    store.add_fact("Org", "established_on", "2021", status="confirmed")

    r = c.get("/report")

    assert r.status_code == 200
    assert "ERRORS" in r.text
    # The old banner falsely claimed promotion and this page's own answers were
    # gated; both are untrue.
    assert "stay gated" not in r.text
    assert "promotion/query" not in r.text
    # The honest banner is present and says promotion is not blocked.
    assert "Promotion is not blocked" in r.text


def test_promotion_succeeds_while_errors_exist(tmp_path):
    """The issue's own repro, pinned as intended behavior (#164).

    A human's accept is a review-tier decision that writes unconditionally with
    respect to the consistency report — an unrelated error must never block it,
    and the promoted fact must reach engine input.
    """
    c = _client(tmp_path)
    store = c.app.state.store
    # A functional conflict on `established_on` (both confirmed): errors > 0.
    store.add_fact("Org", "established_on", "2020", status="confirmed")
    store.add_fact("Org", "established_on", "2021", status="confirmed")
    assert "ERRORS" in c.get("/report").text
    unrelated = store.add_fact("Widget", "is_a", "Gadget", status="needs_review")

    r = c.post(f"/facts/{unrelated}/accept")

    assert r.status_code == 200
    assert store.get_fact(unrelated)["status"] == "confirmed"
    assert unrelated in {int(row["id"]) for row in store.engine_fact_terms()}
    # The error is unchanged — promotion neither cleared nor worsened it.
    assert "ERRORS" in c.get("/report").text


def test_report_shows_answers_even_while_errors_exist(tmp_path):
    """The "query gated" half of the old claim was false: answers still render.

    `verify()` returns findings and answers together and the template renders
    `rep.answers` with no error guard, so an answer and an error coexist on the
    page — the banner must not say otherwise.
    """
    c = _client(tmp_path)
    store = c.app.state.store
    source_id = store.add_source("sources/sample.txt")
    store.add_fact("Ada", "born_in", "London", status="confirmed", source_id=source_id)
    path = query_path(tmp_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        ".decl answer_q1(value: symbol)\n"
        'answer_q1(O) :- relation("Ada", "born_in", O).\n',
        encoding="utf-8",
    )
    # An unrelated functional conflict drives errors > 0 on the same page.
    store.add_fact("Org", "established_on", "2020", status="confirmed")
    store.add_fact("Org", "established_on", "2021", status="confirmed")

    body = unescape(c.get("/report").text)

    assert "ERRORS" in body
    assert "London" in body


def test_ask_still_verifies_a_clean_answer_while_errors_exist(tmp_path, monkeypatch):
    """Ask does not gate its VERIFIED label on the KB consistency report (#164).

    Ask certifies from the answer query's own run (policy is the base relation
    decl), not from `/report`'s error count, so an unrelated functional_conflict
    never downgrades a clean deterministic answer. This pins that the report
    banner must NOT claim Ask withholds VERIFIED while errors stand — it does not.
    If Ask ever starts gating on report errors, this fails and the banner copy
    must be revisited in lockstep.
    """
    class DeterministicOnly:
        name = "deterministic-only"

        def extract_query_intent(self, *, question: str, schema_hint: str = ""):
            raise AssertionError("deterministic question must bypass LLM")

        def translate_query(self, *, question: str, qid: int, schema_hint: str = ""):
            raise AssertionError("Ask must not call direct Datalog fallback")

        def answer_question(self, *, question: str, context: str):
            raise AssertionError("verified engine answer must not call fallback")

    monkeypatch.setattr(webapp, "get_client", lambda cfg: DeterministicOnly())
    c = _client(tmp_path)
    store = c.app.state.store
    source_id = store.add_source("sources/sample.txt")
    store.add_fact("샘플인물", "역할", "검토자", status="confirmed", source_id=source_id)
    # An unrelated functional conflict makes /report show errors > 0.
    store.add_fact("Org", "established_on", "2020", status="confirmed")
    store.add_fact("Org", "established_on", "2021", status="confirmed")
    assert "ERRORS" in c.get("/report").text

    r = c.post("/ask", data={"question": "샘플인물의 역할은 무엇인가?"})

    assert r.status_code == 200
    body = unescape(r.text)
    assert "VERIFIED — engine" in body
    assert "검토자" in body


def test_report_shows_missing_duckdb_message(tmp_path, monkeypatch):
    real_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name == "duckdb":
            raise ImportError("missing")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="anthropic",
        model="m",
        api_key=None,
        base_url=None,
    )
    c = TestClient(create_app(cfg))
    r = c.get("/report")
    assert r.status_code == 200
    assert "DuckDB verification backend is not available" in r.text
    assert "DuckDB is not installed" in r.text


def test_report_surfaces_invalid_query_file(tmp_path):
    c = _client(tmp_path)
    path = query_path(tmp_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(".decl answer_q1(value: symbol)\nanswer_q1(O) :- bogus(O).\n", encoding="utf-8")

    r = c.get("/report")
    assert r.status_code == 200
    assert "ERRORS" in r.text
    assert "bogus" in r.text


def test_report_surfaces_invalid_relation_alias_policy(tmp_path):
    c = _client(tmp_path)
    policy = tmp_path / "policy"
    policy.mkdir()
    (policy / "relation-aliases.md").write_text("- `role` -> `role`\n", encoding="utf-8")
    path = query_path(tmp_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        '.decl answer_q1(value: symbol)\n'
        'answer_q1(O) :- relation("Sample Person", "role", O).\n',
        encoding="utf-8",
    )

    r = c.get("/report")

    assert r.status_code == 200
    assert "ERRORS" in r.text
    assert "policy error" in r.text
    assert "self-map" in r.text


def test_questions_surfaces_invalid_relation_alias_policy(tmp_path):
    c = _client(tmp_path)
    c.app.state.store.add_question("What is the sample role?")
    policy = tmp_path / "policy"
    policy.mkdir()
    (policy / "relation-aliases.md").write_text("- `role` -> `role`\n", encoding="utf-8")
    path = query_path(tmp_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        '.decl answer_q1(value: symbol)\n'
        'answer_q1(O) :- relation("Sample Person", "role", O).\n',
        encoding="utf-8",
    )

    r = c.get("/questions")

    assert r.status_code == 200
    assert "policy error" in r.text
    assert "self-map" in r.text


def test_edit_form_renders(tmp_path):
    c = _client(tmp_path)
    r = c.get(f"/facts/{c.fact_id}/edit")
    assert r.status_code == 200
    assert 'name="subject"' in r.text
    assert 'name="subject_kind"' in r.text
    assert "/amend" in r.text


def test_amend_endpoint_updates_and_audits(tmp_path):
    c = _client(tmp_path)
    r = c.post(
        f"/facts/{c.fact_id}/amend",
        data={
            "subject": "NewSubj",
            "subject_kind": "string",
            "relation": "became",
            "relation_kind": "string",
            "object": "NewObj",
            "object_kind": "string",
            "note": "n",
        },
    )
    assert r.status_code == 200
    assert "NewSubj" in r.text and "NewObj" in r.text
    store = c.app.state.store
    assert store.get_fact(c.fact_id)["subject"] == "NewSubj"
    assert any(e["action"] == "amended" for e in store.fact_log(c.fact_id))


@pytest.mark.parametrize("missing_kind", ["subject_kind", "relation_kind", "object_kind"])
def test_amend_endpoint_rejects_missing_kind_without_writing(tmp_path, missing_kind):
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact(
        structural_term('person("Ada")'),
        structural_term("has_role"),
        structural_term('role(person("Ada"), "PI")'),
        status="needs_review",
        note="original",
    )
    before_row = dict(store.get_fact(fid))
    before_terms = store.get_fact_terms(fid)
    data = {
        "subject": 'person("Ada")',
        "subject_kind": "term",
        "relation": "has_role",
        "relation_kind": "term",
        "object": 'role(person("Ada"), "PI")',
        "object_kind": "term",
        "note": "changed",
    }
    del data[missing_kind]

    r = c.post(f"/facts/{fid}/amend", data=data)

    assert r.status_code == 422
    assert dict(store.get_fact(fid)) == before_row
    assert store.get_fact_terms(fid) == before_terms
    assert store.fact_log(fid) == []


def test_edit_form_preserves_structural_fact_input_kinds(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact(
        structural_term('person("Ada")'),
        structural_term("born_in"),
        "London",
        status="needs_review",
    )

    r = c.get(f"/facts/{fid}/edit")

    assert r.status_code == 200
    assert 'name="subject_kind"' in r.text
    assert '<option value="term" selected>term</option>' in r.text
    assert 'name="object_kind"' in r.text
    assert '<option value="string" selected>string</option>' in r.text


def test_edit_form_uses_duckdb_term_values_not_stale_sqlite_mirrors(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact(
        structural_term('person("Ada")'),
        structural_term("has_role"),
        structural_term('role(person("Ada"), "PI")'),
        status="needs_review",
    )
    store._conn.execute(
        "UPDATE facts SET subject = ?, relation = ?, object = ? WHERE id = ?",
        ("stale_subject", "stale_relation", "stale_object", fid),
    )

    body = unescape(c.get(f"/facts/{fid}/edit").text)

    assert 'value="person("Ada")"' in body
    assert 'value="has_role"' in body
    assert 'value="role(person("Ada"), "PI")"' in body
    assert "stale_subject" not in body


def test_edit_form_uses_raw_values_for_stringlit_inputs(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact(
        'person("Ada")',
        "has_role",
        'role(person("Ada"), "PI")',
        status="needs_review",
    )

    body = unescape(c.get(f"/facts/{fid}/edit").text)

    assert 'value="person("Ada")"' in body
    assert 'value="role(person("Ada"), "PI")"' in body
    assert 'value=""person(' not in body


def test_edit_save_preserves_duckdb_terms_when_sqlite_mirror_is_stale(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact(
        structural_term('person("Ada")'),
        structural_term("has_role"),
        structural_term('role(person("Ada"), "PI")'),
        status="needs_review",
    )
    store._conn.execute(
        "UPDATE facts SET subject = ?, relation = ?, object = ? WHERE id = ?",
        ("stale_subject", "stale_relation", "stale_object", fid),
    )

    r = c.post(
        f"/facts/{fid}/amend",
        data={
            "subject": 'person("Ada")',
            "subject_kind": "term",
            "relation": "has_role",
            "relation_kind": "term",
            "object": 'role(person("Ada"), "PI")',
            "object_kind": "term",
            "note": "",
        },
    )

    assert r.status_code == 200
    assert store.get_fact_terms(fid) == (
        Compound("person", (StringLit("Ada"),)),
        Atom("has_role"),
        Compound("role", (Compound("person", (StringLit("Ada"),)), StringLit("PI"))),
    )


def test_edit_save_preserves_unchanged_stringlit_without_adding_display_quotes(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact(
        'person("Ada")',
        "has_role",
        'role(person("Ada"), "PI")',
        status="needs_review",
    )

    r = c.post(
        f"/facts/{fid}/amend",
        data={
            "subject": 'person("Ada")',
            "subject_kind": "string",
            "relation": "has_role",
            "relation_kind": "string",
            "object": 'role(person("Ada"), "PI")',
            "object_kind": "string",
            "note": "",
        },
    )

    assert r.status_code == 200
    assert store.get_fact_terms(fid) == (
        StringLit('person("Ada")'),
        StringLit("has_role"),
        StringLit('role(person("Ada"), "PI")'),
    )


def test_amend_endpoint_can_save_explicit_structural_terms(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact("A", "r", "B", status="needs_review")

    r = c.post(
        f"/facts/{fid}/amend",
        data={
            "subject": 'person("Ada")',
            "subject_kind": "term",
            "relation": "born_in",
            "relation_kind": "term",
            "object": "London",
            "object_kind": "string",
            "note": "",
        },
    )

    assert r.status_code == 200
    assert store.get_fact_terms(fid) == (
        Compound("person", (StringLit("Ada"),)),
        Atom("born_in"),
        StringLit("London"),
    )


def test_amend_endpoint_saves_term_looking_text_as_stringlit_in_string_mode(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact("A", "r", "B", status="needs_review")

    r = c.post(
        f"/facts/{fid}/amend",
        data={
            "subject": 'person("Ada")',
            "subject_kind": "string",
            "relation": "has_role",
            "relation_kind": "string",
            "object": 'role(person("Ada"), "PI")',
            "object_kind": "string",
            "note": "",
        },
    )

    assert r.status_code == 200
    assert store.get_fact_terms(fid) == (
        StringLit('person("Ada")'),
        StringLit("has_role"),
        StringLit('role(person("Ada"), "PI")'),
    )
    assert 'class="subj term-string">"person(\\"Ada\\")"' in unescape(r.text)


def test_amend_normalizes_an_nfd_string_slot_to_nfc(tmp_path):
    """#200 web boundary: an NFD value posted to the amend form is stored NFC.

    `_fact_input` NFC-normalizes string-kind slots, so the fact is NFC at rest
    and byte-matches an NFC query at the DuckDB engine level. The NFD input is
    proven byte-different from its NFC form first, so the assertion means
    something.
    """
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact("A", "is_a", "B", status="needs_review")
    nfd_subject = unicodedata.normalize("NFD", "café")
    nfc_subject = unicodedata.normalize("NFC", "café")
    assert nfd_subject != nfc_subject

    r = c.post(
        f"/facts/{fid}/amend",
        data={
            "subject": nfd_subject,
            "subject_kind": "string",
            "relation": "is_a",
            "relation_kind": "string",
            "object": "B",
            "object_kind": "string",
            "note": "",
        },
    )

    assert r.status_code == 200
    assert store.get_fact_terms(fid)[0] == StringLit(nfc_subject)
    assert store.get_fact(fid)["subject"] != nfd_subject


def test_amend_normalizes_an_escape_encoded_nfd_term_leaf_to_nfc(tmp_path):
    """#200 web boundary, the discriminating case: an ASCII term source -> NFD leaf.

    The posted object is 100% ASCII (all \\uXXXX escapes) yet DECODES to an NFD
    `StringLit` leaf after parsing. Only `nfc_term` running post-parse on the
    amend path stores it NFC; a pre-parse whole-string `nfc()` would leave the
    escape-encoded leaf in NFD, so this fails on the wrong implementation.
    """
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact("A", "is_a", "B", status="needs_review")
    term_source = 'note("caf\\u0065\\u0301")'
    assert term_source.isascii()
    nfc_cafe = unicodedata.normalize("NFC", "café")
    nfd_cafe = unicodedata.normalize("NFD", "café")
    assert nfc_cafe != nfd_cafe

    r = c.post(
        f"/facts/{fid}/amend",
        data={
            "subject": "A",
            "subject_kind": "string",
            "relation": "is_a",
            "relation_kind": "string",
            "object": term_source,
            "object_kind": "term",
            "note": "",
        },
    )

    assert r.status_code == 200
    obj_term = store.get_fact_terms(fid)[2]
    assert obj_term == Compound("note", (StringLit(nfc_cafe),))
    assert obj_term.args[0].value != nfd_cafe


def _corrupt_sidecar_kb(tmp_path, *, status="needs_review", with_query=False):
    """Build a KB whose facts.duckdb is genuine garbage and hand back a FRESH
    client, so its store lazily opens the corrupt file on first use.

    Seeding through one app then serving through another matters: a store that
    already holds an open DuckDB handle would keep reading the pre-corruption
    file from cache and give a false pass. The fact carries a compound term, so
    the corruption destroys real structural data (not an empty row). Returns
    (client, fact_id, sqlite_term_token).
    """
    cfg = Config(
        root=tmp_path, db_path=tmp_path / "kb.sqlite",
        provider="anthropic", model="m", api_key=None, base_url=None,
    )
    seed = create_app(cfg).state.store
    source_id = seed.add_source("sources/a.txt")
    fact_id = seed.add_fact(
        structural_term('person("Ada")'),
        structural_term("born_in"),
        "London",
        status=status,
        source_id=source_id,
        confidence=0.9,
    )
    token = seed.get_fact(fact_id)["term_token"]
    if with_query:
        qp = query_path(tmp_path)
        qp.parent.mkdir(parents=True, exist_ok=True)
        qp.write_text(
            ".decl answer_q1(value: symbol)\n"
            'answer_q1(O) :- relation("Ada", "born_in", O).\n',
            encoding="utf-8",
        )
    seed.close()
    fact_terms_path(tmp_path).write_bytes(b"not a duckdb database file" * 500)
    return TestClient(create_app(cfg)), fact_id, token


def test_review_halts_loudly_when_the_sidecar_is_unreadable(tmp_path):
    # The dangerous former behavior: `_fact_view`'s `except ValueError` swallowed
    # the corruption and rendered every field as kind="string" with a 200 -- a
    # structural fact made indistinguishable from a genuine string fact. /review
    # is a full-page GET, so the handler renders the halt page inline.
    client, _fid, _token = _corrupt_sidecar_kb(tmp_path)

    r = client.get("/review")

    assert r.status_code == 409
    assert "Fact terms unavailable" in r.text
    assert "born_in" not in r.text  # the fact is never rendered as a plain string


def test_provenance_halts_loudly_when_the_sidecar_is_unreadable(tmp_path):
    client, fid, _token = _corrupt_sidecar_kb(tmp_path)

    r = client.get(f"/facts/{fid}/provenance")

    assert r.status_code == 409
    assert "Fact terms unavailable" in r.text


def test_report_halts_loudly_when_the_sidecar_is_unreadable(tmp_path):
    # `report_trace` used to raise an uncaught 500 here while `verify()` degraded
    # gracefully in the same route -- the two now converge on the halt page.
    client, _fid, _token = _corrupt_sidecar_kb(
        tmp_path, status="confirmed", with_query=True
    )

    r = client.get("/report")

    assert r.status_code == 409
    assert "Fact terms unavailable" in r.text


def test_edit_form_over_htmx_redirects_instead_of_silently_swallowing(tmp_path):
    # The edit form is an htmx partial swap, and htmx will not swap a 4xx/5xx body
    # into the DOM -- an inline halt page would be a *silent* no-op on exactly the
    # path #173 is about. The handler must send HX-Redirect so htmx does a
    # full-page navigation to the halt page instead.
    client, fid, _token = _corrupt_sidecar_kb(tmp_path)

    r = client.get(f"/facts/{fid}/edit", headers={"HX-Request": "true"})

    assert r.status_code == 409
    assert r.headers.get("HX-Redirect") == webapp.FACT_TERMS_UNAVAILABLE_PATH


def test_amend_over_htmx_is_refused_via_redirect_and_writes_nothing(tmp_path):
    # The anti-data-loss guarantee on the real (htmx) save path: without
    # HX-Redirect the swallowed 4xx would leave the user with no halt and no
    # feedback at all. The redirect forces the halt page; nothing is written.
    client, fid, token = _corrupt_sidecar_kb(tmp_path)

    r = client.post(
        f"/facts/{fid}/amend",
        headers={"HX-Request": "true"},
        data={
            "subject": 'person("Ada")',
            "subject_kind": "string",
            "relation": "born_in",
            "relation_kind": "string",
            "object": "London",
            "object_kind": "string",
            "note": "",
        },
    )

    assert r.headers.get("HX-Redirect") == webapp.FACT_TERMS_UNAVAILABLE_PATH
    # get_fact reads SQLite only, so the token is legible through the corruption:
    # an unchanged token proves the amend wrote nothing to either store.
    assert client.app.state.store.get_fact(fid)["term_token"] == token


def test_amend_without_htmx_is_refused_inline_and_writes_nothing(tmp_path):
    # A non-browser client (script/API) sends no HX-Request header, so it gets the
    # inline halt page. The store guard still refuses the write regardless.
    client, fid, token = _corrupt_sidecar_kb(tmp_path)

    r = client.post(
        f"/facts/{fid}/amend",
        data={
            "subject": "A",
            "subject_kind": "string",
            "relation": "r",
            "relation_kind": "string",
            "object": "B",
            "object_kind": "string",
            "note": "",
        },
    )

    assert r.status_code == 409
    assert "Fact terms unavailable" in r.text
    assert client.app.state.store.get_fact(fid)["term_token"] == token


def test_fact_terms_unavailable_page_renders_without_touching_terms(tmp_path):
    # The HX-Redirect target must render even while the term store cannot be read,
    # so the route reads no fact terms of its own.
    client, _fid, _token = _corrupt_sidecar_kb(tmp_path)

    r = client.get(webapp.FACT_TERMS_UNAVAILABLE_PATH)

    assert r.status_code == 409
    assert "Fact terms unavailable" in r.text


@pytest.mark.parametrize(
    ("action", "expected_status"),
    [
        ("toggle", "confirmed"),
        ("accept", "confirmed"),
        ("reject", "superseded"),
    ],
)
def test_committed_decision_over_htmx_carries_saved_notice_to_sidecar_halt(
    tmp_path, action, expected_status
):
    """A corrupt term sidecar cannot hide a status decision SQLite committed."""
    client, fid, _token = _corrupt_sidecar_kb(tmp_path)

    post = client.post(
        f"/facts/{fid}/{action}", headers={"HX-Request": "true"}
    )

    assert post.status_code == 409
    redirect = post.headers["HX-Redirect"]
    assert redirect.startswith(
        webapp.FACT_TERMS_UNAVAILABLE_PATH
        + f"?decision_fact_id={fid}&decision_action={action}&decision_log_id="
    )
    assert client.app.state.store.get_fact(fid)["status"] == expected_status

    halt = client.get(redirect)

    assert halt.status_code == 409
    body = " ".join(halt.text.split())
    assert f"Your {action} decision was already saved in SQLite" in body
    assert "facts.duckdb" in halt.text
    assert "confirm the fact's current status" in body


def test_committed_decision_notice_survives_an_app_restart(tmp_path):
    client, fid, _token = _corrupt_sidecar_kb(tmp_path)

    post = client.post(f"/facts/{fid}/accept", headers={"HX-Request": "true"})

    restarted = TestClient(create_app(client.app.state.cfg))
    halt = restarted.get(post.headers["HX-Redirect"])

    assert halt.status_code == 409
    assert "Your accept decision was already saved in SQLite" in " ".join(
        halt.text.split()
    )


def test_saved_decision_notice_rejects_forged_or_stale_redirect_state(tmp_path):
    client, fid, _token = _corrupt_sidecar_kb(tmp_path)
    post = client.post(f"/facts/{fid}/accept", headers={"HX-Request": "true"})
    redirect = post.headers["HX-Redirect"]

    forged = client.get(redirect.replace("decision_action=accept", "decision_action=reject"))
    assert "decision was already saved in SQLite" not in forged.text

    client.app.state.store.toggle_review(fid)
    stale = client.get(redirect)
    assert "decision was already saved in SQLite" not in stale.text


def test_non_htmx_committed_decision_renders_the_saved_notice_inline(tmp_path):
    client, fid, _token = _corrupt_sidecar_kb(tmp_path)

    halt = client.post(f"/facts/{fid}/accept")

    assert halt.status_code == 409
    assert "Your accept decision was already saved in SQLite" in " ".join(
        halt.text.split()
    )


def test_unchanged_decision_does_not_create_a_saved_notice(tmp_path):
    client, fid, _token = _corrupt_sidecar_kb(tmp_path, status="confirmed")

    post = client.post(f"/facts/{fid}/accept", headers={"HX-Request": "true"})

    assert post.status_code == 409
    assert post.headers["HX-Redirect"] == webapp.FACT_TERMS_UNAVAILABLE_PATH
    halt = client.get(post.headers["HX-Redirect"])
    assert "decision was already saved in SQLite" not in halt.text


def test_generic_sidecar_halt_does_not_claim_a_decision_was_saved(tmp_path):
    client, _fid, _token = _corrupt_sidecar_kb(tmp_path)

    r = client.get(webapp.FACT_TERMS_UNAVAILABLE_PATH)

    assert r.status_code == 409
    assert "decision was already saved in SQLite" not in r.text


def test_review_renders_normally_for_a_string_fact_on_a_healthy_sidecar(tmp_path):
    # False-positive guard: a legitimately string-typed fact must keep rendering
    # as kind="string" with a 200 -- only a genuine raise halts, never the
    # absent/None-terms path.
    c = _client(tmp_path)

    r = c.get("/review")

    assert r.status_code == 200
    assert "Fact terms unavailable" not in r.text
    assert "is_a" in r.text


def test_amend_endpoint_rejects_invalid_structural_terms_without_writing(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact(
        structural_term('person("Ada")'),
        structural_term("born_in"),
        "London",
        status="needs_review",
    )
    before_row = dict(store.get_fact(fid))
    before_terms = store.get_fact_terms(fid)

    r = c.post(
        f"/facts/{fid}/amend",
        data={
            "subject": 'person("Ada"',
            "subject_kind": "term",
            "relation": "born_in",
            "relation_kind": "term",
            "object": "London",
            "object_kind": "string",
            "note": "bad",
        },
    )

    assert r.status_code == 400
    assert "expected" in r.text
    assert dict(store.get_fact(fid)) == before_row
    assert store.get_fact_terms(fid) == before_terms
    assert store.fact_log(fid) == []


def test_amend_endpoint_rejects_nonground_structural_terms_without_writing(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    fid = store.add_fact("A", "r", "B", status="needs_review")
    before_row = dict(store.get_fact(fid))
    before_terms = store.get_fact_terms(fid)

    r = c.post(
        f"/facts/{fid}/amend",
        data={
            "subject": "person(X)",
            "subject_kind": "term",
            "relation": "r",
            "relation_kind": "string",
            "object": "B",
            "object_kind": "string",
            "note": "bad",
        },
    )

    assert r.status_code == 400
    assert "ground" in r.text
    assert dict(store.get_fact(fid)) == before_row
    assert store.get_fact_terms(fid) == before_terms
    assert store.fact_log(fid) == []


def test_provenance_shows_source_and_audit(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    sid = store.add_source("sources/x.txt", kind="text")
    fid = store.add_fact("S", "r", "O", status="needs_review", source_id=sid)
    store.toggle_review(fid)  # leaves an audit entry
    r = c.get(f"/facts/{fid}/provenance")
    assert r.status_code == 200
    assert "sources/x.txt" in r.text
    assert "candidate_created" in r.text
    assert "system" in r.text
    assert "toggled" in r.text
    assert "human" in r.text


def test_provenance_renders_trust_dossier_sections(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    policy = tmp_path / "policy"
    policy.mkdir()
    (policy / "logic-policy.dl").write_text(
        '.decl functional(rel: symbol)\nfunctional("published_year").\n',
        encoding="utf-8",
    )
    source_id = store.add_source("sources/sample-source.txt")
    artifact_id = store.add_source_artifact(
        source_id=source_id,
        kind="original_text",
        path="sources/sample-source.txt",
    )
    job_id = store.create_extraction_job(
        source_id=source_id,
        artifact_id=artifact_id,
        provider="fake",
        model="sample-model",
        total_chunks=1,
    )
    chunk_id = store.add_source_chunks(
        job_id=job_id,
        source_id=source_id,
        chunks=["Sample Report was published in 2024."],
    )[0]
    store.mark_extraction_job_running(job_id)
    store.mark_chunk_running(chunk_id)
    store.mark_chunk_done(chunk_id)
    run_id = store.add_run(provider="fake", model="sample-model", summary="sample run")
    fact_id = store.add_fact(
        "Sample Report",
        "published_year",
        "2024",
        status="needs_review",
        confidence=0.99,
        source_id=source_id,
        run_id=run_id,
        job_id=job_id,
        note="model note",
    )
    store.add_fact_evidence(
        fact_id=fact_id,
        source_id=source_id,
        artifact_id=artifact_id,
        job_id=job_id,
        chunk_id=chunk_id,
        snippet="Sample Report was published in 2024.",
    )
    store.accept_fact(fact_id)
    other_source = store.add_source("sources/sample-conflict.txt")
    store.add_fact(
        "Sample Report",
        "published_year",
        "2025",
        status="accepted",
        source_id=other_source,
    )

    body = unescape(c.get(f"/facts/{fact_id}/provenance").text)

    assert "Trust dossier" in body
    assert "Trust summary" in body
    assert "source backed" in body
    assert "single source" in body
    assert "conflicted" in body
    assert "source support" in body
    assert "sources/sample-source.txt" in body
    assert "Conflict summary" in body
    assert "2024" in body
    assert "2025" in body
    assert "Lifecycle timeline" in body
    assert "candidate_created" in body
    assert "accepted" in body
    assert f"job #{job_id}" in body
    assert "Source evidence" in body
    assert "Sample Report was published in 2024." in body
    assert "Metadata" in body
    assert "model metadata only" in body
    assert body.index("Trust summary") < body.index("Metadata")


def test_provenance_renders_structural_terms_from_duckdb(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    sid = store.add_source("sources/x.txt", kind="text")
    fid = store.add_fact(
        structural_term('person("Ada")'),
        structural_term("has_role"),
        structural_term('role(person("Ada"), "PI")'),
        status="needs_review",
        source_id=sid,
    )

    body = unescape(c.get(f"/facts/{fid}/provenance").text)

    assert 'class="subj term-term">person("Ada")' in body
    assert 'class="rel term-term">has_role' in body
    assert 'class="obj term-term">role(person("Ada"), "PI")' in body
    assert "sources/x.txt" in body


def test_report_renders_compound_fact_input_and_answer(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    store.add_fact(
        structural_term('person("Ada")'),
        structural_term("has_role"),
        structural_term('role(person("Ada"), "PI")'),
        status="confirmed",
    )
    path = query_path(tmp_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        ".decl answer_q1(value: symbol)\n"
        'answer_q1(O) :- relation(person("Ada"), has_role, O).\n',
        encoding="utf-8",
    )

    body = unescape(c.get("/report").text)

    assert 'q1: role(person("Ada"), "PI")' in body
    assert 'relation(person("Ada"), has_role, role(person("Ada"), "PI"))' in body


def test_analytics_page_renders(tmp_path):
    from verinote.store.analytics import duckdb_available

    c = _client(tmp_path)
    c.app.state.store.add_fact("A", "is_a", "B", status="confirmed", confidence=0.95)
    r = c.get("/analytics")
    assert r.status_code == 200
    if duckdb_available():
        assert "By status" in r.text and "confirmed" in r.text
    else:
        assert "DuckDB isn't installed" in r.text


def test_confidence_chart_is_captioned_as_not_a_trust_signal(tmp_path):
    """The confidence buckets are the LLM's opinion of itself, and the page says so.

    Rendered *above* the numbers: a reader who meets the buckets first has already
    decided what they mean, and "lots of 0.9s" reads as "this KB is trustworthy" —
    the exact inference verinote exists to refuse.
    """
    from verinote.store.analytics import duckdb_available

    if not duckdb_available():
        pytest.skip("analytics page is disabled without DuckDB")

    c = _client(tmp_path)
    c.app.state.store.add_fact("A", "is_a", "B", status="confirmed", confidence=0.95)
    body = c.get("/analytics").text

    disclaimer = "verinote never uses it to decide what is true"
    assert disclaimer in body
    # Only the confidence table carries it; the deterministic breakdowns do not.
    assert body.count(disclaimer) == 1
    # heading, then the disclaimer, then the buckets — in that order.
    assert (
        body.index("Confidence distribution")
        < body.index(disclaimer)
        < body.index("0.9–1.0")
    )


def test_add_question_persists(tmp_path):
    c = _client(tmp_path)
    r = c.post("/questions", data={"text": "Where was Ada born?"}, follow_redirects=False)
    assert r.status_code == 303
    assert [q["text"] for q in c.app.state.store.questions()] == ["Where was Ada born?"]


def test_ask_page_renders_and_is_linked(tmp_path):
    c = _client(tmp_path)

    home = c.get("/")
    page = c.get("/ask")

    assert home.status_code == 200
    assert 'href="/ask"' in home.text
    assert page.status_code == 200
    assert "<h1>Ask</h1>" in page.text


def test_ask_post_renders_verified_engine_answer_without_persisting(
    tmp_path, monkeypatch
):
    class DeterministicOnly:
        name = "deterministic-only"

        def extract_query_intent(self, *, question: str, schema_hint: str = ""):
            raise AssertionError("deterministic question must bypass LLM")

        def translate_query(self, *, question: str, qid: int, schema_hint: str = ""):
            raise AssertionError("Ask must not call direct Datalog fallback")

        def answer_question(self, *, question: str, context: str):
            raise AssertionError("verified engine answer must not call fallback")

    monkeypatch.setattr(webapp, "get_client", lambda cfg: DeterministicOnly())
    c = _client(tmp_path)
    store = c.app.state.store
    source_id = store.add_source("sources/sample.txt")
    store.add_fact("샘플인물", "역할", "검토자", status="confirmed", source_id=source_id)

    r = c.post("/ask", data={"question": "샘플인물의 역할은 무엇인가?"})

    assert r.status_code == 200
    body = unescape(r.text)
    assert "VERIFIED — engine" in body
    assert "검토자" in body
    assert body.index("<pre>샘플인물, 역할, 검토자") < body.index(
        "deterministic query matched confirmed/accepted facts"
    )
    assert "Verified source facts" in r.text
    assert "sources/sample.txt" in r.text
    assert store.questions() == []
    assert not query_path(tmp_path).exists()


def test_ask_post_renders_unverified_llm_fallback(tmp_path, monkeypatch):
    class Fallback:
        name = "fallback"

        def extract_query_intent(self, *, question: str, schema_hint: str = ""):
            return parse_query_intent(
                {
                    "kind": "unknown_or_unsupported",
                    "subject": None,
                    "relation": None,
                    "object": None,
                    "relation_candidates": None,
                    "operator": None,
                    "value_type": None,
                    "value": None,
                    "reason": "unsupported synthetic question",
                }
            )

        def translate_query(self, *, question: str, qid: int, schema_hint: str = ""):
            raise AssertionError("Ask fallback must not call direct Datalog")

        def answer_question(self, *, question: str, context: str):
            assert "sources/sample.txt" in context
            return "Synthetic answer from excerpts."

    monkeypatch.setattr(webapp, "get_client", lambda cfg: Fallback())
    c = _client(tmp_path)
    source = tmp_path / "sources" / "sample.txt"
    source.parent.mkdir()
    source.write_text("Sample Entity provides Sample Service.", encoding="utf-8")
    c.app.state.store.add_source("sources/sample.txt")

    r = c.post("/ask", data={"question": "Sample Entity overview"})

    assert r.status_code == 200
    body = unescape(r.text)
    assert "UNVERIFIED — source exploration" in body
    assert "Synthetic answer from excerpts." in body
    assert body.index("<pre>Synthetic answer from excerpts.</pre>") < body.index(
        "unsupported synthetic question"
    )
    assert "sources/sample.txt" in body


def test_delete_question_removes_query_file_entry(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    store.add_fact("Ada", "born_in", "London", status="confirmed")
    qid = store.add_question("Where was Ada born?")
    store.set_question_query(
        qid,
        '.decl answer_q1(value: symbol)\nanswer_q1(O) :- relation("Ada", "born_in", O).',
        "translated",
    )
    query_file = query_path(tmp_path)
    query_file.parent.mkdir(parents=True, exist_ok=True)
    query_file.write_text(store.questions()[0]["query_dl"] + "\n", encoding="utf-8")

    r = c.post(f"/questions/{qid}/delete", follow_redirects=False)

    assert r.status_code == 303
    assert store.questions() == []
    assert query_file.read_text(encoding="utf-8") == ""
    assert "No questions yet" in c.get("/questions").text


def test_translate_and_report_answers(tmp_path, monkeypatch, fake_client, intent_payload):
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client(
            intent=intent_payload(
                "lookup_object", subject="Sample Person", relation="born_in"
            )
        ),
    )
    c = _client(tmp_path)
    store = c.app.state.store
    store.add_fact("Sample Person", "born_in", "Sample Place", status="confirmed")
    store.add_question("Where was Sample Person born?")

    r = c.post("/questions/translate", follow_redirects=False)
    assert r.status_code == 303
    assert store.questions()[0]["status"] == "translated"
    # the report and questions page now surface the engine-evaluated answer
    assert "Sample Place" in c.get("/report").text
    assert "Sample Place" in c.get("/questions").text


def test_questions_translate_relation_discovery_shows_actual_lifecycle_states(
    tmp_path, monkeypatch
):
    from verinote.pipeline.query_candidate_eval import QueryCandidateSetEvaluation
    from verinote.pipeline.query_candidate_eval import QueryCandidateSetOutcome

    def intent_for(question: str):
        if question.startswith("Translated"):
            return _intent(
                "discover_entity_relations", subject="Synthetic Web Entity"
            )
        if question.startswith("Review"):
            return _intent(
                "discover_entity_relations", subject="Synthetic Web Review Entity"
            )
        if question.startswith("Ambiguous"):
            return _intent(
                "discover_entity_relations", subject="Synthetic Web Ambiguous Entity"
            )
        return _intent("discover_entity_relations", subject="Synthetic Web Missing Entity")

    client = IntentOnlyClient(intent_for)
    monkeypatch.setattr(webapp, "get_client", lambda cfg: client)
    from verinote.pipeline.query import evaluate_query_candidate_plan as real_eval

    def no_answer_for_empty_plan(store, plan):
        if plan.reason == "no relation discovery candidates matched the schema":
            return QueryCandidateSetEvaluation(
                plan=plan,
                outcome=QueryCandidateSetOutcome.NO_ANSWER,
            )
        return real_eval(store, plan)

    monkeypatch.setattr(
        "verinote.pipeline.query.evaluate_query_candidate_plan",
        no_answer_for_empty_plan,
    )
    c = _client(tmp_path)
    store = c.app.state.store
    store.add_fact(
        "Synthetic Web Entity",
        "synthetic_web_relation",
        "Synthetic Web Value",
        status="confirmed",
    )
    store.add_fact(
        "Synthetic Web Review Entity",
        "source",
        "Synthetic Review Value",
        status="confirmed",
    )
    store.add_fact(
        "Synthetic Web Ambiguous Entity",
        "subject_relation",
        "Synthetic Subject Value",
        status="confirmed",
    )
    store.add_fact(
        "Synthetic Web Source",
        "object_relation",
        "Synthetic Web Ambiguous Entity",
        status="confirmed",
    )
    store.add_question("Translated relation discovery?")
    store.add_question("Review relation discovery?")
    store.add_question("Ambiguous relation discovery?")
    store.add_question("No answer relation discovery?")

    r = c.post("/questions/translate", follow_redirects=False)

    assert r.status_code == 303
    assert [q["status"] for q in store.questions()] == [
        "translated",
        "review_required",
        "ambiguous",
        "no_answer",
    ]
    body = unescape(c.get("/questions").text)
    assert "Translated" in body
    assert "Review required" in body
    assert "Ambiguous" in body
    assert "No answer" in body
    assert "synthetic_web_relation" in body
    assert "relation label requires review: source" in body
    assert "multiple query candidates returned conflicting answers" in body
    assert "no confirmed facts match" in body
    assert client.direct_datalog_calls == 0
    query_dl = query_path(tmp_path).read_text(encoding="utf-8")
    assert (
        'answer_q1("synthetic_web_relation") :- '
        'relation("Synthetic Web Entity", "synthetic_web_relation", O).'
    ) in query_dl
    assert "review_required" not in query_dl
    assert "ambiguous" not in query_dl
    assert "no_answer" not in query_dl


def test_translate_persists_llm_error_reason(tmp_path, monkeypatch, fake_client):
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client(error=LLMError("provider unavailable")),
    )
    c = _client(tmp_path)
    c.app.state.store.add_question("What is the sample answer?")
    r = c.post("/questions/translate", follow_redirects=False)

    assert r.status_code == 303
    q = c.app.state.store.questions()[0]
    assert q["status"] == "translation_failed"
    assert q["reason"] == "provider unavailable"
    page = c.get("/questions").text
    assert "translation_failed" in page
    assert "provider unavailable" in page


def test_translate_persists_get_client_failure_reason(tmp_path, monkeypatch):
    def raise_client_error(cfg):
        raise LLMError("missing provider credentials")

    monkeypatch.setattr(webapp, "get_client", raise_client_error)
    c = _client(tmp_path)
    c.app.state.store.add_question("What is the sample answer?")
    r = c.post("/questions/translate", follow_redirects=False)

    assert r.status_code == 303
    q = c.app.state.store.questions()[0]
    assert q["status"] == "translation_failed"
    assert q["reason"] == "missing provider credentials"
    assert (tmp_path / "facts" / "query.dl").read_text(encoding="utf-8") == ""
    page = c.get("/questions").text
    assert "translation_failed" in page
    assert "missing provider credentials" in page


def test_translate_leaves_the_reason_blank_when_the_llm_error_has_no_message(
    tmp_path, monkeypatch
):
    """`_fail_pending_translations` (S6) is deliberately OUT of #551's scope: its
    `reason` is a standalone column, not interpolated after a separator, and
    `question_outcome_view` already renders a per-status default sentence when it
    is blank — there is no dangling colon here to fix. Naming the exception's type
    at this site would REPLACE that sentence ("The provider output could not be
    used.") with a bare class name, which is a regression, not an improvement.
    This pins the unchanged behavior so nothing re-routes this site through
    `_error_cause` later."""

    def raise_blank_llm_error(cfg):
        raise LLMError("")

    monkeypatch.setattr(webapp, "get_client", raise_blank_llm_error)
    c = _client(tmp_path)
    c.app.state.store.add_question("What is the sample answer?")
    r = c.post("/questions/translate", follow_redirects=False)

    assert r.status_code == 303
    q = c.app.state.store.questions()[0]
    assert q["status"] == "translation_failed"
    assert q["reason"] == ""
    page = c.get("/questions").text
    assert "The provider output could not be used." in page


def test_translate_collapses_whitespace_and_bounds_the_reason_to_240_chars(
    tmp_path, monkeypatch
):
    """S6's inline normalisation (`" ".join(str(exc).split())[:240]`) is a bare
    expression, not a named helper, since #551 cut this site off from
    `_short_error` — easier to lose a piece of by accident than a helper call
    would be. `questions.reason` is unbounded `TEXT`, written for every pending
    question at once, so both halves need a test: the internal-whitespace
    collapse, and the 240-character bound.

    Message built so the expected result can be computed by direct slicing,
    not by re-deriving the production logic: 100 `a`s, an irregular whitespace
    run that must collapse to one space, then 200 `b`s. Collapsed that is 301
    characters (100 + 1 + 200); truncated to 240 it is 100 `a`s, one space, and
    139 `b`s — one message that exercises both the collapse and the
    truncation boundary.
    """
    message = "a" * 100 + "  \n\t  " + "b" * 200
    collapsed = "a" * 100 + " " + "b" * 200
    assert len(collapsed) == 301
    expected_reason = collapsed[:240]
    assert expected_reason == "a" * 100 + " " + "b" * 139

    def raise_long_llm_error(cfg):
        raise LLMError(message)

    monkeypatch.setattr(webapp, "get_client", raise_long_llm_error)
    c = _client(tmp_path)
    c.app.state.store.add_question("What is the sample answer?")
    r = c.post("/questions/translate", follow_redirects=False)

    assert r.status_code == 303
    q = c.app.state.store.questions()[0]
    assert q["status"] == "translation_failed"
    assert q["reason"] == expected_reason
    assert len(q["reason"]) == 240


def test_translate_shows_invalid_model_output_reason_in_question_row(
    tmp_path, monkeypatch
):
    monkeypatch.setattr(webapp, "get_client", lambda cfg: object())

    def translate(store, client, *, root):
        q = store.questions(pending_only=True)[0]
        reason = "invalid model output: missing answer rule"
        store.set_question_query(
            q["id"], f'review_required("{reason}")', "review_required", reason
        )
        return [{"id": q["id"], "status": "review_required", "reason": reason}]

    monkeypatch.setattr(webapp, "translate_questions", translate)
    c = _client(tmp_path)
    c.app.state.store.add_question("Which synthetic result is available?")

    r = c.post("/questions/translate", follow_redirects=False)

    assert r.status_code == 303
    body = unescape(c.get("/questions").text)
    assert "Review required" in body
    assert "invalid model output: missing answer rule" in body


def test_questions_page_shows_non_executable_reason(tmp_path):
    c = _client(tmp_path)
    qid = c.app.state.store.add_question("Which sample item is current?")
    c.app.state.store.set_question_query(
        qid,
        'ambiguous("multiple sample entities match")',
        "ambiguous",
        "multiple sample entities match",
    )

    r = c.get("/questions")

    assert r.status_code == 200
    assert "ambiguous" in r.text
    assert "multiple sample entities match" in r.text


def test_questions_page_shows_all_visible_outcomes(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    rows = [
        (
            "Translated synthetic question?",
            '.decl answer_q1(value: symbol)\nanswer_q1(O) :- relation("S", "r", O).',
            "translated",
            "",
        ),
        (
            "Review synthetic question?",
            'review_required("unsupported synthetic question")',
            "review_required",
            "unsupported synthetic question",
        ),
        (
            "No synthetic answer?",
            'no_answer("no confirmed facts match")',
            "no_answer",
            "no confirmed facts match",
        ),
        (
            "Failed synthetic translation?",
            None,
            "translation_failed",
            "provider returned invalid schema",
        ),
        (
            "Ambiguous synthetic question?",
            'ambiguous("multiple synthetic candidates matched")',
            "ambiguous",
            "multiple synthetic candidates matched",
        ),
    ]
    for text, query_dl, status, reason in rows:
        qid = store.add_question(text)
        store.set_question_query(qid, query_dl, status, reason)

    body = unescape(c.get("/questions").text)

    assert "Translated" in body
    assert "Review required" in body
    assert "No answer" in body
    assert "Translation failed" in body
    assert "Ambiguous" in body
    assert "badge-question-no-answer" in body
    assert "badge-question-translation-failed" in body
    assert "unsupported synthetic question" in body
    assert "provider returned invalid schema" in body
    assert "multiple synthetic candidates matched" in body
    assert "No engine answers yet" in body
    assert "Check each question outcome above" in body


def test_questions_page_recovers_reason_from_non_executable_query(tmp_path):
    c = _client(tmp_path)
    qid = c.app.state.store.add_question("Which synthetic item matches?")
    c.app.state.store.set_question_query(
        qid,
        'ambiguous("legacy synthetic ambiguity")',
        "ambiguous",
        "",
    )

    body = unescape(c.get("/questions").text)

    assert "legacy synthetic ambiguity" in body


def test_repair_action_accepts_valid_fix(tmp_path, monkeypatch, fake_client, intent_payload):
    import time
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client(
            intent=intent_payload(
                "lookup_object", subject="Sample Person", relation="born_in"
            )
        ),
    )
    c = _client(tmp_path)
    store = c.app.state.store
    store.add_fact("Sample Person", "born_in", "Sample Place", status="confirmed")
    qid = store.add_question("Where was Sample Person born?")
    store.set_question_query(
        qid, 'review_required("Where was Sample Person born?")', "review_required"
    )

    r = c.post("/questions/repair", follow_redirects=False)
    assert r.status_code == 303
    for _ in range(100):
        if store.questions()[0]["status"] == "translated":
            break
        time.sleep(0.01)
    assert store.questions()[0]["status"] == "translated"


def test_repair_post_enqueues_without_constructing_a_client(tmp_path, monkeypatch):
    class ThreadRecorder:
        def __init__(self):
            self.started = []

        def Thread(self, *, target, name, daemon):  # noqa: N802 - threading API
            recorder = self

            class Handle:
                def start(self):
                    recorder.started.append(name)

            return Handle()

    recorder = ThreadRecorder()
    monkeypatch.setattr(webapp, "threading", recorder)
    monkeypatch.setattr(
        webapp, "get_client", lambda cfg: (_ for _ in ()).throw(AssertionError("request called LLM"))
    )
    c = _client(tmp_path)
    store = c.app.state.store
    _id = store.add_question("What is synthetic?")
    store.set_question_query(_id, 'review_required("synthetic")', "review_required")

    response = c.post("/questions/repair", follow_redirects=False)

    assert response.status_code == 303
    assert recorder.started == ["verinote-question-repair-1"]
    assert store.latest_repair_job()["status"] == "pending"


def _repair_kb(tmp_path) -> tuple[TestClient, "Store"]:
    """A KB with one `review_required` question, ready to be repaired."""
    c = _client(tmp_path)
    store = c.app.state.store
    qid = store.add_question("What is synthetic?")
    store.set_question_query(qid, 'review_required("synthetic")', "review_required")
    return c, store


def test_repair_worker_names_the_type_when_the_config_error_has_no_message(
    tmp_path, monkeypatch
):
    """#551, site S3: `except (ConfigCorruptError, CredentialsCorruptError)` in the
    repair worker's `run()` must not leave "repair failed: " with nothing after
    the colon for a blank exception.

    `assert_writable` is called on the worker's OWN `Store`, while every request
    (including this test's own POST) calls it only on `app.state.store` — so the
    patch is keyed on store identity, not on call order, to reach the worker
    without disturbing the route.
    """
    c, store = _repair_kb(tmp_path)
    real_assert_writable = webapp.assert_writable

    def boom(store_arg):
        if store_arg is not c.app.state.store:
            raise CredentialsCorruptError("")
        return real_assert_writable(store_arg)

    monkeypatch.setattr(webapp, "assert_writable", boom)

    response = c.post("/questions/repair", follow_redirects=False)
    assert response.status_code == 303

    def failed():
        job = store.latest_repair_job()
        assert job is not None and job["status"] == "failed"

    _wait_for(failed)
    assert store.latest_repair_job()["message"] == "repair failed: CredentialsCorruptError"


def test_repair_worker_names_the_type_when_the_llm_error_has_no_message(
    tmp_path, monkeypatch
):
    """#551, site S4: `except LLMError` in the repair worker has the same
    blank-cause symptom as the config/credentials clause beside it."""
    c, store = _repair_kb(tmp_path)

    def raise_blank_llm_error(cfg):
        raise LLMError("")

    monkeypatch.setattr(webapp, "get_client", raise_blank_llm_error)

    response = c.post("/questions/repair", follow_redirects=False)
    assert response.status_code == 303

    def failed():
        job = store.latest_repair_job()
        assert job is not None and job["status"] == "failed"

    _wait_for(failed)
    assert store.latest_repair_job()["message"] == "repair failed: LLMError"


def test_repair_worker_names_the_type_when_the_generic_error_has_no_message(
    tmp_path, monkeypatch
):
    """#551, site S5: the repair worker's broad `except Exception` clause already
    composes `_short_error`, which now routes through `_error_cause` — a bare
    `ValueError()` must not leave "repair failed: "."""
    c, store = _repair_kb(tmp_path)

    def boom(*args, **kwargs):
        raise ValueError()

    monkeypatch.setattr(webapp, "process_repair_job", boom)

    response = c.post("/questions/repair", follow_redirects=False)
    assert response.status_code == 303

    def failed():
        job = store.latest_repair_job()
        assert job is not None and job["status"] == "failed"

    _wait_for(failed)
    assert store.latest_repair_job()["message"] == "repair failed: ValueError"


def test_repair_worker_names_the_type_when_the_generic_error_has_only_whitespace(
    tmp_path, monkeypatch
):
    """`.strip()`, not truthiness. `ValueError("   ")` has a truthy `str()`, so
    without `.strip()` in `_error_cause` this still degrades to
    "repair failed: " with a trailing space and no cause."""
    c, store = _repair_kb(tmp_path)

    def boom(*args, **kwargs):
        raise ValueError("   ")

    monkeypatch.setattr(webapp, "process_repair_job", boom)

    response = c.post("/questions/repair", follow_redirects=False)
    assert response.status_code == 303

    def failed():
        job = store.latest_repair_job()
        assert job is not None and job["status"] == "failed"

    _wait_for(failed)
    assert store.latest_repair_job()["message"] == "repair failed: ValueError"


def test_questions_page_shows_live_repair_progress_and_terminal_failure(tmp_path):
    c = _client(tmp_path)
    store = c.app.state.store
    qid = store.add_question("What is synthetic?")
    store.set_question_query(qid, 'review_required("synthetic")', "review_required")
    job, _ = store.enqueue_repair_job(provider="fake", model="m")

    live = c.get("/questions").text
    assert "Repair job #1: pending" in live
    assert 'hx-trigger="every 2s"' in live

    store.fail_pending_repair_job(int(job["id"]), "Repair failed: synthetic provider outage")
    terminal = c.get("/questions").text
    assert "synthetic provider outage" in terminal
    assert 'hx-trigger="every 2s"' not in terminal


def test_questions_progress_counts_real_failed_item(tmp_path):
    from verinote.pipeline.repair import process_repair_job

    class OutageClient:
        def extract_query_intent(self, *, question, schema_hint):
            raise LLMError("synthetic provider outage")

    c = _client(tmp_path)
    store = c.app.state.store
    for text in ("What is synthetic one?", "What is synthetic two?"):
        qid = store.add_question(text)
        store.set_question_query(qid, 'review_required("synthetic")', "review_required")
    job, _ = store.enqueue_repair_job(provider="fake", model="m")
    process_repair_job(store, OutageClient(), job_id=int(job["id"]), root=tmp_path)

    body = c.get("/questions").text
    assert "Repair job #1: failed" in body
    assert "1/2 processed" in body
    assert "synthetic provider outage" in body


def test_settings_page_renders(tmp_path):
    c = _client(tmp_path)
    r = c.get("/settings")
    assert r.status_code == 200
    assert "Provider" in r.text and "Anthropic" in r.text
    assert "ClaudeCLI" in r.text
    assert str(tmp_path) in r.text
    assert 'name="extraction_chunk_chars"' in r.text
    assert 'name="extraction_chunk_overlap_chars"' in r.text
    assert 'name="extraction_max_facts_per_chunk"' in r.text
    assert 'name="auto_accept_recommendations"' in r.text
    assert 'name="relation_aliases_text"' in r.text
    assert 'href="/prompts"' in r.text


def test_settings_saves_app_theme_across_kb_switches(tmp_path, monkeypatch):
    monkeypatch.setenv("HOME", str(tmp_path / "home"))
    monkeypatch.setenv("XDG_CONFIG_HOME", str(tmp_path / "xdg"))
    monkeypatch.setenv("APPDATA", str(tmp_path / "appdata"))
    c = _client(tmp_path)
    other = tmp_path / "other-kb"

    initial = c.get("/settings")
    assert 'data-theme="system"' in initial.text
    assert 'action="/settings/theme"' in initial.text
    assert 'name="theme"' in initial.text

    r = c.post("/settings/theme", data={"theme": "dark"}, follow_redirects=False)

    assert r.status_code == 303
    assert r.headers["location"] == "/settings"
    assert read_app_config()["theme"] == "dark"
    assert 'data-theme="dark"' in c.get("/settings").text

    assert c.post("/settings/root", data={"root": str(other)}).status_code == 200
    assert read_app_config()["theme"] == "dark"
    assert 'data-theme="dark"' in c.get("/settings").text


def test_halted_kb_disables_and_refuses_app_theme_changes(tmp_path, monkeypatch):
    monkeypatch.setenv("HOME", str(tmp_path / "home"))
    monkeypatch.setenv("XDG_CONFIG_HOME", str(tmp_path / "xdg"))
    monkeypatch.setenv("APPDATA", str(tmp_path / "appdata"))
    cfg, _, _ = _job_kb(tmp_path, with_policy=False)
    c = TestClient(create_app(cfg))

    settings = c.get("/settings")

    assert settings.status_code == 200
    assert 'action="/settings/theme"' in settings.text
    assert 'name="theme" disabled' in settings.text
    assert "Theme changes are unavailable while this KB's logic policy is halted." in settings.text

    r = c.post("/settings/theme", data={"theme": "light"}, follow_redirects=False)

    assert r.status_code == 409
    assert "Verification halted" in r.text
    assert read_app_config() == {}


def test_prompts_page_renders_default_prompt(tmp_path):
    c = _client(tmp_path)

    r = c.get("/prompts")

    assert r.status_code == 200
    assert "Prompts" in r.text
    assert "Extraction" in r.text
    assert "semantic subject-predicate-object statement" in r.text
    assert "API keys are not shown" in r.text


def test_prompts_page_switches_prompt_key(tmp_path):
    c = _client(tmp_path)

    r = c.get("/prompts", params={"prompt": "query-intent"})

    assert r.status_code == 200
    assert "Query intent" in r.text
    assert "Classify one natural-language question" in r.text
    assert "semantic subject-predicate-object statement" not in r.text


def test_prompt_save_writes_kb_policy_file(tmp_path):
    c = _client(tmp_path)

    r = c.post(
        "/prompts",
        data={
            "prompt_id": "extraction",
            "prompt_text": "Use only supplied synthetic text.",
        },
        follow_redirects=False,
    )

    assert r.status_code == 303
    assert r.headers["location"] == "/prompts?prompt=extraction"
    path = tmp_path / "policy" / "prompts" / "extraction.md"
    assert path.read_text(encoding="utf-8") == "Use only supplied synthetic text.\n"
    assert "Use only supplied synthetic text." in c.get("/prompts").text


def test_prompt_reset_deletes_override(tmp_path):
    c = _client(tmp_path)
    path = tmp_path / "policy" / "prompts" / "extraction.md"
    path.parent.mkdir(parents=True)
    path.write_text("Custom prompt.\n", encoding="utf-8")

    r = c.post(
        "/prompts/reset",
        data={"prompt_id": "extraction"},
        follow_redirects=False,
    )

    assert r.status_code == 303
    assert not path.exists()
    assert "semantic subject-predicate-object statement" in c.get("/prompts").text


def test_prompt_save_rejects_empty_text(tmp_path):
    c = _client(tmp_path)

    r = c.post(
        "/prompts",
        data={"prompt_id": "extraction", "prompt_text": "   "},
    )

    assert r.status_code == 400
    assert "prompt text is required" in r.text
    assert not (tmp_path / "policy" / "prompts" / "extraction.md").exists()


def test_prompt_save_rejects_missing_required_placeholder(tmp_path):
    c = _client(tmp_path)
    submitted = "Return a query for the supplied question."

    r = c.post(
        "/prompts",
        data={
            "prompt_id": "query-translation",
            "prompt_text": submitted,
        },
    )

    assert r.status_code == 400
    assert "{qid}" in r.text
    assert submitted in r.text
    assert not (tmp_path / "policy" / "prompts" / "query-translation.md").exists()


def test_prompt_routes_reject_unknown_key(tmp_path):
    c = _client(tmp_path)

    assert c.get("/prompts", params={"prompt": "../secret"}).status_code == 400
    r = c.post(
        "/prompts",
        data={"prompt_id": "../secret", "prompt_text": "No."},
    )
    assert r.status_code == 400


def test_prompt_editor_never_renders_api_key(tmp_path):
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="anthropic",
        model="m",
        api_key="supersecret",
        base_url=None,
    )
    client = TestClient(create_app(cfg))

    r = client.get("/prompts")

    assert r.status_code == 200
    assert "supersecret" not in r.text


def test_settings_saves_relation_aliases(tmp_path):
    c = _client(tmp_path)

    initial_body = c.get("/settings").text
    assert "- `제공 요소` -&gt; `provides`" in initial_body

    r = c.post(
        "/settings/relation-aliases",
        data={"relation_aliases_text": "- `title` -> `role`"},
        follow_redirects=False,
    )

    assert r.status_code == 303
    alias_path = tmp_path / "policy" / "relation-aliases.md"
    assert alias_path.read_text(encoding="utf-8") == "- `title` -> `role`\n"
    body = c.get("/settings").text
    assert "- `title` -&gt; `role`" in body
    assert "- `제공 요소` -&gt; `provides`" in body


def test_settings_saves_plain_relation_aliases(tmp_path):
    c = _client(tmp_path)

    r = c.post(
        "/settings/relation-aliases",
        data={"relation_aliases_text": "- role -> 역할"},
        follow_redirects=False,
    )

    assert r.status_code == 303
    alias_path = tmp_path / "policy" / "relation-aliases.md"
    assert alias_path.read_text(encoding="utf-8") == "- role -> 역할\n"


def test_settings_omits_conflicting_default_alias_direction(tmp_path):
    policy = tmp_path / "policy"
    policy.mkdir()
    (policy / "relation-aliases.md").write_text("- `role` -> `역할`\n", encoding="utf-8")
    c = _client(tmp_path)

    body = c.get("/settings").text

    assert "- `role` -&gt; `역할`" in body
    assert "- `역할` -&gt; `role`" not in body


def test_settings_rejects_invalid_relation_aliases(tmp_path):
    c = _client(tmp_path)

    r = c.post(
        "/settings/relation-aliases",
        data={"relation_aliases_text": "- `role` -> `role`"},
    )

    assert r.status_code == 400
    assert "self-map" in r.text
    assert not (tmp_path / "policy" / "relation-aliases.md").exists()


def test_settings_rejects_malformed_relation_aliases(tmp_path):
    c = _client(tmp_path)

    r = c.post(
        "/settings/relation-aliases",
        data={"relation_aliases_text": "- role 역할"},
    )

    assert r.status_code == 400
    assert "expected `raw` -&gt; `canonical`" in r.text
    assert not (tmp_path / "policy" / "relation-aliases.md").exists()


def test_settings_save_changes_active_provider(tmp_path, monkeypatch):
    for var in ("VERINOTE_PROVIDER", "VERINOTE_MODEL", "VERINOTE_BASE_URL"):
        monkeypatch.delenv(var, raising=False)
    c = _client(tmp_path)
    r = c.post(
        "/settings",
        data={
            "provider": "ollama",
            "model": "llama3.1",
            "base_url": "",
            "extraction_chunk_chars": "500",
            "extraction_chunk_overlap_chars": "20",
            "extraction_max_facts_per_chunk": "5",
            "auto_accept_recommendations": "on",
        },
        follow_redirects=False,
    )
    assert r.status_code == 303
    # the next get_client would pick the ollama adapter — no code change
    assert c.app.state.cfg.provider == "ollama"
    assert c.app.state.cfg.extraction_chunk_chars == 500
    assert c.app.state.cfg.extraction_chunk_overlap_chars == 20
    assert c.app.state.cfg.extraction_max_facts_per_chunk == 5
    assert c.app.state.cfg.auto_accept_recommendations is True
    assert (tmp_path / "config.json").is_file()


def test_settings_disables_connection_test_for_untestable_provider(tmp_path):
    """Every provider verinote ships can be connection-tested, so the disabled
    state is reached by a `config.json` naming one it does not -- the case where
    running the test would mean resolving a provider that has no adapter."""
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="madeup",
        model="",
        api_key=None,
        base_url=None,
    )
    client = TestClient(create_app(cfg))

    r = client.get("/settings")

    assert "Test connection" in r.text
    assert 'aria-disabled="true"' in r.text
    assert "Connection test is not available for this provider." not in r.text


def test_test_connection_rejects_untestable_provider(tmp_path):
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="madeup",
        model="",
        api_key=None,
        base_url=None,
    )
    client = TestClient(create_app(cfg))

    r = client.post("/settings/test")

    assert r.status_code == 400
    assert "Connection test is not available for this provider." in r.text


def test_settings_enables_connection_test_for_claude_cli(tmp_path):
    """The CLI's models cannot be listed, so actually running the chosen one is
    the only evidence available that it works -- the button matters most here."""
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="claudecli",
        model="opus",
        api_key=None,
        base_url=None,
    )
    client = TestClient(create_app(cfg))

    r = client.get("/settings")

    assert "ClaudeCLI" in r.text
    assert "Test connection" in r.text
    assert 'aria-disabled="true"' not in r.text


def test_settings_enables_connection_test_for_ollama(tmp_path):
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="ollama",
        model="llama3.1",
        api_key=None,
        base_url=None,
    )
    client = TestClient(create_app(cfg))

    r = client.get("/settings")

    assert "Ollama" in r.text
    assert "Test connection" in r.text
    assert 'aria-disabled="true"' not in r.text


def test_settings_switches_active_kb_root(tmp_path, monkeypatch):
    monkeypatch.setenv("HOME", str(tmp_path / "home"))
    monkeypatch.setenv("XDG_CONFIG_HOME", str(tmp_path / "xdg"))
    monkeypatch.setenv("APPDATA", str(tmp_path / "appdata"))
    c = _client(tmp_path)
    other = tmp_path / "other-kb"

    r = c.post("/settings/root", data={"root": str(other)}, follow_redirects=False)

    assert r.status_code == 303
    assert r.headers["location"] == "/"
    assert c.app.state.cfg.root == other.resolve()
    assert c.app.state.store.db_path == other.resolve() / "kb.sqlite"
    assert (other / "kb.sqlite").is_file()
    assert not app_config_path().exists()
    assert "Review queue is empty" in c.get("/review").text
    assert str(other.resolve()) in c.get("/").text


def test_settings_rejects_empty_kb_root(tmp_path):
    c = _client(tmp_path)

    r = c.post("/settings/root", data={"root": "   "})

    assert r.status_code == 400
    assert "KB directory is required" in r.text


@pytest.mark.parametrize(
    "path",
    ["/kb/select", "/settings/root", "/settings/root/persist"],
)
def test_web_root_selection_errors_leave_active_root_and_app_config_unchanged(
    tmp_path, path
):
    c = _client(tmp_path)
    before_root = c.app.state.cfg.root
    config_path = app_config_path()
    config_path.parent.mkdir(parents=True)
    config_path.write_text(
        '{"active_root": "/synthetic/saved-kb", "extra": "keep"}\n',
        encoding="utf-8",
    )
    before_config = config_path.read_bytes()

    r = c.post(
        path,
        data={"root": "relative-kb", "confirm_persistence": "on"},
        follow_redirects=False,
    )

    assert r.status_code == 400
    assert "absolute path" in r.text
    assert c.app.state.cfg.root == before_root
    assert config_path.read_bytes() == before_config


@pytest.mark.parametrize("path", ["/kb/select", "/settings/root"])
def test_reopening_selected_kb_does_not_modify_saved_active_root(tmp_path, path):
    c = _client(tmp_path)
    config_path = app_config_path()
    config_path.parent.mkdir(parents=True)
    config_path.write_text(
        '{"active_root": "/synthetic/saved-kb", "extra": "keep"}\n',
        encoding="utf-8",
    )
    before = config_path.read_bytes()

    r = c.post(path, data={"root": str(tmp_path.resolve())}, follow_redirects=False)

    assert r.status_code == 303
    assert c.app.state.cfg.root == tmp_path.resolve()
    assert config_path.read_bytes() == before


def test_settings_renders_explicit_machine_wide_kb_confirmation(tmp_path):
    c = _client(tmp_path)

    r = c.get("/settings")

    assert r.status_code == 200
    assert 'action="/settings/root/persist"' in r.text
    assert 'name="confirm_persistence"' in r.text
    assert str(tmp_path.resolve()) in r.text
    assert "future verinote web processes" in r.text


def test_settings_persist_active_kb_requires_confirmation(tmp_path):
    c = _client(tmp_path)

    r = c.post(
        "/settings/root/persist",
        data={"root": str(tmp_path.resolve())},
        follow_redirects=False,
    )

    assert r.status_code == 400
    assert "Confirm the machine-wide KB change" in r.text
    assert not app_config_path().exists()


def test_settings_persist_active_kb_after_explicit_confirmation(tmp_path):
    c = _client(tmp_path)

    r = c.post(
        "/settings/root/persist",
        data={"root": str(tmp_path.resolve()), "confirm_persistence": "on"},
        follow_redirects=False,
    )

    assert r.status_code == 303
    assert r.headers["location"] == "/settings"
    assert read_app_config()["active_root"] == str(tmp_path.resolve())


def test_settings_persist_refuses_unsafe_root_without_touching_app_config(tmp_path):
    c = _client(tmp_path)
    root, expected = _unsafe_ui_root(tmp_path, "normal-worktree")

    r = c.post(
        "/settings/root/persist",
        data={"root": str(root), "confirm_persistence": "on"},
        follow_redirects=False,
    )

    assert r.status_code == 400
    assert "inside Git worktree" in r.text
    assert c.app.state.cfg.root == tmp_path.resolve()
    assert not expected.exists()
    assert not app_config_path().exists()


def test_settings_persist_refuses_when_verinote_root_overrides(tmp_path, monkeypatch):
    monkeypatch.setenv("VERINOTE_ROOT", str(tmp_path))
    c = _client(tmp_path)

    r = c.post(
        "/settings/root/persist",
        data={"root": str(tmp_path.resolve()), "confirm_persistence": "on"},
        follow_redirects=False,
    )

    assert r.status_code == 400
    assert "VERINOTE_ROOT controls this process" in r.text
    assert not app_config_path().exists()


def test_settings_never_renders_api_key(tmp_path):
    cfg = Config(
        root=tmp_path, db_path=tmp_path / "kb.sqlite",
        provider="anthropic", model="m", api_key="supersecret", base_url=None,
    )
    client = TestClient(create_app(cfg))
    r = client.get("/settings")
    assert "supersecret" not in r.text
    # `"API keys"` alone would be the unconditional <h2> and would pass whatever
    # the rows said. Assert the row's actual state for the configured provider.
    assert "Anthropic" in r.text
    assert "not set" in r.text


def test_test_connection_reports_adapter(tmp_path, monkeypatch, fake_client):
    monkeypatch.setattr(
        webapp, "get_client", lambda c: fake_client([ExtractedFact("A", "is_a", "B", 0.9)])
    )
    c = _client(tmp_path)
    r = c.post("/settings/test")
    assert r.status_code == 200
    assert "fake answered with 1 fact" in r.text


def test_test_connection_surfaces_llm_error(tmp_path, monkeypatch, fake_client):
    monkeypatch.setattr(webapp, "get_client", lambda c: fake_client(error=LLMError("no key")))
    c = _client(tmp_path)
    r = c.post("/settings/test")
    assert r.status_code == 502
    assert "connection failed: no key" in r.text


def test_worker_demotes_a_stale_fact_and_does_not_re_promote_it_same_request(
    tmp_path, monkeypatch, fake_client
):
    # #329 web path: overwriting a source's text (a new artifact + a fresh job) runs
    # the sweep as a sibling of extraction, returning the now-unsupported confirmed
    # citation to review. The stale flag plus exclude_fact_ids keep the SAME
    # request's auto-accept pass from immediately re-promoting it, even though two
    # other sources still corroborate the value.
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="anthropic",
        model="m",
        api_key=None,
        base_url=None,
        auto_accept_recommendations=True,
    )
    policy = tmp_path / POLICY_RELPATH
    with Store(cfg.db_path) as store:
        store.init_schema()
        src_a = store.add_source("sources/a.txt")
        old = store.add_source_artifact(
            source_id=src_a, kind="original_text", path="sources/a-v1.txt", checksum="a1"
        )
        new = store.add_source_artifact(
            source_id=src_a, kind="original_text", path="sources/a-v2.txt", checksum="a2"
        )
        old_job = store.create_extraction_job(
            source_id=src_a, artifact_id=old, provider="anthropic", model="m", total_chunks=1
        )
        oc = store.add_source_chunks(
            job_id=old_job, source_id=src_a, chunks=["London body"]
        )[0]
        store.mark_extraction_job_running(old_job)
        store.mark_chunk_running(oc)
        store.mark_chunk_done(oc)
        store.finish_extraction_job(old_job)
        london_a = store.add_fact(
            "Ada", "born_in", "London",
            status="confirmed", source_id=src_a, job_id=old_job,
        )
        store.add_fact_evidence(
            fact_id=london_a, source_id=src_a, artifact_id=old, snippet="London"
        )
        # Two other sources corroborate the value, so it WOULD be auto-accept
        # eligible were it not stale + excluded this pass.
        for path in ("sources/b.txt", "sources/c.txt"):
            witness = store.add_source(path)
            wj = store.create_extraction_job(
                source_id=witness, provider="anthropic", model="m", total_chunks=1
            )
            wc = store.add_source_chunks(job_id=wj, source_id=witness, chunks=["x"])[0]
            store.mark_extraction_job_running(wj)
            store.mark_chunk_running(wc)
            store.mark_chunk_done(wc)
            store.finish_extraction_job(wj)
            store.add_fact(
                "Ada", "born_in", "London",
                status="confirmed", source_id=witness, job_id=wj,
            )
        # The overwrite: a pending job at the NEW artifact whose text says Paris.
        new_job = store.create_extraction_job(
            source_id=src_a, artifact_id=new, provider="anthropic", model="m", total_chunks=1
        )
        store.add_source_chunks(
            job_id=new_job, source_id=src_a, chunks=["Ada was born in Paris."]
        )
        policy.parent.mkdir(parents=True, exist_ok=True)
        policy.write_text(DEFAULT_POLICY, encoding="utf-8")
        store.record_policy_marker(policy_sha256(DEFAULT_POLICY), origin="scaffold")

    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("Ada", "born_in", "Paris", 0.9)]),
    )

    create_app(cfg)  # resumes the pending job -> _start_source_extraction

    def swept():
        assert _job_row(cfg, new_job)["status"] == "done"
        with Store(cfg.db_path) as store:
            store.init_schema()
            london = next(
                f for f in store.facts()
                if f["object"] == "London" and f["source_id"] == src_a
            )
            assert london["status"] == "needs_review" and london["stale"] == 1

    _wait_for(swept)
    time.sleep(0.2)  # let an unguarded re-promotion land, if the wiring regressed
    with Store(cfg.db_path) as store:
        store.init_schema()
        london = next(
            f for f in store.facts()
            if f["object"] == "London" and f["source_id"] == src_a
        )
        assert london["status"] == "needs_review"  # demoted, not re-promoted
        assert london["stale"] == 1
        paris = next(f for f in store.facts() if f["object"] == "Paris")
        assert paris["status"] == "candidate"


def test_worker_leaves_a_done_job_done_when_the_stale_sweep_raises(
    tmp_path, monkeypatch, fake_client
):
    # #329 exception-safety: the sweep is a sibling call inside the worker's
    # try/except, so it USED TO BE that without the local guard a sweep error
    # reached the outer `except Exception -> fail_extraction_job` and retroactively
    # flipped a completed extraction to `failed` (the "KB lies about its own run
    # state" class #194/#239 closed). Since #525 the outer clause re-reads and
    # declines a `done` job, so this test no longer distinguishes the local guard:
    # MEASURED, removing that guard leaves it green. What it still pins is that a
    # sweep error leaves the completed job `done`, whichever layer contains it.
    # Note the Config below leaves `auto_accept_recommendations` at its `False`
    # default, so this fixture does NOT exercise the sweep guard's one remaining
    # effect — letting auto-accept run after a failed sweep.
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="anthropic",
        model="m",
        api_key=None,
        base_url=None,
    )
    policy = tmp_path / POLICY_RELPATH
    with Store(cfg.db_path) as store:
        store.init_schema()
        sid = store.add_source("sources/a.txt")
        art = store.add_source_artifact(
            source_id=sid, kind="original_text", path="sources/a.txt", checksum="v1"
        )
        job_id = store.create_extraction_job(
            source_id=sid, artifact_id=art, provider="anthropic", model="m", total_chunks=1
        )
        store.add_source_chunks(job_id=job_id, source_id=sid, chunks=["some text"])
        policy.parent.mkdir(parents=True, exist_ok=True)
        policy.write_text(DEFAULT_POLICY, encoding="utf-8")
        store.record_policy_marker(policy_sha256(DEFAULT_POLICY), origin="scaffold")

    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("A", "is_a", "B", 0.9)]),
    )

    def boom(self, job_id):
        raise RuntimeError("sweep exploded")

    monkeypatch.setattr(store_db.Store, "surface_stale_engine_facts", boom)

    create_app(cfg)  # resumes the pending job -> _start_source_extraction

    def job_finished():
        assert _job_row(cfg, job_id)["status"] == "done"

    _wait_for(job_finished)
    time.sleep(0.2)  # let a late fail_extraction_job land, if the guard regressed
    row = _job_row(cfg, job_id)
    assert row["status"] == "done"  # the sweep error did NOT flip it to failed
    assert row["failed_chunks"] == 0
    # extraction genuinely completed; only the (now contained) sweep failed.
    with Store(cfg.db_path) as store:
        store.init_schema()
        assert any(f["subject"] == "A" for f in store.facts())


# --- #269: a corrupt config.json halts the web app instead of silently
#     defaulting extraction/ask/test to the cloud provider the user never chose ---


def _config_web_kb(tmp_path, *, corrupt, with_job=False):
    """A KB with a *recorded* policy (so the policy guard is inert) whose
    config.json is corrupt when `corrupt`, else a valid ollama config. Returns
    (cfg, fact_id, job_id); job_id is None unless `with_job`. cfg is resolved via
    Config.for_root so it carries the real settings_error the app will see."""
    policy = tmp_path / POLICY_RELPATH
    job_id = None
    with Store(tmp_path / "kb.sqlite") as store:
        store.init_schema()
        fact_id = store.add_fact("A", "is_a", "B", status="needs_review", confidence=0.9)
        policy.parent.mkdir(parents=True, exist_ok=True)
        policy.write_text(DEFAULT_POLICY, encoding="utf-8")
        store.record_policy_marker(policy_sha256(DEFAULT_POLICY), origin="scaffold")
        if with_job:
            sid = store.add_source("sources/a.txt")
            job_id = store.create_extraction_job(
                source_id=sid, provider="ollama", model="m", total_chunks=1
            )
            store.add_source_chunks(job_id=job_id, source_id=sid, chunks=["some text"])
    if corrupt:
        (tmp_path / "config.json").write_text("{bad json", encoding="utf-8")
    else:
        save_settings(tmp_path, provider="ollama", model="llama3.1")
    return Config.for_root(tmp_path), fact_id, job_id


def _spy_anthropic(monkeypatch):
    """Record every construction of the concrete cloud adapter class.

    The corrupt-config fallback provider is `anthropic`, so a non-empty list means
    an adapter was built and bytes were about to leave the machine. A 409 alone
    would not prove they didn't — only that the *response* was a halt."""
    import verinote.llm.anthropic_adapter as aa

    built = []
    orig = aa.AnthropicAdapter.__init__

    def spy(self, cfg):
        built.append(cfg)
        orig(self, cfg)

    monkeypatch.setattr(aa.AnthropicAdapter, "__init__", spy)
    return built


def test_ask_on_corrupt_config_halts_and_never_builds_the_adapter(tmp_path, monkeypatch):
    """A provider-calling POST (`/ask`) refuses with 409 AND the concrete adapter
    is never constructed — the confidentiality guarantee, not just a halt page."""
    cfg, _, _ = _config_web_kb(tmp_path, corrupt=True)
    built = _spy_anthropic(monkeypatch)
    c = TestClient(create_app(cfg))

    r = c.post("/ask", data={"question": "who is A?"})

    assert r.status_code == 409
    assert "did not choose" in r.text  # the config_corrupt page, not a stack trace
    assert built == []  # no bytes left the machine


def test_sources_upload_on_corrupt_config_halts_before_extraction(tmp_path, monkeypatch):
    """`/sources` refuses on the triggering request itself (the synchronous hoist
    in `_start_source_extraction`), so no adapter is built and no doomed worker
    runs."""
    cfg, _, _ = _config_web_kb(tmp_path, corrupt=True)
    built = _spy_anthropic(monkeypatch)
    c = TestClient(create_app(cfg))

    r = c.post("/sources", files={"file": ("note.txt", b"some text", "text/plain")})

    assert r.status_code == 409
    assert built == []


def test_read_only_route_stays_reachable_under_corrupt_config(tmp_path):
    """No over-blocking: `/review` reads no provider, so a deliberate absence of a
    config middleware leaves it reachable at 200 while config is corrupt."""
    cfg, _, _ = _config_web_kb(tmp_path, corrupt=True)
    c = TestClient(create_app(cfg))

    r = c.get("/review")

    assert r.status_code == 200
    assert "is_a" in r.text  # the reviewable fact renders normally


def test_settings_test_over_htmx_redirects_not_inline(tmp_path):
    """htmx never swaps a 4xx into the DOM (#173), so the handler answers an htmx
    request with HX-Redirect to the full-page halt, not an inline 409 body."""
    cfg, _, _ = _config_web_kb(tmp_path, corrupt=True)
    c = TestClient(create_app(cfg))

    r = c.post("/settings/test", headers={"HX-Request": "true"})

    assert r.status_code == 409
    assert r.headers.get("HX-Redirect") == webapp.CONFIG_UNAVAILABLE_PATH
    # An htmx halt sends no inline body to be (silently) swallowed.
    assert "did not choose" not in r.text


def test_settings_page_reachable_and_warns_under_corrupt_config(tmp_path):
    """`/settings` is the recovery page: it must stay reachable AND must not present
    the fallback provider as if it were the user's saved choice."""
    cfg, _, _ = _config_web_kb(tmp_path, corrupt=True)
    c = TestClient(create_app(cfg))

    r = c.get("/settings")

    assert r.status_code == 200
    body = " ".join(unescape(r.text).split())  # collapse template line-breaks
    assert "built-in defaults" in body  # the warning banner
    assert "not your saved choice" in body


def test_resaving_valid_settings_clears_the_config_halt(tmp_path, monkeypatch, fake_client):
    """Re-saving a provider writes a fresh valid file; the app re-resolves and the
    next provider call goes through instead of halting."""
    cfg, _, _ = _config_web_kb(tmp_path, corrupt=True)
    app = create_app(cfg)
    c = TestClient(app)

    r = c.post("/settings", data={"provider": "ollama", "model": "llama3.1"})
    assert r.status_code == 200  # followed the 303 redirect back to /settings
    assert app.state.cfg.settings_error is None  # halt cleared on re-resolution

    # And a provider call now reaches the client instead of 409-ing.
    monkeypatch.setattr(
        webapp, "get_client", lambda cfg: fake_client([ExtractedFact("A", "is_a", "B", 0.9)])
    )
    assert c.post("/settings/test").status_code == 200
    assert "built-in defaults" not in c.get("/settings").text  # warning gone


def test_mid_session_corruption_is_caught_at_the_next_resolution(tmp_path, monkeypatch):
    """Healthy at create_app, corrupted on disk afterward: the next re-resolution
    (re-opening the same root) must read DISK and catch it, not trust a cached
    healthy dict from launch time.

    The same action is the discriminator: re-opening the root succeeds (303) while
    healthy, and is refused (400) once the file is corrupt.

    This covers the explicit re-open path specifically. A passive on-disk corruption
    with NO re-open keeps serving the cached (last-good) cfg, which is safe precisely
    because `cfg.settings_error` is a one-time snapshot — the same frozen-snapshot
    semantics that make the worker-thread ConfigCorruptError clause unreachable today
    (see `test_worker_config_corrupt_does_not_mark_the_job_failed`) are what keep the
    stale window from ever silently switching provider."""
    cfg, _, _ = _config_web_kb(tmp_path, corrupt=False)  # valid ollama config
    app = create_app(cfg)
    c = TestClient(app)
    assert app.state.cfg.settings_error is None

    # Baseline: re-opening the (healthy) current root works.
    r = c.post("/settings/root", data={"root": str(tmp_path)}, follow_redirects=False)
    assert r.status_code == 303
    assert app.state.cfg.provider == "ollama"

    # Corrupt the current KB's config on disk, then re-open the same root.
    (tmp_path / "config.json").write_text("{bad json", encoding="utf-8")
    built = _spy_anthropic(monkeypatch)
    r = c.post("/settings/root", data={"root": str(tmp_path)}, follow_redirects=False)

    assert r.status_code == 400  # caught, not a stale/cached 303 false-negative
    assert "config.json is corrupt" in unescape(r.text)
    assert app.state.cfg.provider == "ollama"  # old healthy cfg still active
    assert built == []


def test_switching_root_to_a_corrupt_kb_is_refused_inline(tmp_path, monkeypatch):
    """Switching from a healthy KB into a corrupt-config KB is refused inline (400)
    with the OLD KB still active — never a transient swap into an untrusted cfg."""
    healthy = tmp_path / "healthy"
    corrupt = tmp_path / "corrupt"
    healthy.mkdir()
    corrupt.mkdir()
    active_cfg, _, _ = _config_web_kb(healthy, corrupt=False)  # provider=ollama
    _config_web_kb(corrupt, corrupt=True)

    app = create_app(active_cfg)
    c = TestClient(app)
    built = _spy_anthropic(monkeypatch)

    r = c.post("/settings/root", data={"root": str(corrupt)}, follow_redirects=False)

    assert r.status_code == 400
    assert "config.json is corrupt" in unescape(r.text)
    # The old KB stays active: root unchanged, provider not swapped to the fallback.
    assert app.state.cfg.root == healthy.resolve()
    assert app.state.cfg.provider == "ollama"
    assert built == []


def test_launching_the_ui_on_a_corrupt_config_resumes_nothing(tmp_path, monkeypatch):
    """A corrupt config discovered at launch resumes no jobs and touches nothing —
    zero HTTP requests, and still the pending job must not be started or failed.
    Mirrors `test_launching_the_ui_on_a_halted_kb_resumes_nothing` for #269."""
    cfg, _, job_id = _config_web_kb(tmp_path, corrupt=True, with_job=True)
    clients = []
    monkeypatch.setattr(webapp, "get_client", lambda cfg: clients.append(cfg))
    before = _job_row(cfg, job_id)

    create_app(cfg)

    time.sleep(0.2)  # a worker, had one started, would have written by now
    assert clients == []  # no worker was started at all
    assert _job_row(cfg, job_id) == before  # job untouched: still pending
    assert before["status"] == "pending"
    assert _job_event_types(cfg, job_id) == []  # no started / failed / rolled_back


def test_launching_the_ui_on_a_valid_config_still_resumes(tmp_path, monkeypatch, fake_client):
    """The config gate above refuses a *corrupt* config, not every config: a valid
    one with a pending job still resumes normally."""
    cfg, _, job_id = _config_web_kb(tmp_path, corrupt=False, with_job=True)
    monkeypatch.setattr(
        webapp, "get_client", lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)])
    )

    create_app(cfg)  # the resume worker runs off a daemon thread started here

    def resumed():
        assert _job_row(cfg, job_id)["status"] == "done"

    _wait_for(resumed)


def test_worker_config_corrupt_does_not_mark_the_job_failed(tmp_path, monkeypatch, fake_client):
    """`except ConfigCorruptError` must stay ABOVE `except Exception` in the worker.

    This exercises an INJECTED scenario, not a naturally-occurring race: the worker's
    get_client(cfg) reads the SAME frozen cfg that already passed the synchronous
    hoist, and `cfg.settings_error` is a one-time `Config.for_root` snapshot, so today
    the thread can never independently observe a fresher corruption. We monkeypatch
    get_client to raise ConfigCorruptError to prove the clause is load-bearing IF such
    an error were ever raised there (a future get_client that re-reads disk, or a
    caller that passes a different cfg into the thread). A corrupt config is a
    host/environment condition, not content-attributable, so the worker must write
    NOTHING — not bury the job in `failed` (a misleading "analysis failed") and burn
    this session's retry budget. Below `except Exception`, the generic handler would
    call `fail_extraction_job`; this test pins that it does not."""
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)  # healthy cfg -> hoist passes
    called = threading.Event()

    def corrupt_at_get_client(cfg):
        # Stand in for "config.json went corrupt after the hoist check": the worker
        # reaches get_client and it raises the same error get_client would.
        called.set()
        raise ConfigCorruptError("config.json is not valid JSON")

    monkeypatch.setattr(webapp, "get_client", corrupt_at_get_client)

    create_app(cfg)  # resumes the pending job -> worker thread -> get_client raises

    assert called.wait(timeout=2.0)
    time.sleep(0.2)  # let a (wrong) fail_extraction_job land, if the order regressed
    job = _job_row(cfg, job_id)
    assert job["status"] == "pending", "a corrupt-config race buried the job in `failed`"
    assert "analysis failed" not in (job["message"] or "")
    assert "extraction_job_failed" not in _job_event_types(cfg, job_id)


def test_worker_sidecar_lock_does_not_overwrite_the_rollback(tmp_path, monkeypatch, fake_client):
    """`except DuckDBFactTermStoreLockedError` must stay ABOVE `except Exception`.

    It is a `ValueError` subclass, so without a clause of its own the generic
    handler takes it and calls `fail_extraction_job`. That is worse here than in
    the sibling cases: `process_extraction_job` has ALREADY rolled the job back to
    `pending` so the next pass resumes it, and the `failed` write lands on top of
    that rollback. The result is a `failed` job with no failed chunk. Since #524
    `plan_source_extraction` continues such a job on the strength of the chunks it
    finished rather than rebuilding from scratch, so those chunks are no longer
    sent to the LLM again; what the overwrite still does is file a condition of
    the host as this job's own failure — in the job row, in the
    `extraction_job_failed` event, and on the Sources page — for a pause that
    clears when the other process lets go.

    The stand-in raises after rewinding the job, which is the state the real
    pipeline hands the worker (`_back_off_from_locked_sidecar`).
    """
    cfg, job_id, _ = _job_kb(tmp_path, with_policy=True)
    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )
    called = threading.Event()

    def locked(store, *args, **kwargs):
        store.rollback_extraction_job(job_id, "Paused: another process is holding it.")
        called.set()
        raise DuckDBFactTermStoreLockedError("the fact-term store is locked")

    monkeypatch.setattr(webapp, "process_extraction_job", locked)

    create_app(cfg)

    assert called.wait(timeout=2.0)
    time.sleep(0.2)  # let a (wrong) fail_extraction_job land, if the order regressed
    job = _job_row(cfg, job_id)
    assert job["status"] == "pending", "the rollback was overwritten by a `failed` write"
    assert "analysis failed" not in (job["message"] or "")
    assert "extraction_job_failed" not in _job_event_types(cfg, job_id)


# --- the Ollama model picker: choose from what the server actually has ---


def _ollama_client(
    tmp_path,
    *,
    provider="ollama",
    model="qwen3:8b",
    base_url=None,
    api_key=None,
    llm_timeout_seconds=None,
):
    # No default of its own: an unset timeout must be Config's own default, so
    # this helper cannot silently retune every test that uses it.
    timeout = (
        {} if llm_timeout_seconds is None else {"llm_timeout_seconds": llm_timeout_seconds}
    )
    return TestClient(
        create_app(
            Config(
                root=tmp_path,
                db_path=tmp_path / "kb.sqlite",
                provider=provider,
                model=model,
                api_key=api_key,
                base_url=base_url,
                **timeout,
            )
        )
    )


class _FakeLister:
    """Stands in for the keyless lister, recording the base URL and timeout it was
    called with so a test can prove which endpoint got asked, and how patiently.

    Installed over `webapp._MODEL_LISTERS` rather than over `webapp.get_client`:
    the listing path deliberately no longer builds a client, and a fake wired to
    the factory would sit unused while the real endpoint was dialled for real.

    `structured_output_ids` defaults to `None` -- what a listing that does not
    report the property returns, which is Ollama's case and the one most of these
    tests are about. A listing that does report it passes a frozenset, including
    an empty one, because "reported and none" is a different fact from "not
    reported" and the picker renders the two differently.
    """

    def __init__(self, models=(), error=None, structured_output_ids=None):
        self.models = tuple(models)
        self.structured_output_ids = structured_output_ids
        self.error = error
        self.seen = []
        self.timeouts = []

    def install(self, monkeypatch, provider="ollama"):
        monkeypatch.setitem(webapp._MODEL_LISTERS, provider, self.lister)
        return self

    def lister(self, base_url, timeout):
        self.seen.append(base_url)
        self.timeouts.append(timeout)
        if self.error is not None:
            raise self.error
        return ModelListing(
            models=self.models, structured_output_ids=self.structured_output_ids
        )


def test_settings_defers_the_ollama_model_list_off_the_page_load(tmp_path, monkeypatch):
    """`/settings` is also the recovery page for a corrupt config and a halted
    policy, so rendering it must never wait on a provider. The field ships as a
    working text input that lazily swaps itself for the picker."""
    monkeypatch.setitem(
        webapp._MODEL_LISTERS,
        "ollama",
        lambda base_url, timeout: pytest.fail("/settings must not call the provider"),
    )

    r = _ollama_client(tmp_path).get("/settings")

    assert r.status_code == 200
    assert 'id="model-field"' in r.text
    assert 'hx-get="/settings/model-field"' in r.text
    assert 'hx-trigger="load"' in r.text
    assert 'name="model"' in r.text  # usable without htmx, never a dead field


def test_settings_does_not_defer_for_a_provider_with_nothing_to_enumerate(tmp_path):
    """A cloud catalogue is not a property of the endpoint, so anthropic keeps the
    text input and must NOT sit in a lazy state waiting for a list forever."""
    r = _ollama_client(tmp_path, provider="anthropic", model="claude-opus-4-8").get("/settings")

    assert 'id="model-field"' in r.text
    assert 'hx-trigger="load"' not in r.text
    assert '<input type="text" name="model" value="claude-opus-4-8"' in r.text


def test_model_field_offers_the_installed_models_as_a_picker(tmp_path, monkeypatch):
    _FakeLister(models=["llava:7b", "qwen3:8b"]).install(monkeypatch)

    r = _ollama_client(tmp_path).get("/settings/model-field?provider=ollama&model=qwen3:8b")

    assert r.status_code == 200
    assert '<select name="model">' in r.text
    assert '<option value="llava:7b" >llava:7b</option>' in unescape(r.text)
    assert '<option value="qwen3:8b" selected>qwen3:8b</option>' in unescape(r.text)
    assert 'type="text" name="model"' not in r.text  # no free-text field alongside


def test_model_field_asks_the_endpoint_being_edited_not_the_saved_one(tmp_path, monkeypatch):
    """The form's Base URL is the endpoint whose models the user is choosing from.
    Listing against the saved config instead would show models from a server the
    user is in the middle of switching away from."""
    fake = _FakeLister(models=["qwen3:8b"]).install(monkeypatch)

    r = _ollama_client(tmp_path, base_url="http://saved:11434").get(
        "/settings/model-field?provider=ollama&model=qwen3:8b&base_url=http://edited:9999"
    )

    assert fake.seen == ["http://edited:9999"]
    assert "http://edited:9999" in r.text


def test_model_field_keeps_a_configured_model_the_server_does_not_have(tmp_path, monkeypatch):
    """config.json really does name this model. Dropping it from the picker would
    silently select a different one and make the page misreport the KB's state."""
    _FakeLister(models=["llava:7b", "qwen3:8b"]).install(monkeypatch)

    r = _ollama_client(tmp_path, model="llama3.1").get(
        "/settings/model-field?provider=ollama&model=llama3.1"
    )

    assert '<option value="llama3.1" selected>llama3.1 — not installed</option>' in unescape(r.text)
    assert "which is not installed on" in r.text
    assert "ollama pull llama3.1" in r.text


def test_model_field_falls_back_to_typing_when_the_server_is_unreachable(tmp_path, monkeypatch):
    """An empty picker would read as 'nothing to choose' and, worse, leave the user
    unable to set a model at all. Say what failed and keep the field usable."""
    _FakeLister(error=LLMError("ollama request failed: connection refused")).install(monkeypatch)

    r = _ollama_client(tmp_path).get("/settings/model-field?provider=ollama&model=qwen3:8b")

    assert r.status_code == 200
    assert "<select" not in r.text
    assert '<input type="text" name="model" value="qwen3:8b"' in r.text
    assert "Could not load the model list" in r.text
    assert "connection refused" in r.text


def test_model_field_distinguishes_an_empty_server_from_an_unreachable_one(tmp_path, monkeypatch):
    """Reachable-with-nothing-pulled and could-not-be-reached are different facts
    about the user's machine, and the fix for each is different."""
    _FakeLister(models=[]).install(monkeypatch)

    r = _ollama_client(tmp_path).get("/settings/model-field?provider=ollama&model=")

    assert "has no models installed" in r.text
    assert "Could not load the model list" not in r.text
    assert '<input type="text" name="model"' in r.text


def test_model_field_names_the_default_endpoint_it_will_actually_dial(tmp_path, monkeypatch):
    fake = _FakeLister(models=[]).install(monkeypatch)

    r = _ollama_client(tmp_path).get("/settings/model-field?provider=ollama&model=")

    assert fake.seen == [None]  # an unset Base URL stays unset in config
    assert "http://localhost:11434" in r.text  # ...but the page names the real target
    # Ollama's entry in the default-endpoint map, not just *some* entry: a map
    # whose two values were swapped would otherwise satisfy the openrouter test
    # and this one both, while telling every user the wrong host.
    assert "openrouter.ai" not in r.text


def test_model_field_does_not_reach_the_provider_for_a_text_input_provider(tmp_path, monkeypatch):
    monkeypatch.setattr(
        webapp,
        "_list_models_for",
        lambda *a, **k: pytest.fail("no provider call for anthropic"),
    )

    r = _ollama_client(tmp_path).get("/settings/model-field?provider=anthropic&model=gpt")

    assert '<input type="text" name="model" value="gpt"' in r.text
    assert 'hx-trigger="load"' not in r.text


def test_model_field_halts_under_a_corrupt_config(tmp_path, monkeypatch):
    """The picker is a provider call like any other: a corrupt config.json makes the
    resolved provider untrustworthy, so it halts rather than listing against a
    provider the user never chose (#269).

    The lister is replaced with a tripwire for the same reason its credentials-halt
    sibling does it: a regressed halt would otherwise dial localhost:11434 for real
    on any dev box running Ollama, and the assertions below would still pass or
    fail on the status code alone without saying that happened.
    """
    cfg, _, _ = _config_web_kb(tmp_path, corrupt=True)
    monkeypatch.setattr(
        webapp,
        "_MODEL_LISTERS",
        {"ollama": lambda base_url, timeout: pytest.fail("dialled despite the halt")},
    )
    c = TestClient(create_app(cfg))

    r = c.get(
        "/settings/model-field?provider=ollama&model=m", headers={"HX-Request": "true"}
    )

    assert r.status_code == 409
    assert r.headers.get("HX-Redirect") == webapp.CONFIG_UNAVAILABLE_PATH


def test_model_field_halts_under_unreadable_credentials(tmp_path, monkeypatch):
    """The other halt `get_client` used to supply for free on this path.

    Dropping the factory call is what puts the API key out of the listing's
    reach, but it also removed both halts, so each is now a hand-written line
    with nothing else guarding it. Without this test `assert_credentials_intact`
    could be deleted from `_list_models_for` and everything else stays green.
    """
    cfg, _, _ = _config_web_kb(tmp_path, corrupt=False)
    monkeypatch.setattr(
        webapp,
        "_MODEL_LISTERS",
        {"ollama": lambda base_url, timeout: pytest.fail("dialled despite the halt")},
    )
    c = TestClient(create_app(replace(cfg, credentials_error="credentials.json is not valid JSON")))

    r = c.get(
        "/settings/model-field?provider=ollama&model=m", headers={"HX-Request": "true"}
    )

    assert r.status_code == 409
    assert r.headers.get("HX-Redirect") == webapp.CREDENTIALS_UNAVAILABLE_PATH


def test_model_field_halts_on_config_first_when_both_errors_are_set(tmp_path, monkeypatch):
    """Config-first order: settings_error must win when both halts apply.

    A corrupt config.json falls back to the built-in cloud provider, which
    requires a key, so credentials_error can be set alongside settings_error
    on the very same Config. `_list_models_for` hand-writes the same
    settings-then-credentials order `get_client` centralises; nothing keeps
    the two surfaces agreeing. Swapping the two assert lines produces zero
    failures elsewhere in the suite -- this is the one test that would catch it.
    """
    cfg, _, _ = _config_web_kb(tmp_path, corrupt=True)
    monkeypatch.setattr(
        webapp,
        "_MODEL_LISTERS",
        {"ollama": lambda base_url, timeout: pytest.fail("dialled despite the halt")},
    )
    c = TestClient(
        create_app(replace(cfg, credentials_error="credentials.json is not valid JSON"))
    )

    r = c.get(
        "/settings/model-field?provider=ollama&model=m", headers={"HX-Request": "true"}
    )

    assert r.status_code == 409
    assert r.headers.get("HX-Redirect") == webapp.CONFIG_UNAVAILABLE_PATH


def test_model_field_never_constructs_a_provider_adapter(tmp_path, monkeypatch):
    """The load-bearing test for the keyless listing seam.

    A `Config` carries a resolved `api_key` — and, after `replace(provider=...)`,
    the *saved* provider's key. So the guarantee is not "the lister happens not
    to read it" but "no object holding it is ever built on this path". Counting
    constructions, rather than asserting an error, is what keeps this true
    against a future edit that gives some adapter a `list_models` again.

    `OpenAIAdapter.__init__` is patched rather than `OpenRouterAdapter`'s, which
    it inherits; patching the subclass would miss a listing routed through the base.
    `ClaudeCliAdapter` never reads `cfg.api_key`, but it still stores the whole
    `Config` (`self.cfg = cfg`), so constructing one here would put an object
    holding the saved provider's key on this path. All four are checked because
    the name says *no* adapter and that has to be literally what is checked.
    """
    built = []
    for cls in (OllamaAdapter, OpenAIAdapter, AnthropicAdapter, ClaudeCliAdapter):
        monkeypatch.setattr(
            cls, "__init__", lambda self, cfg, _n=cls.__name__: built.append(_n)
        )
    monkeypatch.setattr(
        webapp, "get_client", lambda _cfg: pytest.fail("the listing path built a client")
    )
    fake = _FakeLister(models=["qwen3:8b"]).install(monkeypatch)

    r = _ollama_client(tmp_path).get("/settings/model-field?provider=ollama&model=qwen3:8b")

    assert r.status_code == 200
    assert '<select name="model">' in r.text  # it really did render the picker
    assert fake.seen == [None]  # ...and really did go through the lister
    assert built == []


def test_model_field_listing_clamps_the_page_load_timeout(tmp_path, monkeypatch):
    """A page-load call must not inherit the minutes-long completion budget.

    The clamp moved here with the seam: the lister takes a float and cannot
    re-derive it, so this is the only place left that can get it wrong.
    """
    fake = _FakeLister(models=[]).install(monkeypatch)

    _ollama_client(tmp_path, llm_timeout_seconds=900.0).get(
        "/settings/model-field?provider=ollama&model="
    )

    assert fake.timeouts == [5.0]


def test_model_field_listing_keeps_an_even_shorter_configured_timeout(tmp_path, monkeypatch):
    """The bound is the *smaller* of the two, so a tighter user setting still wins.
    Without this, clamping to a flat 5.0 would pass the test above."""
    fake = _FakeLister(models=[]).install(monkeypatch)

    _ollama_client(tmp_path, llm_timeout_seconds=1.5).get(
        "/settings/model-field?provider=ollama&model="
    )

    assert fake.timeouts == [1.5]


@pytest.mark.parametrize(
    "providers, listers, expected",
    [
        ({"ollama", "openai"}, {"ollama": lambda base_url, timeout: []}, "openai"),
        ({"ollama"}, {"ollama": lambda base_url, timeout: [], "openai": None}, "openai"),
    ],
)
def test_the_check_rejects_a_provider_set_and_lister_table_that_disagree(
    providers, listers, expected
):
    """Both directions: a picker that can never fill, and dead code a later edit
    could wire up without passing this gate."""
    with pytest.raises(RuntimeError, match=expected):
        webapp._check_every_listable_provider_has_a_keyless_lister(providers, listers)


def test_the_check_rejects_a_lister_that_could_be_handed_a_config():
    """Set equality alone would accept this. The parameter-name clause is what
    keeps 'keyless' a fact about the code rather than a claim in a comment.

    Mutation: delete the `_LISTER_SIGNATURE` clause and this no longer raises --
    a plain non-closing lambda passes every other clause.
    """
    with pytest.raises(RuntimeError, match="cannot be handed an API key"):
        webapp._check_every_listable_provider_has_a_keyless_lister(
            {"ollama"}, {"ollama": lambda cfg, base_url, timeout: []}
        )


# The next three are the shapes that satisfy the parameter names *while already
# carrying a Config*, which is what makes the names alone insufficient. Each was
# accepted by the check before the `isfunction` and `__closure__` clauses existed.


def test_the_check_rejects_a_partial_that_has_already_bound_a_config():
    """`functools.partial` hides the bound first argument from `signature`, so a
    keyed lister with `cfg` pre-filled reports exactly ('base_url', 'timeout').

    Mutation: delete the `inspect.isfunction` clause and this stops raising the
    RuntimeError asserted here -- the check falls through to `fn.__closure__`,
    which a partial does not have at all.
    """
    def keyed_lister(cfg, base_url, timeout):
        return [cfg.api_key]

    bound = functools.partial(keyed_lister, object())
    # The premise: it looks exactly like a conforming lister.
    assert tuple(inspect.signature(bound).parameters) == webapp._LISTER_SIGNATURE

    with pytest.raises(RuntimeError, match="must be a plain function"):
        webapp._check_every_listable_provider_has_a_keyless_lister(
            {"ollama"}, {"ollama": bound}
        )


def test_the_check_rejects_a_closure_over_a_captured_config():
    """A closure cell holds the Config, and no parameter ever mentions it. This is
    the shape a maintainer reaches for most naturally -- defining a lister inside
    a factory that already has the config in scope.

    Mutation: delete the `fn.__closure__` clause and this stops raising -- a
    closure is a plain function with the right parameter names.
    """
    cfg = object()

    def lister(base_url, timeout):
        return [cfg]

    # The premise: it looks exactly like a conforming lister.
    assert tuple(inspect.signature(lister).parameters) == webapp._LISTER_SIGNATURE
    assert lister.__closure__ is not None

    with pytest.raises(RuntimeError, match="must not close over anything"):
        webapp._check_every_listable_provider_has_a_keyless_lister(
            {"ollama"}, {"ollama": lister}
        )


def test_the_check_rejects_a_callable_object_holding_a_config():
    """`__call__(self, base_url, timeout)` presents the right parameters (`signature`
    drops `self`) while the instance carries the Config as an attribute.

    Mutation: delete the `inspect.isfunction` clause and this stops raising the
    RuntimeError asserted here -- the check falls through to `fn.__closure__`,
    which an instance does not have.
    """
    class KeyedLister:
        def __init__(self, cfg):
            self.cfg = cfg

        def __call__(self, base_url, timeout):
            return [self.cfg]

    lister = KeyedLister(object())
    # The premise: it looks exactly like a conforming lister.
    assert tuple(inspect.signature(lister).parameters) == webapp._LISTER_SIGNATURE

    with pytest.raises(RuntimeError, match="must be a plain function"):
        webapp._check_every_listable_provider_has_a_keyless_lister(
            {"ollama"}, {"ollama": lister}
        )


# And the names themselves are read off the function rather than trusted from
# `inspect.signature`'s default view, because a lister can lie about them without
# taking any of the three shapes above.


def test_the_check_rejects_a_lister_whose_wrapped_attribute_hides_a_config():
    """`inspect.signature` follows `__wrapped__`, so a plain, non-closing function
    that really takes `(cfg, base_url, timeout)` can report the conforming two by
    pointing that attribute at a decoy.

    Mutation: drop `follow_wrapped=False` from the parameter-name clause and this
    stops raising -- the decoy's names are then what the check sees, and every
    other clause already passes.
    """
    def decoy(base_url, timeout):
        return []

    def keyed_lister(cfg, base_url, timeout):
        return [cfg.api_key]

    keyed_lister.__wrapped__ = decoy
    # The premise: under the default it looks exactly like a conforming lister.
    assert tuple(inspect.signature(keyed_lister).parameters) == webapp._LISTER_SIGNATURE

    with pytest.raises(RuntimeError, match="cannot be handed an API key"):
        webapp._check_every_listable_provider_has_a_keyless_lister(
            {"ollama"}, {"ollama": keyed_lister}
        )


def test_the_module_body_runs_the_check(monkeypatch):
    """The predicate is only a gate if the module body actually calls it.

    The tests above call it themselves, and `sys.modules` caching means
    the real import already happened before any of them ran -- so deleting the
    module-level call leaves every one of them green. This re-executes
    `app.py`'s body into a fresh module object with a provider set the shipped
    `_MODEL_LISTERS` cannot satisfy, which raises only if that call site exists.

    The match names the check, not just the provider, and that is what makes the
    sentence above true: an unsatisfiable provider set fails the default-endpoint
    check too, so matching on `anthropic` alone passed with THIS call site
    deleted -- the sibling's error standing in for it.

    A separate module object rather than `importlib.reload`: reloading the live
    `verinote.web.app` would swap out the identity of `Config`-facing handlers
    and helper objects every other test in this session imported, and a reload
    that raises part-way (which is the point here) leaves it half-rebuilt.

    The set ADDS to the shipped one rather than replacing it, so this stays a
    test of the lister call site: dropping `openrouter` from listable would make
    the clearing-set check -- which runs before this one -- raise first, about a
    provider that clears the Base URL with no endpoint to name.
    """
    monkeypatch.setattr(
        verinote.config,
        "MODEL_LISTING_PROVIDERS",
        frozenset({"ollama", "openrouter", "anthropic"}),
    )
    name = "verinote_web_app_under_test"
    spec = importlib.util.spec_from_file_location(name, webapp.__file__)
    module = importlib.util.module_from_spec(spec)
    monkeypatch.setitem(sys.modules, name, module)

    with pytest.raises(RuntimeError, match=r"model lister.*anthropic"):
        spec.loader.exec_module(module)


# --- the Claude CLI alias picker: suggestions you can type past ---


def test_model_field_offers_cli_aliases_as_a_real_select(tmp_path, monkeypatch):
    """Curated does not mean typed. The CLI has no listing endpoint, so the
    options are curated from the adapter -- rendering them must not require (or
    attempt) a provider call, and must still be a control you pick from."""
    monkeypatch.setattr(
        webapp, "get_client", lambda _cfg: pytest.fail("claudecli has nothing to enumerate")
    )

    r = _ollama_client(tmp_path).get("/settings/model-field?provider=claudecli&model=opus")

    assert r.status_code == 200
    assert '<select name="model">' in r.text
    assert '<option value="opus" selected>opus — latest</option>' in unescape(r.text)
    assert 'type="text" name="model"' not in r.text  # choosing costs a click, not typing
    assert 'hx-trigger="load"' not in r.text  # nothing to lazily fetch


def test_model_field_keeps_the_cli_default_selectable(tmp_path):
    """An empty model means "pass no --model, let the CLI decide" -- this KB's
    shipped default. Without an option for it the browser auto-selects the first
    alias, so merely opening Settings and pressing Save would switch the KB to
    `fable` without anyone choosing it."""
    r = _ollama_client(tmp_path).get("/settings/model-field?provider=claudecli&model=")

    assert '<option value="" selected>CLI default (no --model)</option>' in r.text
    first_option = re.search(r"<option [^>]*>", r.text).group(0)
    assert 'value=""' in first_option  # and it is the one a browser lands on


def test_model_field_keeps_a_pinned_id_in_the_list_and_selected(tmp_path):
    """A pin is not an alias, but it is still what config.json says. Dropping it
    from the options would make the page report a model the KB is not using."""
    r = _ollama_client(tmp_path).get(
        "/settings/model-field?provider=claudecli&model=claude-opus-4-8"
    )

    assert (
        '<option value="claude-opus-4-8" selected>claude-opus-4-8 — pinned</option>'
        in unescape(r.text)
    )
    assert '<option value="" >CLI default (no --model)</option>' in r.text  # not selected


def test_model_field_alias_options_match_what_the_adapter_resolves(tmp_path):
    """A listed alias the adapter did not recognise would be offered as a choice
    and then silently mean something else.

    `_cli_model(alias) == alias` is NOT the property to assert: unknown strings
    are passed through unchanged, so any invented alias satisfies it. What only
    a real alias satisfies is absorbing a decorated display name of its family.
    """
    from verinote.llm.claude_cli_adapter import CLI_MODEL_ALIASES, _cli_model

    r = _ollama_client(tmp_path).get("/settings/model-field?provider=claudecli&model=")

    offered = re.findall(r'<option value="([^"]*)"[^>]*>([^<]*) — latest</option>', r.text)
    assert [value for value, _ in offered] == list(CLI_MODEL_ALIASES)
    for alias, _ in offered:
        assert _cli_model(f"Claude {alias.title()} 9.9") == alias


def test_model_field_swaps_to_a_text_input_for_a_full_id(tmp_path, monkeypatch):
    """The escape hatch sits beside the control, not inside it: the list stays a
    list, and a full id is still reachable in one click."""
    monkeypatch.setattr(
        webapp, "get_client", lambda _cfg: pytest.fail("claudecli has nothing to enumerate")
    )

    listed = _ollama_client(tmp_path).get("/settings/model-field?provider=claudecli&model=opus")
    typed = _ollama_client(tmp_path).get(
        "/settings/model-field?provider=claudecli&model=opus&custom=1"
    )

    assert "Enter a model id" in listed.text
    assert '<input type="text" name="model" value="opus"' in typed.text
    assert "<select" not in typed.text
    assert "passed to the CLI unchanged" in typed.text
    assert "Back to the list" in typed.text  # and the swap is reversible


def test_settings_page_renders_the_cli_picker_as_a_list(tmp_path):
    """`custom` is a view state the partial owns; the page render must never
    start in it. A Settings page that opened on a text input would put every
    user back to typing -- the exact thing the select replaced."""
    cfg = Config(
        root=tmp_path,
        db_path=tmp_path / "kb.sqlite",
        provider="claudecli",
        model="opus",
        api_key=None,
        base_url=None,
    )

    r = TestClient(create_app(cfg)).get("/settings")

    assert '<select name="model">' in r.text
    assert 'type="text" name="model"' not in r.text


def test_model_field_custom_never_downgrades_the_discovered_picker(tmp_path, monkeypatch):
    """A stray `custom=1` must not turn Ollama's discovered list into free text --
    that list is evidence, and there is no reason to type past it."""
    _FakeLister(models=["qwen3:8b"]).install(monkeypatch)

    r = _ollama_client(tmp_path).get(
        "/settings/model-field?provider=ollama&model=qwen3:8b&custom=1"
    )

    assert '<select name="model">' in r.text
    assert 'type="text" name="model"' not in r.text


def test_model_field_keeps_aliases_off_the_ollama_picker(tmp_path, monkeypatch):
    """Curated aliases and a discovered list are different claims; the Ollama
    field must never quietly gain entries no server reported."""
    _FakeLister(models=["qwen3:8b"]).install(monkeypatch)

    r = _ollama_client(tmp_path).get("/settings/model-field?provider=ollama&model=qwen3:8b")

    assert "<datalist" not in r.text
    for alias in ("fable", "sonnet"):
        assert f'value="{alias}"' not in r.text


def test_model_field_gives_no_aliases_to_a_plain_text_provider(tmp_path):
    r = _ollama_client(tmp_path).get("/settings/model-field?provider=anthropic&model=claude")

    assert "<datalist" not in r.text
    assert '<input type="text" name="model" value="claude" placeholder="model id">' in r.text


# --- the OpenRouter catalogue picker: grouped by what each entry advertises ---


_CATALOGUE = {
    "schema": ("openai/gpt-oss-20b:free", "z/vendor-schema"),
    "plain": ("a/vendor-plain",),
}


def _openrouter_lister(monkeypatch, **kwargs):
    """A stub OpenRouter catalogue: two advertising entries and one that is not."""
    models = _CATALOGUE["schema"] + _CATALOGUE["plain"]
    defaults = {
        "models": sorted(models),
        "structured_output_ids": frozenset(_CATALOGUE["schema"]),
    }
    defaults.update(kwargs)
    return _FakeLister(**defaults).install(monkeypatch, provider="openrouter")


def _optgroups(text: str) -> list[str]:
    return re.findall(r'<optgroup label="([^"]*)">', text)


def test_settings_defers_the_openrouter_catalogue_off_the_page_load(tmp_path, monkeypatch):
    """OpenRouter is listable now, so its Model field must behave like Ollama's:
    a working text input that lazily swaps itself for the picker, never a network
    call on the critical path of the page that also recovers a corrupt config."""
    monkeypatch.setitem(
        webapp._MODEL_LISTERS,
        "openrouter",
        lambda base_url, timeout: pytest.fail("/settings must not call the provider"),
    )

    r = _ollama_client(tmp_path, provider="openrouter", model="openai/gpt-oss-20b:free").get(
        "/settings"
    )

    assert r.status_code == 200
    assert 'hx-trigger="load"' in r.text
    assert 'name="model"' in r.text


def test_openrouter_model_field_groups_the_catalogue_by_advertised_structured_output(
    tmp_path, monkeypatch
):
    """The catalogue says which entries advertise structured output, and verinote
    always asks for a JSON-schema `response_format` -- so which group a model is
    in is the one thing about it that predicts a failure, and it has to be visible
    while choosing.

    Mutation: render the ids flat (drop the `structured_output_ids is none`
    branch's else arm) and the two `<optgroup>`s disappear while every option
    stays, so a user picks a non-advertising model with nothing said about it.
    """
    _openrouter_lister(monkeypatch)

    r = _ollama_client(tmp_path, provider="openrouter", model="openai/gpt-oss-20b:free").get(
        "/settings/model-field?provider=openrouter&model=openai/gpt-oss-20b:free"
    )

    assert r.status_code == 200
    assert '<select name="model">' in r.text
    assert _optgroups(r.text) == [
        "Advertises structured output",
        "Does not advertise structured output",
    ]
    advertising, unadvertising = r.text.split('<optgroup label="Does not advertise')
    for model in _CATALOGUE["schema"]:
        assert f'value="{model}"' in advertising
    for model in _CATALOGUE["plain"]:
        assert f'value="{model}"' in unadvertising
        assert f'value="{model}"' not in advertising


def test_ollama_model_field_stays_flat_because_its_listing_reports_no_capability(
    tmp_path, monkeypatch
):
    """`/api/tags` says nothing about structured output, so grouping Ollama's
    models would attribute a claim to a server that never made it.

    Mutation: collapse the two in `_model_field_context`
    (`listing.structured_output_ids or frozenset()`), which is the one-character
    version of treating "does not report it" as "reports none" -- every installed
    model then lands under "Does not advertise structured output".
    """
    _FakeLister(models=["llava:7b", "qwen3:8b"]).install(monkeypatch)

    r = _ollama_client(tmp_path).get("/settings/model-field?provider=ollama&model=qwen3:8b")

    assert '<select name="model">' in r.text
    assert _optgroups(r.text) == []
    assert "advertise" not in r.text


def test_openrouter_model_field_groups_a_catalogue_where_nothing_advertises(
    tmp_path, monkeypatch
):
    """A listing that reports the property and finds nothing advertising it is not
    a listing that does not report the property: the first is a catalogue saying
    "none of these", the second (Ollama) is silence. Both groups render whenever
    the property IS reported, so an empty "Advertises structured output" heading
    is itself the honest answer -- and this state is reachable, being exactly what
    OpenRouter renaming `structured_outputs` would produce.

    Mutation: `{% if structured_output_ids is none %}` -> `{% if not
    structured_output_ids %}` in `model_field.html`, which folds the empty
    frozenset into the `None` case; the picker goes flat and every model loses the
    heading that says its own entry does not advertise what verinote will ask it
    for. The two sibling grouping tests miss it -- OpenRouter's set is non-empty
    there and Ollama's is `None`, so neither tells truthiness from `is none`.
    """
    models = sorted(_CATALOGUE["schema"] + _CATALOGUE["plain"])
    _FakeLister(models=models, structured_output_ids=frozenset()).install(
        monkeypatch, provider="openrouter"
    )

    r = _ollama_client(tmp_path, provider="openrouter", model="openai/gpt-oss-20b:free").get(
        "/settings/model-field?provider=openrouter&model=openai/gpt-oss-20b:free"
    )

    assert r.status_code == 200
    assert '<select name="model">' in r.text
    assert _optgroups(r.text) == [
        "Advertises structured output",
        "Does not advertise structured output",
    ]
    advertising, unadvertising = r.text.split('<optgroup label="Does not advertise')
    for model in models:
        assert f'value="{model}"' in unadvertising
        assert f'value="{model}"' not in advertising


def test_openrouter_model_field_groups_a_catalogue_where_everything_advertises(
    tmp_path, monkeypatch
):
    """The mirror of the sibling above, for the other empty group.

    "Both groups are rendered whenever the property IS reported, even if one is
    empty" is a claim about either group, and only one half of it was pinned. A
    catalogue in which every entry advertises structured output leaves the SECOND
    group empty, and its heading still has to appear: dropping it would turn the
    remaining heading into an unlabelled section over the whole list, so a user
    reading it has no way to know a "does not advertise" group exists at all --
    and the next entry to land in it would arrive under a heading that had never
    been shown.

    Mutation: wrap either `<optgroup>` in a non-empty guard (for this one,
    `{% if models | reject('in', structured_output_ids) | list %}`). The sibling
    above catches that guard on the first group only, because its first group is
    the empty one; this catches it on the second. Neither of the two grouping
    tests with a mixed catalogue sees either, since both of their groups are
    non-empty.
    """
    models = sorted(_CATALOGUE["schema"] + _CATALOGUE["plain"])
    _FakeLister(models=models, structured_output_ids=frozenset(models)).install(
        monkeypatch, provider="openrouter"
    )

    r = _ollama_client(tmp_path, provider="openrouter", model="openai/gpt-oss-20b:free").get(
        "/settings/model-field?provider=openrouter&model=openai/gpt-oss-20b:free"
    )

    assert r.status_code == 200
    assert '<select name="model">' in r.text
    assert _optgroups(r.text) == [
        "Advertises structured output",
        "Does not advertise structured output",
    ]
    advertising, unadvertising = r.text.split('<optgroup label="Does not advertise')
    for model in models:
        assert f'value="{model}"' in advertising
        assert f'value="{model}"' not in unadvertising


def test_openrouter_model_field_names_openrouters_endpoint_not_ollamas(tmp_path, monkeypatch):
    """With a second listable provider, a hardcoded `OLLAMA_DEFAULT_BASE_URL`
    would tell an OpenRouter user the page dialled `http://localhost:11434`.

    Mutation: drop `_LISTABLE_DEFAULT_ENDPOINTS` and go back to the Ollama
    constant -- the positive assertion is what catches it, since the note would
    still render, just naming a host nothing here ever contacted.
    """
    fake = _openrouter_lister(monkeypatch)

    r = _ollama_client(tmp_path, provider="openrouter", model="openai/gpt-oss-20b:free").get(
        "/settings/model-field?provider=openrouter&model=openai/gpt-oss-20b:free"
    )

    assert fake.seen == [None]  # an unset Base URL stays unset in config
    assert "https://openrouter.ai/api/v1" in r.text
    assert "localhost:11434" not in r.text


def test_openrouter_model_field_note_makes_no_ollama_claim(tmp_path, monkeypatch):
    """The note under the picker is what stops the catalogue reading as account
    entitlement and the two groups reading as a measurement. Both claims are
    OpenRouter-specific and both are load-bearing: the request carried no key, so
    this is what the endpoint publishes rather than what this account may call,
    and the groups repeat `supported_parameters` rather than reporting a run.

    Mutation: the note's `{% if provider == 'ollama' %}` -> `{% if 1 %}`, one
    token no other test sees; the page then tells an OpenRouter user that N models
    are "installed on" `https://openrouter.ai/api/v1` and the entire keyless
    caveat disappears.
    """
    _openrouter_lister(monkeypatch)

    r = _ollama_client(tmp_path, provider="openrouter", model="openai/gpt-oss-20b:free").get(
        "/settings/model-field?provider=openrouter&model=openai/gpt-oss-20b:free"
    )

    assert "3 models listed by" in r.text
    assert "answered without an API key" in r.text
    assert "not a list of what your account can reach" in r.text
    # The sentence that keeps the grouping a repetition rather than a measurement.
    assert "has not run any of them" in r.text
    assert "installed on" not in r.text
    assert "ollama pull" not in r.text


def test_ollama_model_field_note_makes_no_openrouter_claim(tmp_path, monkeypatch):
    """The other half of that branch, so an inversion fails in both directions.
    Ollama's list is what is installed on a machine the user controls, where a
    missing model is fixed with `ollama pull`; the catalogue copy would tell that
    user their own server "answered without an API key" and that what it reported
    is a published list rather than what they pulled.

    Mutation: invert the note's `{% if provider == 'ollama' %}` the other way (to
    `{% if 0 %}`, or swap the two arms) -- the OpenRouter sibling above still
    passes, and only this catches it.
    """
    _FakeLister(models=["llava:7b", "qwen3:8b"]).install(monkeypatch)

    r = _ollama_client(tmp_path).get("/settings/model-field?provider=ollama&model=qwen3:8b")

    assert "2 models installed on" in r.text
    assert "http://localhost:11434" in r.text
    assert "answered without an API key" not in r.text
    assert "your account can reach" not in r.text
    assert "listed by" not in r.text


def test_openrouter_model_field_names_a_configured_proxy_not_the_literal_default(
    tmp_path, monkeypatch
):
    """The note must interpolate the endpoint actually dialled. Hardcoding
    `openrouter.ai` into the copy would pass the test above and then misreport
    every gateway or proxy deployment.

    Mutation: write the literal URL into `model_field.html` instead of
    `{{ endpoint }}` -- that mutation is invisible to the sibling test and fails
    only here.
    """
    fake = _openrouter_lister(monkeypatch)

    r = _ollama_client(tmp_path, provider="openrouter", model="openai/gpt-oss-20b:free").get(
        "/settings/model-field?provider=openrouter&model=openai/gpt-oss-20b:free"
        "&base_url=https://proxy.internal/v1"
    )

    assert fake.seen == ["https://proxy.internal/v1"]
    assert "proxy.internal" in r.text
    assert "openrouter.ai" not in r.text


def test_openrouter_model_field_keeps_a_configured_model_absent_from_the_catalogue(
    tmp_path, monkeypatch
):
    """config.json really does name this model, and the catalogue is a published
    list, not a verdict on the id. Dropping it would silently select a different
    model and make the page misreport the KB's own state.

    Mutation: drop the "keep the configured model" branch and the option vanishes
    while the picker still renders -- the browser then selects the first entry,
    so the field shows a model this KB is not configured for. Mutation 2: leave
    the label at Ollama's `— not installed` and an OpenRouter user is told a
    catalogue entry is an installation.
    """
    _openrouter_lister(monkeypatch)

    r = _ollama_client(tmp_path, provider="openrouter", model="vendor/private-model").get(
        "/settings/model-field?provider=openrouter&model=vendor/private-model"
    )

    assert (
        '<option value="vendor/private-model" selected>vendor/private-model'
        " — not in this catalogue</option>" in unescape(r.text)
    )
    assert "did not list" in r.text
    assert "not installed" not in r.text
    assert "ollama pull" not in r.text


def test_openrouter_model_field_falls_back_to_typing_when_the_catalogue_fails(
    tmp_path, monkeypatch
):
    """An empty picker would read as "the catalogue offers nothing" and leave the
    user unable to set a model at all.

    The real lister runs here, against an `urlopen` that refuses, rather than a
    stub raising `LLMError`: the mutation this exists for lives in the adapter.
    Mutation: return an empty `ModelListing` from
    `openrouter_adapter.list_models` on a transport error instead of raising --
    the page then renders the empty-catalogue copy, blaming OpenRouter for
    listing nothing when it was never reached. A stubbed lister would pass that
    mutation, since the stub is what raises.
    """

    def refuse(req, *, timeout):
        raise OSError("connection refused")

    monkeypatch.setattr("urllib.request.urlopen", refuse)

    r = _ollama_client(tmp_path, provider="openrouter", model="openai/gpt-oss-20b:free").get(
        "/settings/model-field?provider=openrouter&model=openai/gpt-oss-20b:free"
    )

    assert r.status_code == 200
    assert "<select" not in r.text
    assert '<input type="text" name="model" value="openai/gpt-oss-20b:free"' in r.text
    assert "Could not load the model list" in r.text
    assert "connection refused" in r.text
    assert "is reachable but" not in r.text


def test_openrouter_empty_catalogue_copy_is_not_ollamas(tmp_path, monkeypatch):
    """`ollama pull` fires for any listable provider whose listing came back
    empty, so the moment OpenRouter joined, this state told an OpenRouter user to
    run a command for a program they may not have installed.

    Mutation: reuse Ollama's string for both and `ollama pull` appears here.
    """
    _FakeLister(models=[], structured_output_ids=frozenset()).install(
        monkeypatch, provider="openrouter"
    )

    r = _ollama_client(tmp_path, provider="openrouter", model="").get(
        "/settings/model-field?provider=openrouter&model="
    )

    assert "ollama pull" not in r.text
    assert "installed" not in r.text
    assert "listed no models" in r.text
    assert "https://openrouter.ai/api/v1" in r.text
    assert '<input type="text" name="model"' in r.text


def test_openrouter_model_field_offers_no_blank_option(tmp_path, monkeypatch):
    """Unlike the CLI picker, a blank model here is not "no model": it falls
    through to `_MODEL_DEFAULTS['openrouter']`, so an option labelled as an
    absence would misreport the model this KB would actually use.

    `<option value=""` rather than a bare `value=""`, so this keeps asserting
    what its name says: an empty Base URL input carries `value=""` too, and a
    response that started including one would fail this test for a reason that
    has nothing to do with the picker's options.
    """
    _openrouter_lister(monkeypatch)

    r = _ollama_client(tmp_path, provider="openrouter", model="openai/gpt-oss-20b:free").get(
        "/settings/model-field?provider=openrouter&model=openai/gpt-oss-20b:free"
    )

    assert '<option value=""' not in r.text


def test_openrouter_model_field_puts_no_api_key_on_the_wire(tmp_path, monkeypatch):
    """The egress site itself, not a proxy for it.

    `test_model_field_never_constructs_a_provider_adapter` counts constructions,
    which is one step removed: a lister body that read a key directly would build
    nothing and still ship it. This runs the real `openrouter_adapter.list_models`
    against an intercepted `urlopen` with a key in `app.state.cfg`, and asserts the
    sentinel appears in none of the three places a request can carry it.

    Mutation: add `Authorization: Bearer {cfg.api_key}` anywhere on that path --
    which requires handing the lister a Config, so the import-time shape check is
    the other half of this -- and the header assertion fails.
    """
    sentinel = "sk-or-v1-SENTINEL-DO-NOT-SEND-0123456789"
    requests = []

    class _Catalogue:
        def __enter__(self):
            return self

        def __exit__(self, *exc):
            return None

        def read(self):
            entry = {"id": "a/b", "supported_parameters": ["structured_outputs"]}
            return json.dumps({"data": [entry]}).encode("utf-8")

    def fake_urlopen(req, *, timeout):
        requests.append(req)
        return _Catalogue()

    monkeypatch.setattr("urllib.request.urlopen", fake_urlopen)
    client = _ollama_client(
        tmp_path, provider="openrouter", model="a/b", api_key=sentinel
    )
    # The premise: the app really is holding the key while this request runs.
    assert client.app.state.cfg.api_key == sentinel

    r = client.get("/settings/model-field?provider=openrouter&model=a/b")

    assert r.status_code == 200
    assert '<select name="model">' in r.text  # it really did render the picker
    assert len(requests) == 1  # ...and really did dial the catalogue
    req = requests[0]
    assert sentinel not in req.full_url
    assert not [item for item in req.header_items() if sentinel in str(item)]
    assert req.data is None or sentinel.encode("utf-8") not in req.data


# --- switching to OpenRouter clears the Base URL, and says what it discarded ---


def test_switching_to_openrouter_clears_the_base_url(tmp_path, monkeypatch):
    """The Base URL field's only job is to point verinote somewhere other than
    the provider's own endpoint, so the value carried over from the provider
    being left is not an OpenRouter endpoint at all -- `http://localhost:11434`
    would be dialled as one. Clear it, and list against OpenRouter's own.

    Mutation: drop the rule in `model_field` and the Ollama endpoint is still in
    the field, and is what the catalogue is listed from.
    """
    fake = _openrouter_lister(monkeypatch)

    r = _ollama_client(tmp_path, base_url="http://localhost:11434").get(
        "/settings/model-field?provider_changed=1&provider=openrouter"
        "&model=&base_url=http://localhost:11434"
    )

    assert r.status_code == 200
    assert '<input type="text" name="base_url" value=""' in r.text
    assert fake.seen == [None]  # ...and the listing followed the clear


def test_switching_to_another_provider_keeps_the_typed_base_url(tmp_path, monkeypatch):
    """The falsifier for the test above. "Always clear" satisfies the OpenRouter
    request and destroys the endpoint of everyone pointing Ollama at another box
    -- and silently, since Save then writes it away.

    Mutation: clear unconditionally (drop the `provider == "openrouter"` clause)
    and this typed endpoint is gone.
    """
    fake = _FakeLister(models=["qwen3:8b"]).install(monkeypatch)

    r = _ollama_client(tmp_path, provider="openrouter", model="").get(
        "/settings/model-field?provider_changed=1&provider=ollama"
        "&model=&base_url=http://box.lan:11434"
    )

    assert r.status_code == 200
    assert '<input type="text" name="base_url" value="http://box.lan:11434"' in r.text
    assert fake.seen == ["http://box.lan:11434"]
    assert "cleared it" not in r.text  # nothing was discarded, so nothing is claimed


def test_settings_renders_a_saved_openrouter_base_url(tmp_path):
    """A page load is not a provider change. This KB's config.json really does
    name a proxy, and blanking the field on load would make the page misreport
    the KB's own state -- then destroy the value on the next Save.

    Mutation: move the clear into the template and this saved proxy renders as
    an empty field.

    NOT the route's `provider_changed` condition, though, and this test must not
    be read as covering it: `/settings` renders through `_settings`, which calls
    `_model_field_context` directly and never enters `model_field`, so the
    condition is not on this path at all and dropping it leaves this green. What
    that condition protects is the plain model-field GET -- the one the Base URL
    input itself fires -- and dropping it is caught by
    `test_openrouter_model_field_names_a_configured_proxy_not_the_literal_default`,
    where the clear would then take the configured proxy away before listing and
    the catalogue would be read from OpenRouter's own endpoint instead.
    """
    r = _ollama_client(
        tmp_path, provider="openrouter", model="a/b", base_url="https://proxy.internal/v1"
    ).get("/settings")

    assert r.status_code == 200
    assert '<input type="text" name="base_url" value="https://proxy.internal/v1"' in r.text
    assert "cleared it" not in r.text


def test_the_provider_change_response_carries_the_model_and_the_base_url(tmp_path, monkeypatch):
    """The swap has to land on a region holding both fields. Before this, the
    provider select targeted `#model-field`, which the Base URL input sits
    outside of -- so no provider change could reach it however the route decided.

    Mutation: leave the base URL input out of the wrapper and the form loses the
    field entirely after one provider change.
    """
    _openrouter_lister(monkeypatch)

    r = _ollama_client(tmp_path).get(
        "/settings/model-field?provider_changed=1&provider=openrouter"
        "&model=openai/gpt-oss-20b:free&base_url="
    )

    assert r.status_code == 200
    assert 'id="provider-fields"' in r.text
    assert 'name="model"' in r.text
    assert 'name="base_url"' in r.text


def test_a_base_url_change_re_renders_only_the_model_field(tmp_path, monkeypatch):
    """This response is swapped into a DOM that already holds the Base URL input,
    so it must not carry a second one: two `name="base_url"` inputs POST two
    values, and it would also re-render the field being typed into.

    Mutation: always render the wrapper and this count is 1, i.e. 2 in the page.
    """
    _FakeLister(models=["qwen3:8b"]).install(monkeypatch)

    r = _ollama_client(tmp_path).get(
        "/settings/model-field?provider=ollama&model=qwen3:8b&base_url=http://box.lan:11434"
    )

    assert r.status_code == 200
    assert 'id="model-field"' in r.text
    assert 'id="provider-fields"' not in r.text
    assert r.text.count('name="base_url"') == 0


def test_the_cleared_base_url_is_announced_not_silently_dropped(tmp_path, monkeypatch):
    """`POST /settings` maps `base_url or None`, so Save on an empty field
    DESTROYS the stored value. Between the clear and Save the page would
    otherwise show an empty Base URL while config.json still holds one -- the
    page asserting a state it does not have. Name what is going, and what would
    be dialled instead.

    Mutation: drop the note and the discard is silent.
    """
    _openrouter_lister(monkeypatch)

    r = _ollama_client(tmp_path, base_url="https://proxy.internal/v1").get(
        "/settings/model-field?provider_changed=1&provider=openrouter"
        "&model=&base_url=https://proxy.internal/v1"
    )

    assert r.status_code == 200
    assert "https://proxy.internal/v1" in r.text  # what is being discarded
    assert "https://openrouter.ai/api/v1" in r.text  # ...and what replaces it
    assert "Nothing has been saved yet" in r.text  # a proposal, not a done deal


def test_no_discard_note_when_there_was_nothing_to_discard(tmp_path, monkeypatch):
    """An empty Base URL loses nothing, so a note would claim a discard that did
    not happen -- and name a URL the user never had.

    Mutation: render the note unconditionally and it appears here with an empty
    `<code></code>` where the discarded URL should be.
    """
    _openrouter_lister(monkeypatch)

    r = _ollama_client(tmp_path).get(
        "/settings/model-field?provider_changed=1&provider=openrouter"
        "&model=openai/gpt-oss-20b:free&base_url="
    )

    assert r.status_code == 200
    assert "cleared it" not in r.text
    assert "Nothing has been saved yet" not in r.text


def test_settings_wires_the_provider_select_to_the_clearing_route(tmp_path):
    """Every test above calls `/settings/model-field?provider_changed=1` on the
    route directly. None of them reads the shipped page, so a broken `<select>`
    could sit there inert -- switched from Ollama to OpenRouter by hand in a
    browser -- while the whole suite around it stays green.

    One test, not two, even though it names two mutations: both are the same
    promise -- "the shipped select safely reaches the clearing route" -- read out
    of `provider_fields.html`'s comment ("this wrapper is served only for a
    provider change" via `hx-get`, "the provider select itself stays outside this
    div" via `hx-target`). `test_the_provider_change_response_carries_the_model_and_the_base_url`
    above bundles three assertions the same way, under one promise about what one
    response carries, rather than one test per assertion.

    Parses `<select name="provider" ...>` out of the page first and reads
    attributes from within that match, the way
    `test_the_provider_change_url_the_page_actually_uses_is_gated` in
    `test_same_origin_guard.py` was just hardened to: a single pattern anchored on
    `name="provider"` spanning the whole tag would go `None` the moment someone
    reorders the template's attributes, which is a tidy-up, not a bug, and would
    fail closed rather than reporting what actually broke.

    Mutation (M7): drop `provider_changed=1` from the select's `hx-get`. The
    route then returns `model_field.html` instead of the `#provider-fields`
    wrapper, so switching from Ollama to OpenRouter never clears
    `http://localhost:11434` -- Save stores it as OpenRouter's own endpoint.
    Caught by the `hx-get` assertion below.

    Mutation (M6): change the select's `hx-target` from `#provider-fields` to
    `#model-field`. `#model-field` sits inside the wrapper, so the wrapper --
    which carries its own `base_url` input and the discard note -- gets swapped
    into a region it contains, and the original input outside it survives. The
    page then holds two `name="base_url"` inputs, one still showing the old
    proxy, sitting directly beneath a note claiming the field was cleared; Save
    posts both. Caught by the `hx-target` assertion below.

    Mutation (M8): drop `[name='base_url']` from the select's `hx-include`.
    The route then receives `base_url=""` on every provider change, so
    `discarded_base_url` is always empty and no note renders -- even when the
    field held a real endpoint -- and `POST /settings` maps `base_url or
    None`, so Save silently destroys the stored value with nothing on the page
    ever having said so. This is worse than M6 or M7: those corrupt or
    misroute the response, but this one makes the loss invisible. Caught by
    the `hx-include` assertion below.

    Mutation (M9): drop `[name='model']` from the select's `hx-include`. Same
    class of bug, on the sibling field the same doc comment
    (`settings.html:54-55`) also names. The route then receives `model=""` on
    every provider change, so whatever model was configured is silently
    dropped from the field -- and unlike the Base URL there is no discard note
    for Model at all, so this loss is even less visible: `POST /settings`
    stores `model` as given (no `or None` guard), so Save overwrites
    config.json's model with an empty string with no warning shown anywhere.
    Caught by the `hx-include` assertion below.
    """
    r = _ollama_client(tmp_path).get("/settings")
    assert r.status_code == 200

    select = re.search(r'<select\b[^>]*\bname="provider"[^>]*>', r.text)
    assert select is not None, "the settings page no longer has a provider select"
    tag = select.group(0)

    hx_get = re.search(r'\bhx-get="([^"]*)"', tag)
    assert hx_get is not None, "the provider select no longer has an hx-get"
    assert "provider_changed=1" in hx_get.group(1)  # M7

    hx_target = re.search(r'\bhx-target="([^"]*)"', tag)
    assert hx_target is not None, "the provider select no longer has an hx-target"
    assert hx_target.group(1) == "#provider-fields"  # M6

    hx_include = re.search(r'\bhx-include="([^"]*)"', tag)
    assert hx_include is not None, "the provider select no longer has an hx-include"
    assert "[name='base_url']" in hx_include.group(1)  # M8
    assert "[name='model']" in hx_include.group(1)  # M9

    assert r.text.count('name="base_url"') == 1
    # ...and it must stay outside the region it targets, never inside it.
    assert select.start() < r.text.index('id="provider-fields"')


# --- the default-endpoint map is checked at import, like the lister table ---


@pytest.mark.parametrize(
    "providers, endpoints, expected",
    [
        ({"ollama", "openrouter"}, {"ollama": "http://localhost:11434"}, "openrouter"),
        (
            {"ollama"},
            {"ollama": "http://localhost:11434", "openrouter": "https://openrouter.ai/api/v1"},
            "openrouter",
        ),
    ],
)
def test_the_check_rejects_a_listable_provider_without_a_default_endpoint(
    providers, endpoints, expected
):
    """Both directions. A listable provider with no default endpoint would print
    some other provider's host as the one that answered; an endpoint for a
    provider that is not listable is an unused URL the next edit adopts without
    anyone re-checking it.

    Mutation: make the check a subset test (`set(providers) - set(endpoints)`)
    and the second case stops raising.
    """
    with pytest.raises(RuntimeError, match=expected):
        webapp._check_every_listable_provider_names_its_default_endpoint(providers, endpoints)


def test_the_check_rejects_a_blank_default_endpoint():
    """A present-but-empty entry passes the membership clause and then renders as
    `<code></code>` under the picker -- a name for the dialled host that names
    nothing.

    Mutation: drop the blank clause and this stops raising.
    """
    with pytest.raises(RuntimeError, match="ollama"):
        webapp._check_every_listable_provider_names_its_default_endpoint(
            {"ollama"}, {"ollama": "   "}
        )


def test_the_module_body_runs_the_default_endpoint_check(monkeypatch):
    """The predicate is only a gate if the module body calls it, and the tests
    above call it themselves -- so deleting the module-level call leaves them all
    green (`sys.modules` cached the real import long before they ran).

    Blanking the adapter constant rather than adding a provider is what isolates
    THIS check: an unknown listable provider is rejected by the lister check
    first, so it could never prove this one runs. A fresh module object rather
    than `importlib.reload`, for the reason its sibling states.
    """
    monkeypatch.setattr(
        verinote.llm.openrouter_adapter, "OPENROUTER_DEFAULT_BASE_URL", ""
    )
    name = "verinote_web_app_endpoint_check_under_test"
    spec = importlib.util.spec_from_file_location(name, webapp.__file__)
    module = importlib.util.module_from_spec(spec)
    monkeypatch.setitem(sys.modules, name, module)

    with pytest.raises(RuntimeError, match="openrouter"):
        spec.loader.exec_module(module)


# --- the Model field's per-provider copy is checked at import too ---


_ONE_CHAIN = "{% if provider == 'ollama' %}a{% elif provider == 'openrouter' %}b{% endif %}"


@pytest.mark.parametrize(
    "providers, source, expected",
    [
        ({"ollama", "openrouter"}, "{% if provider == 'ollama' %}a{% endif %}", "openrouter"),
        ({"ollama"}, _ONE_CHAIN, "openrouter"),
    ],
)
def test_the_check_rejects_a_listable_provider_without_model_field_copy(
    providers, source, expected
):
    """Both directions. A listable provider with no arm gets whatever the chain's
    catch-all says -- which is another provider's copy, the failure this replaces
    -- or, with no catch-all, nothing at all. An arm for a provider that is not
    listable is copy no render reaches, and the next edit to
    `MODEL_LISTING_PROVIDERS` would adopt it without anyone re-reading it.

    Mutation: drop the `unlistable` clause and the second case stops raising;
    drop `missing` and the first does.
    """
    with pytest.raises(RuntimeError, match=expected):
        webapp._check_every_listable_provider_has_model_field_copy(providers, source)


@pytest.mark.parametrize(
    "source",
    [
        # Lopsided: `openrouter` appears fewer times than `ollama` overall.
        _ONE_CHAIN + "{% if provider == 'ollama' %}c{% endif %}",
        # Balanced: three arms each across the file, and still one chain that
        # renders `openrouter` nothing, paid for by a duplicate arm elsewhere
        # that no render can reach.
        _ONE_CHAIN
        + "{% if provider == 'ollama' %}c{% endif %}"
        + "{% if provider == 'ollama' %}d{% elif provider == 'openrouter' %}e"
        "{% elif provider == 'openrouter' %}f{% endif %}",
    ],
)
def test_the_check_rejects_copy_written_at_only_some_of_the_branch_sites(source):
    """Membership alone would pass copy written at one chain and forgotten at the
    rest, which is the likely shape of the mistake: the template branches on the
    provider in four separate places, and a new provider needs an arm in each.

    The second source is why membership per chain is the rule rather than equal
    totals: the totals balance at three apiece, so a check that counted names
    across the whole file would pass it while the second chain rendered
    `openrouter` nothing at all.

    Mutation: compare totals per name instead of membership per chain -- the
    shape this check had -- and the second source stops raising.
    """
    with pytest.raises(RuntimeError, match="openrouter"):
        webapp._check_every_listable_provider_has_model_field_copy(
            {"ollama", "openrouter"}, source
        )


def test_the_check_rejects_a_template_with_no_provider_chain_left_in_it():
    """A per-chain rule says nothing about a file with no chains: flatten every
    arm back into shared prose and there is no chain left to be missing a
    provider, so the loop would pass a template that tells every provider's users
    the same sentence -- the exact misreport this check exists for.

    Mutation: delete the empty-`chains` guard and this stops raising.
    """
    source = "<p>{{ models|length }} models installed on {{ endpoint }}.</p>"

    with pytest.raises(RuntimeError, match="no longer branches on `provider`"):
        webapp._check_every_listable_provider_has_model_field_copy(
            {"ollama", "openrouter"}, source
        )


def test_the_check_rejects_a_second_arm_for_the_same_provider_in_one_chain():
    """"Exactly one arm" is the rule, and the second half of it is its own clause:
    a chain that tests the same provider twice has a second arm no render can
    reach, so the copy in it is never seen and never re-read. Membership alone
    would pass this -- both names are present -- which is why the counts inside
    the chain are checked as well as the names.

    Mutation: stop collecting `repeated` and this stops raising, while a
    maintainer could write a new sentence into a dead arm and see nothing.
    """
    source = (
        "{% if provider == 'ollama' %}a{% elif provider == 'openrouter' %}b"
        "{% elif provider == 'ollama' %}c{% endif %}"
    )

    with pytest.raises(RuntimeError, match=r"duplicate arms \['ollama'\]"):
        webapp._check_every_listable_provider_has_model_field_copy(
            {"ollama", "openrouter"}, source
        )


def test_the_check_does_not_count_a_provider_named_only_in_a_comment():
    """The header comment above the markup discusses these arms at length, so a
    name can appear there with no arm anywhere. Reading the file as text, four
    such mentions look exactly like four arms; parsing drops `{# ... #}` in the
    lexer, so they are not arms and the missing copy is still reported.

    Mutation: find the arms with `re.findall(r"provider == '([^']*)'")` over the
    source -- the shape this check had -- and this stops raising, because
    `newguy` then "appears" in all four chains without a line of copy written for
    it anywhere.
    """
    shipped = webapp._TEMPLATES.joinpath(webapp._MODEL_FIELD_TEMPLATE).read_text(
        encoding="utf-8"
    )
    source = shipped.replace(
        "#}",
        "\n   e.g. provider == 'newguy', provider == 'newguy',\n"
        "   provider == 'newguy', provider == 'newguy' #}",
        1,
    )
    assert source != shipped, "the header comment's closing `#}` moved"

    with pytest.raises(RuntimeError, match=r"missing \['newguy'\]"):
        webapp._check_every_listable_provider_has_model_field_copy(
            {"ollama", "openrouter", "newguy"}, source
        )


def test_a_comment_that_illustrates_an_arm_does_not_brick_the_check():
    """The other direction of the same property, and the live one: the header
    comment invites quoting one of these tests to explain it, and doing so must
    not stop the app from starting.

    Mutation: find the arms with `re.findall(r"provider == '([^']*)'")` over the
    source and this raises -- `ollama` then "appears" five times to
    `openrouter`'s four, so the app refuses to start and the error names
    `openrouter`, whose four arms are all present and correct.
    """
    shipped = webapp._TEMPLATES.joinpath(webapp._MODEL_FIELD_TEMPLATE).read_text(
        encoding="utf-8"
    )
    source = shipped.replace(
        "#}", "\n   Each chain opens with provider == 'ollama'. #}", 1
    )
    assert source != shipped, "the header comment's closing `#}` moved"

    webapp._check_every_listable_provider_has_model_field_copy(
        {"ollama", "openrouter"}, source
    )


def test_the_check_catches_reverting_the_shipped_arms_to_a_catch_all_else():
    """The realistic regression, run against the shipped file rather than a
    fixture: someone folds `{% elif provider == 'openrouter' %}` back into
    `{% else %}` because two providers make the elif look redundant. Every chain
    then hands Ollama's copy -- "installed", "ollama pull" -- to whatever joins
    next, which is the failure the whole check exists to stop.

    Mutation: return early from the check when a chain has any arm at all, rather
    than requiring each listable provider to have one, and this stops raising.
    """
    shipped = webapp._TEMPLATES.joinpath(webapp._MODEL_FIELD_TEMPLATE).read_text(
        encoding="utf-8"
    )
    source = shipped.replace("{% elif provider == 'openrouter' %}", "{% else %}")
    assert source != shipped, "the shipped arms no longer read `elif provider == ...`"

    with pytest.raises(RuntimeError, match=r"missing \['openrouter'\]"):
        webapp._check_every_listable_provider_has_model_field_copy(
            {"ollama", "openrouter"}, source
        )


def test_the_check_finds_the_inline_chain_inside_the_option_tag():
    """Four chains, not three. One of them is inline in an `<option>` tag rather
    than a block of its own, and it carries the sentence that says why a
    configured model is absent -- "not installed" versus "not in this catalogue".
    A walk that only found block-level chains would let that one drift.

    Mutation: skip `nodes.If` whose body is not a block and this returns three
    chains, so the inline arms could be reverted to an `{% else %}` unnoticed.
    """
    shipped = webapp._TEMPLATES.joinpath(webapp._MODEL_FIELD_TEMPLATE).read_text(
        encoding="utf-8"
    )
    anchor = "{% if provider == 'ollama' %} — not installed"
    inline_lineno = shipped[: shipped.index(anchor)].count("\n") + 1

    chains = webapp._model_field_provider_chains(shipped)

    assert [names for _, names in chains] == [["ollama", "openrouter"]] * 4
    assert inline_lineno in [lineno for lineno, _ in chains]


def test_the_walk_ignores_chains_that_do_not_test_the_provider():
    """Only chains that branch on `provider` are places a provider needs an arm.
    The template's `{% if models %}` and `{% if structured_output_ids is none %}`
    are not, and dragging them in would demand arms where the distinction being
    made is not per-provider at all.

    Mutation: treat every `nodes.If` as a provider chain and this returns extra
    entries with no names, which the check would then report as missing both
    providers -- the shipped template would not import.
    """
    source = "{% if models %}x{% endif %}" + _ONE_CHAIN

    assert webapp._model_field_provider_chains(source) == [(1, ["ollama", "openrouter"])]


def test_a_compound_guard_still_counts_as_that_providers_arm():
    """`_provider_names_compared` walks every `Compare` node inside the test, not
    just the test node itself, so an arm guarded by `provider == 'x' and models`
    -- rather than a bare `provider == 'x'` -- is still `x`'s arm, not a chain
    the walk fails to recognise.

    One test rather than two: the chain-level assertion pins that `openrouter`
    is attributed to *that* arm specifically (not merely that something passed),
    and the check-level call is the positive-acceptance case the surrounding
    tests lack -- every other guard test here asserts a raise, which a walk that
    cannot see the compound arm would still produce (a chain missing an arm
    always raises), so a rejection test cannot tell a correctly-attributed
    compound arm apart from an invisible one. Passing is the only outcome a
    narrowed walk cannot fake, and both assertions are cheap enough on one
    `source` that splitting them would only duplicate the fixture.

    Mutation: narrow the walk in `_provider_names_compared` to
    `for node in (test,):` and the compound `openrouter` arm goes unseen, so the
    chain reads as `['ollama']` -- missing `openrouter` -- and both assertions
    fail: the chain list no longer matches, and the check raises.
    """
    source = (
        "{% if provider == 'ollama' %}a"
        "{% elif provider == 'openrouter' and models %}b"
        "{% endif %}"
    )

    assert webapp._model_field_provider_chains(source) == [(1, ["ollama", "openrouter"])]
    webapp._check_every_listable_provider_has_model_field_copy(
        {"ollama", "openrouter"}, source
    )


def test_the_module_body_runs_the_model_field_copy_check(tmp_path, monkeypatch):
    """The predicate is only a gate if the module body calls it; the tests above
    call it themselves, so deleting the module-level call leaves them green.

    Doctoring a copy of the shipped template is what isolates THIS check: it is
    the only import-time check that reads that file, so nothing else can raise
    first. `resources.files` is redirected rather than the repo's template being
    edited, so the real one is never touched and a crashed run cannot leave it
    mutated. A fresh module object rather than `importlib.reload`, for the reason
    its sibling states.
    """
    shipped = webapp._TEMPLATES.joinpath(webapp._MODEL_FIELD_TEMPLATE).read_text(
        encoding="utf-8"
    )
    partials = tmp_path / "templates" / "partials"
    partials.mkdir(parents=True)
    (partials / "model_field.html").write_text(
        shipped.replace("provider == 'openrouter'", "provider == 'ollama'"), encoding="utf-8"
    )
    real_files = importlib.resources.files
    monkeypatch.setattr(
        importlib.resources,
        "files",
        lambda package: tmp_path if package == "verinote.web" else real_files(package),
    )
    name = "verinote_web_app_copy_check_under_test"
    spec = importlib.util.spec_from_file_location(name, webapp.__file__)
    module = importlib.util.module_from_spec(spec)
    monkeypatch.setitem(sys.modules, name, module)

    with pytest.raises(RuntimeError, match="openrouter"):
        spec.loader.exec_module(module)


# --- which providers clear the Base URL is checked at import too ---


def test_the_check_rejects_a_clearing_provider_that_cannot_name_its_replacement():
    """The note the clear renders promises what verinote would dial instead, and
    `_model_field_context` fills `endpoint` only for a listable provider -- so a
    clearing provider outside `MODEL_LISTING_PROVIDERS` renders that promise
    around an empty `<code></code>`. Same misreport the blank-endpoint clause
    rejects, reached by a one-line edit to either set instead.

    Mutation: drop the check body and this stops raising.
    """
    with pytest.raises(RuntimeError, match="anthropic"):
        webapp._check_every_clearing_provider_can_name_its_replacement(
            {"anthropic"}, {"ollama", "openrouter"}
        )


def test_the_check_allows_a_listable_provider_that_does_not_clear():
    """The falsifier for the test above, and the reason this one is a subset test
    and not the symmetric difference its three siblings use: Ollama is listable
    and must NOT clear, since clearing on every switch would wipe an endpoint the
    user typed. Asserted here rather than left to import, so the rule is stated
    where a maintainer copying a sibling would read it.

    Mutation: make it `set(clearing) ^ set(listable)` and this raises.
    """
    webapp._check_every_clearing_provider_can_name_its_replacement(
        {"openrouter"}, {"ollama", "openrouter"}
    )


def test_the_module_body_runs_the_clearing_provider_check(monkeypatch):
    """The predicate is only a gate if the module body calls it, and the tests
    above call it themselves -- so deleting the module-level call leaves them
    green (`sys.modules` cached the real import long before they ran).

    Narrowing `MODEL_LISTING_PROVIDERS` is the only lever that can reach this
    check, because `_BASE_URL_CLEARING_PROVIDERS` is a literal rebuilt by the
    exec below. That narrowing also breaks the three table checks, so what
    isolates THIS one is that its call site runs before theirs and that the match
    below is a phrase only its message contains: move the call after them, or
    delete it, and the error that arrives is the lister table's instead. A fresh
    module object rather than `importlib.reload`, for the reason its siblings
    state.
    """
    monkeypatch.setattr(verinote.config, "MODEL_LISTING_PROVIDERS", frozenset({"ollama"}))
    name = "verinote_web_app_clearing_check_under_test"
    spec = importlib.util.spec_from_file_location(name, webapp.__file__)
    module = importlib.util.module_from_spec(spec)
    monkeypatch.setitem(sys.modules, name, module)

    with pytest.raises(RuntimeError, match="clears the Base URL"):
        spec.loader.exec_module(module)


@contextmanager
def _broken_override(path: Path, mode: str):
    """Make a saved prompt override unreadable, the two ways these tests use.

    Copy per test file rather than a shared import: the repo keeps its chmod
    probe local (`tests/test_cloud_adapters.py`). The skip is a RUNTIME one after
    an actual read attempt, not a `geteuid` marker, so it also covers a
    filesystem that ignores the mode bit.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    if mode == "non_utf8":
        path.write_bytes(b"at most {max_facts} facts \xff\xfe and a bad byte")
        yield
        return
    path.write_text("at most {max_facts} facts\n", encoding="utf-8")
    path.chmod(0o000)
    try:
        try:
            path.read_text(encoding="utf-8")
        except PermissionError:
            pass
        else:
            pytest.skip("this user reads straight through mode 0o000")
        yield
    finally:
        # `exists()` because a caller may legitimately REMOVE the file inside
        # the block (`tests/test_web.py`'s reset test does).
        if path.exists():
            path.chmod(0o600)


def _prompts_client(tmp_path) -> TestClient:
    """`_client`'s app, but surfacing a server error as a status, not a raise.

    `_client` builds `TestClient(app)`, i.e. `raise_server_exceptions=True`, under
    which an unhandled render failure comes back as the exception itself rather
    than the 500 these tests are about. That would make a mutant look like an
    error instead of a status regression.
    """
    cfg = Config(
        root=tmp_path, db_path=tmp_path / "kb.sqlite",
        provider="anthropic", model="m", api_key=None, base_url=None,
    )
    return TestClient(create_app(cfg), raise_server_exceptions=False)


@pytest.mark.parametrize("mode", ["non_utf8", "chmod"])
def test_a_broken_extraction_limit_hint_is_extraction_failed_not_analysis_failed(
    tmp_path, monkeypatch, fake_client, mode
):
    """An unreadable override is this machine's condition, not the provider's.

    `_extraction_schema_hint(cfg)` is an argument expression inside the worker's
    `try`, so without #539's clause the read failure lands on the generic handler
    and the job says `analysis failed` — blaming the analysis for a file the user
    saved here. Measured without the clause, the message is `analysis failed: `
    plus whatever the read raised: the codec error under `non_utf8`,
    `[Errno 13] Permission denied: <the override>` under `chmod`. Neither of the
    two assertions at the end of this test holds of such a string (the
    `startswith` one aborts first, so only it is observed failing; the `not in`
    one is false of the same message). The `not in` one is spelled out because
    the misattribution, not the wording, is what #539 and #474 are about.
    """
    from verinote.prompts.library import prompt_override_path

    monkeypatch.setattr(
        webapp,
        "get_client",
        lambda cfg: fake_client([ExtractedFact("X", "is_a", "Y", 0.9)]),
    )
    c = _client(tmp_path)

    with _broken_override(prompt_override_path(tmp_path, "extraction-limit-hint"), mode):
        upload = c.post(
            "/sources",
            files={"file": ("note.txt", b"some text", "text/plain")},
            follow_redirects=False,
        )
        assert upload.status_code == 303

        def failed():
            jobs = c.app.state.store.source_extraction_jobs()
            assert jobs and jobs[0]["status"] == "failed"

        _wait_for(failed)
        message = c.app.state.store.source_extraction_jobs()[0]["message"]

    assert message.startswith(
        "extraction failed: prompt extraction-limit-hint could not be loaded"
    )
    assert "analysis failed" not in message


@pytest.mark.parametrize("mode", ["non_utf8", "chmod"])
def test_the_prompts_page_survives_the_override_it_exists_to_repair(tmp_path, mode):
    """The page you go to in order to fix a broken override must open (#539).

    200 and not the 409 of `sidecar_unreadable.html`, `config_corrupt.html` and
    `credentials_corrupt.html`: those three replace the page the caller asked
    for with a refusal template and end there, while this one is still a working
    page. The selector and reset assertions are what that distinction turns on,
    so they are pinned here rather than left implied by the status: the selector
    still routes to the other prompts, and the broken prompt itself still
    carries a control that repairs it. What this GET does not carry is that
    prompt's text or its Save form — see
    `test_a_broken_override_page_offers_no_editor_to_overwrite_it_with`. A
    refused save POST does get a textarea, holding the bytes that request
    submitted and nothing else (#545).
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)

    with _broken_override(prompt_override_path(tmp_path, "extraction"), mode):
        for url in ("/prompts", "/prompts?prompt=extraction"):
            r = c.get(url)
            assert r.status_code == 200, url
            assert "prompt extraction could not be loaded" in r.text, url
            # Still a delivered page: the selector lists the other prompts, so the
            # user can reach one that works and see which file is broken.
            assert 'value="query-translation"' in r.text, url
            assert 'value="ask-fallback"' in r.text, url
            # And it can still be repaired from here: `delete_prompt_override`
            # only `unlink()`s, so reset survives a file this process could not
            # read.
            assert "Reset to default" in r.text, url
            assert 'action="/prompts/reset"' in r.text, url
            assert 'name="prompt_id" value="extraction"' in r.text, url
            # The copy says the button "deletes the override file named above",
            # so the page has to name it. Measured: deleting that span reddens
            # this assertion for `non_utf8` and not for `chmod`, because the
            # codec error in the banner carries no path while the
            # `PermissionError` does. So for a non-UTF-8 user the span is the
            # whole of what tells them which file they are about to destroy.
            assert str(prompt_override_path(tmp_path, "extraction")) in r.text, url


def test_a_broken_override_does_not_blank_the_other_prompts(tmp_path):
    """The new clause must not swallow the page, nor the `unknown prompt` 400.

    Green before #539 as well — this row pins what must NOT change. The
    mutations these assertions exist for: a `_prompts_page` that returned
    `prompt=None` on the SUCCESS path too would blank the editor; folding the
    existing `except PromptError` into the new clause would turn `unknown
    prompt: nope` from 400 into 200; gating the reset control on `prompt is
    None` instead of on `reset_only` would offer to reset `nope`, an id
    `prompt_definition` has already rejected.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)

    with _broken_override(prompt_override_path(tmp_path, "extraction"), "non_utf8"):
        good = c.get("/prompts?prompt=query-translation")

        assert good.status_code == 200
        assert 'name="prompt_text"' in good.text  # the editor is still there
        # The query parameter is `prompt` (not `prompt_id`), and it is honoured:
        # this is query-translation's own text, not the broken default's.
        assert "<code>query-translation</code>" in good.text
        assert "You translate a natural-language question" in good.text
        assert "You are a fact extractor" not in good.text

        unknown = c.get("/prompts?prompt=nope")

        assert unknown.status_code == 400
        assert "unknown prompt: nope" in unknown.text
        # `prompt is None` for `nope` too, and for `nope` there is nothing to
        # reset — no definition and no file. The reset affordance keys off
        # `reset_only`, not off `prompt`, which is what keeps them apart.
        assert 'action="/prompts/reset"' not in unknown.text


@pytest.mark.parametrize("mode", ["non_utf8", "chmod"])
def test_a_refused_save_keeps_its_400_and_still_says_why(tmp_path, mode):
    """The load banner is composed with the caller's, and does not relabel a 400.

    Two independent properties, one per mutation (measured). A hardcoded 200 in
    the new clause would answer 200 to a request the app declined to perform,
    while the banner still says "prompt text is required". A clause that
    ASSIGNED `error` instead of composing it keeps the 400 and drops that
    string, telling the user the wrong reason their save failed.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)

    with _broken_override(prompt_override_path(tmp_path, "extraction"), mode):
        r = c.post(
            "/prompts",
            data={"prompt_id": "extraction", "prompt_text": ""},
            follow_redirects=False,
        )

    assert r.status_code == 400
    assert "prompt text is required" in r.text
    assert "prompt extraction could not be loaded" in r.text


@pytest.mark.parametrize("mode", ["non_utf8", "chmod"])
def test_a_broken_override_can_be_reset_from_the_page_that_names_it(tmp_path, mode):
    """#539's opening symptom: no way out of a broken override from inside the UI.

    Serving 200 with a banner answers the issue's non-regression list but not
    that complaint — a page that names the broken file and offers nothing that
    acts on it is still a dead end. Reset survives an unreadable override —
    `delete_prompt_override` only `unlink()`s and never reads it — so reset is
    the repair this page carries.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "extraction")

    with _broken_override(override, mode):
        broken = c.get("/prompts?prompt=extraction")
        assert broken.status_code == 200
        assert 'action="/prompts/reset"' in broken.text

        reset = c.post(
            "/prompts/reset",
            data={"prompt_id": "extraction"},
            follow_redirects=False,
        )

        assert reset.status_code == 303
        assert not override.exists()

    healthy = c.get("/prompts?prompt=extraction")

    assert healthy.status_code == 200
    assert "could not be loaded" not in healthy.text
    assert 'name="prompt_text"' in healthy.text  # the editor is back
    assert "You are a fact extractor" in healthy.text  # showing the default again


@pytest.mark.parametrize("mode", ["non_utf8", "chmod"])
def test_a_broken_override_page_offers_no_editor_to_overwrite_it_with(tmp_path, mode):
    """Reset, deliberately, and NOT a textarea seeded with the default.

    A Save form pre-filled with default text turns one careless click into a
    silent overwrite of a file this process could not read — the user's own
    customisation, gone, with nothing having displayed it. So on a GET the
    broken prompt's page carries the reset control and no editor, and the
    opening line of the default text is not in the response.

    Unchanged by the echo-back editor #545 added, and this row is what pins that
    it is unchanged: that block is gated on `prompt_text is not none`, which a
    GET never sets, and what it can show is the bytes of the request being
    answered, never a default. Mutate the gate to `not prompt` and this test
    reddens.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)

    with _broken_override(prompt_override_path(tmp_path, "extraction"), mode):
        r = c.get("/prompts?prompt=extraction")

    assert r.status_code == 200
    assert 'name="prompt_text"' not in r.text
    assert "Save prompt" not in r.text
    assert "You are a fact extractor" not in r.text


@pytest.mark.parametrize("kind", ["non_utf8", "chmod"])
def test_no_reset_is_offered_when_the_override_is_not_what_failed(
    tmp_path, monkeypatch, kind
):
    """This page's one destructive control must not be offered on a guess.

    `get_prompt` reads the packaged default as well as the KB override, and the
    clause above it is deliberately broad, so a load failure does not by itself
    say which file broke. Both states below were driven against a REAL damaged
    install (packaged `defaults/extraction.md` at `0o000`) before the gate
    existed: with nothing saved the page offered a Reset that deleted nothing
    and fixed nothing, and with a healthy override beside it the same button
    destroyed the user's file and left the page just as broken. The exception is
    raised here rather than by corrupting the installed package, because what is
    under test is the gate, which reads the disk and not the exception.
    """
    from verinote.prompts import library
    from verinote.prompts.library import prompt_override_path

    # The state this stands in for exists because `get_prompt` reads the
    # packaged default too: `prompts/library.py` calls `default_prompt_text`
    # before it looks at the override. (It is not the only way into the clause —
    # see `test_an_unreadable_prompts_directory_still_renders_the_page` — but it
    # is the one this test describes.) Couple the two, so this test cannot
    # quietly outlive that read.
    reached = []
    real_default_text = library.default_prompt_text

    def spy_default_text(prompt_id):
        reached.append(prompt_id)
        return real_default_text(prompt_id)

    monkeypatch.setattr(library, "default_prompt_text", spy_default_text)
    library.get_prompt(tmp_path, "extraction")

    assert reached == ["extraction"]

    if kind == "non_utf8":
        exc = UnicodeDecodeError("utf-8", b"\xff\xfe", 0, 1, "invalid start byte")
    else:
        exc = PermissionError(13, "Permission denied")

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "extraction")

    def raise_it(root, prompt_id):
        raise exc

    monkeypatch.setattr(webapp, "get_prompt", raise_it)

    nothing_saved = c.get("/prompts?prompt=extraction")

    assert nothing_saved.status_code == 200
    assert "prompt extraction could not be loaded" in nothing_saved.text
    assert 'action="/prompts/reset"' not in nothing_saved.text

    override.parent.mkdir(parents=True, exist_ok=True)
    override.write_text("My own extraction prompt.\n", encoding="utf-8")

    healthy_override = c.get("/prompts?prompt=extraction")

    assert healthy_override.status_code == 200
    assert 'action="/prompts/reset"' not in healthy_override.text
    assert override.read_text(encoding="utf-8") == "My own extraction prompt.\n"


def test_no_reset_is_offered_for_a_directory_at_the_override_path(tmp_path, monkeypatch):
    """The gate answers about a file `get_prompt` would open, not about any path.

    `get_prompt` opens the override only behind `override_path.is_file()`, so a
    directory there is not something it ever read, and `delete_prompt_override`
    could not remove it either — `unlink()` on a directory raises. Without the
    same guard on the gate the read fails, the page takes that for "unreadable",
    and it offers a Reset whose POST is a bare 500.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "extraction")
    override.mkdir(parents=True)

    def raise_it(root, prompt_id):
        raise PermissionError(13, "Permission denied")

    monkeypatch.setattr(webapp, "get_prompt", raise_it)

    r = c.get("/prompts?prompt=extraction")

    assert r.status_code == 200
    assert "prompt extraction could not be loaded" in r.text
    assert 'action="/prompts/reset"' not in r.text
    assert override.is_dir()  # the GET left it alone


def test_an_unreadable_prompts_directory_still_renders_the_page(tmp_path):
    """The gate must not raise: it runs inside the handler that builds this page.

    `Path.is_file()` propagates a `PermissionError` from an unreadable parent
    directory rather than answering False, and `get_prompt` reaches it too — it
    guards its own read with `is_file()` — so this state gets to the load
    failure clause with no corrupt file and no damaged install. The clause then
    calls the gate, from a line inside no `try` of its own. Hoisting the gate's
    `is_file()` guard above its `try` makes the raise escape both, and
    `GET /prompts` answers 500: the page #539 exists to keep alive, killed by
    the guard added to make its Reset control honest. Measured over every test
    in this file, nothing else catches that mutation, which is why this test is
    here.

    The POST from this same state is driven by
    `test_a_reset_that_cannot_unlink_is_a_page_not_a_crash[dir_0o000_no_override]`
    (#545): `delete_prompt_override` calls `path.exists()`, which raises the same
    way, and that POST now renders a page naming the file it could not delete
    instead of escaping the handler. What is pinned HERE is the GET, and 200 is
    what it answers today.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    prompts_dir = prompt_override_path(tmp_path, "extraction").parent
    prompts_dir.mkdir(parents=True, exist_ok=True)
    prompts_dir.chmod(0o000)
    try:
        try:
            (prompts_dir / "extraction.md").is_file()
        except PermissionError:
            pass
        else:
            pytest.skip("this user stats straight through mode 0o000")

        r = c.get("/prompts?prompt=extraction")

        assert r.status_code == 200
        assert "prompt extraction could not be loaded" in r.text
        assert 'value="query-translation"' in r.text  # the selector still routes away
    finally:
        prompts_dir.chmod(0o700)


def test_a_refusal_over_an_invalid_stored_override_keeps_both_reasons(tmp_path):
    """A refused save has two things to say when the page cannot load either.

    Reachable with no write failure at all: a stored override that is READABLE
    and fails validation — someone hand-edited `query-translation.md` and lost
    `{qid}`. `_prompts_page` then takes its `except PromptError` branch, which
    used to REPLACE the caller's `error` with its own. The page told the user
    their submitted text must include `{qid}` while the complaint was about the
    file on disk, and the reason their save was actually refused was gone.

    Mutation: put `str(exc)` back in that branch and `prompt text is required`
    leaves the page.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "query-translation")
    override.parent.mkdir(parents=True, exist_ok=True)
    override.write_text("Translate the question into Datalog.\n", encoding="utf-8")

    r = c.post(
        "/prompts",
        data={"prompt_id": "query-translation", "prompt_text": "   "},
        follow_redirects=False,
    )

    assert r.status_code == 400
    assert "prompt text is required" in r.text  # the route's own refusal
    assert "{qid}" in r.text  # and why the page could not be loaded either


@pytest.mark.parametrize("route", ["/prompts", "/prompts/reset"])
def test_an_unknown_prompt_id_is_reported_once_not_twice(tmp_path, route):
    """Composing the two diagnoses must not concatenate a string with itself.

    An unknown `prompt_id` raises the same `PromptError` in the route's own
    library call and again in `_prompts_page`'s `get_prompt`, out of the same
    `prompt_definition` lookup. Composing them unconditionally printed
    `unknown prompt: nope; unknown prompt: nope` on the most ordinary error page
    these routes have.

    Equality on the banner text, not a substring: `unknown prompt: nope` is a
    substring of the duplicate, so every other assertion in this file — and
    `test_prompt_routes_reject_unknown_key`, which drives this exact request —
    passes on it. Mutation: drop `error == load_error` from that branch.
    """
    c = _prompts_client(tmp_path)

    r = c.post(
        route,
        data={"prompt_id": "nope", "prompt_text": "No."},
        follow_redirects=False,
    )

    assert r.status_code == 400
    banner = re.search(r'<p class="error" role="alert">(.*?)</p>', r.text, re.S)
    assert banner is not None
    assert unescape(banner.group(1)).strip() == "unknown prompt: nope"


@pytest.mark.parametrize("mode", ["non_utf8", "chmod"])
def test_a_refused_save_keeps_the_typed_text_when_the_prompt_cannot_load(tmp_path, mode):
    """A refusal must hand the user's own paragraph back, not eat it.

    `{% if prompt %}` collapses the editor whenever `get_prompt` fails, so a
    save refused over a broken override took the submitted text down with the
    textarea — retype it or lose it. The loss is not caused by the write: it
    lands on any refused POST whose page cannot load, which is why this row is
    red on its own.

    `query-translation` because its required `{qid}` makes a NON-EMPTY text
    refusable; an empty-text refusal preserves nothing observable.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    submitted = "Return a query for the supplied question."

    with _broken_override(prompt_override_path(tmp_path, "query-translation"), mode):
        r = c.post(
            "/prompts",
            data={"prompt_id": "query-translation", "prompt_text": submitted},
            follow_redirects=False,
        )

    assert r.status_code == 400
    assert submitted in r.text  # the paragraph they typed
    assert 'name="prompt_text"' in r.text  # in a field they can resubmit
    assert "{qid}" in r.text  # still told why it was refused
    assert "prompt query-translation could not be loaded" in r.text


def test_a_save_into_an_unwritable_prompts_dir_is_a_page_not_a_crash(tmp_path):
    """The bare 500 #545 opens with: `mkdir`/`write_text` raise outside `PromptError`.

    `policy/prompts` at `0o500` with NO override file yet — the state the issue
    measured. With one present the save succeeds, because `write_text` truncates
    an existing inode and needs write permission on the FILE, not on the
    directory (measured; `..._reset_that_cannot_unlink...[dir_0o500_readable_override]`
    pins that it still does).

    Here `get_prompt` succeeds — `is_file()` answers False through a `0o500`
    directory — so the page is the ordinary editor and the text comes back in
    it. What is new is the banner and the status.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "extraction")
    override.parent.mkdir(parents=True, exist_ok=True)
    override.parent.chmod(0o500)
    sentinel = "Extract every fact you are sure of, and nothing else."
    try:
        probe = override.parent / ".write-probe"
        try:
            probe.write_text("x", encoding="utf-8")
        except PermissionError:
            pass
        else:
            probe.unlink()
            pytest.skip("this user writes straight through mode 0o500")

        r = c.post(
            "/prompts",
            data={"prompt_id": "extraction", "prompt_text": sentinel},
            follow_redirects=False,
        )
    finally:
        override.parent.chmod(0o700)

    assert r.status_code == 500  # plan §3.1: the write did not happen
    assert f"prompt extraction could not be saved to {override}:" in r.text
    assert sentinel in r.text
    assert 'name="prompt_text"' in r.text
    assert not override.exists()  # outside the restricted directory: `stat()` needs it


def test_a_save_over_an_unwritable_override_keeps_the_typed_text(tmp_path):
    """A failed write over an unreadable override: banner AND textarea.

    `chmod` only and never `non_utf8`: a non-UTF-8 override is fully writable,
    so that save succeeds with 303 and the row would prove nothing.

    Here `get_prompt` fails as well as the write, so `prompt` is None and the
    pre-existing editor is gone — the submitted text can only come back through
    the block #545 added to `prompts.html`.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "extraction")
    sentinel = "at most {max_facts} facts, and this sentence."

    with _broken_override(override, "chmod"):
        r = c.post(
            "/prompts",
            data={"prompt_id": "extraction", "prompt_text": sentinel},
            follow_redirects=False,
        )

    assert r.status_code == 500  # plan §3.1: the write did not happen
    assert f"prompt extraction could not be saved to {override}:" in r.text
    assert sentinel in r.text
    assert 'name="prompt_text"' in r.text
    assert "prompt extraction could not be loaded" in r.text  # both reasons


def test_a_save_that_cannot_be_written_over_an_invalid_override_still_names_the_file(
    tmp_path,
):
    """The state where the diagnosis used to be thrown away entirely.

    A READABLE override that fails validation, and a write that fails: the
    `except PromptError` branch of `_prompts_page` renders this page, and that
    branch used to replace the caller's `error` and hardcode its status. So the
    page said the user's text was missing `{qid}` — text that contains `{qid}` —
    at 400, naming no file, while the actual failure was the disk.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "query-translation")
    override.parent.mkdir(parents=True, exist_ok=True)
    override.write_text("Translate the question into Datalog.\n", encoding="utf-8")
    override.chmod(0o444)
    sentinel = "Translate the question {qid} into Datalog."
    try:
        try:
            with override.open("a", encoding="utf-8"):
                pass
        except PermissionError:
            pass
        else:
            pytest.skip("this user writes straight through mode 0o444")

        r = c.post(
            "/prompts",
            data={"prompt_id": "query-translation", "prompt_text": sentinel},
            follow_redirects=False,
        )
    finally:
        override.chmod(0o600)

    assert r.status_code == 500  # plan §3.1: the write did not happen
    assert f"prompt query-translation could not be saved to {override}:" in r.text
    assert sentinel in r.text
    assert 'name="prompt_text"' in r.text


@pytest.mark.parametrize(
    "mode",
    ["dir_0o500_readable_override", "dir_0o000_override", "dir_0o000_no_override"],
)
def test_a_reset_that_cannot_unlink_is_a_page_not_a_crash(tmp_path, mode):
    """Reset is not the working half: `exists()` and `unlink()` fail too.

    The issue measured only an unreadable FILE, where reset legitimately
    redirects. Restrict the DIRECTORY and both calls raise — `unlink()` under
    `0o500`, `Path.exists()` under `0o000`, which propagates `EACCES` instead of
    answering False.

    The usability assertion is the reset form, not `name="prompt_text"`: under
    `0o500` a textarea is present but it belongs to the pre-existing editor and
    says nothing about the reset, and under `0o000` the page is the reset-only
    shape and has none. What the user needs in every one of the three states is
    the control that retries the reset.

    `override.exists()` runs after the `finally` restores the mode, because
    under `0o000` it WOULD raise `PermissionError` from the test itself — the
    same `pathlib` fact the production code is here for. Where it stands it does
    not raise, and the row passes.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "extraction")
    override.parent.mkdir(parents=True, exist_ok=True)
    if mode != "dir_0o000_no_override":
        override.write_text("at most {max_facts} facts\n", encoding="utf-8")
    override.parent.chmod(0o500 if mode == "dir_0o500_readable_override" else 0o000)
    try:
        probe = override.parent / ".write-probe"
        try:
            probe.write_text("x", encoding="utf-8")
        except PermissionError:
            pass
        else:
            probe.unlink()
            pytest.skip("this user writes straight through a restricted directory")

        if mode == "dir_0o500_readable_override":
            # Pinned as a positive because it is the trap in the issue body: a
            # save into a `0o500` directory SUCCEEDS while an override exists,
            # since `write_text` truncates the existing inode. A future
            # `save_prompt_override` switching to temp-file + `os.replace` would
            # need directory write permission and start failing here.
            saved = c.post(
                "/prompts",
                data={
                    "prompt_id": "extraction",
                    "prompt_text": "at most {max_facts} facts, rewritten in place",
                },
                follow_redirects=False,
            )
            assert saved.status_code == 303
            assert "rewritten in place" in override.read_text(encoding="utf-8")

        r = c.post(
            "/prompts/reset",
            data={"prompt_id": "extraction"},
            follow_redirects=False,
        )
    finally:
        override.parent.chmod(0o700)

    assert r.status_code == 500  # plan §3.1: the delete did not happen
    assert f"prompt extraction override could not be deleted from {override}:" in r.text
    assert 'action="/prompts/reset"' in r.text  # the retry the user needs
    if mode != "dir_0o000_no_override":
        assert override.exists()  # nothing was destroyed on the way to the failure
    # WHICH section carried that control is the other half of the sentence in
    # `reset_prompt_route`'s clause, so it is pinned here rather than implied: an
    # override that loads gets the full editor, one that cannot be READ gets the
    # reset-only section. The third shape that sentence names — an override that
    # reads and fails validation — has no control at all and is pinned by
    # `test_a_failed_reset_over_an_invalid_override_offers_no_control`.
    if mode == "dir_0o500_readable_override":
        assert 'name="prompt_text"' in r.text  # the full editor
        assert "Could not load" not in r.text
    else:
        assert 'name="prompt_text"' not in r.text  # the reset-only shape
        assert "Could not load" in r.text


class _Unlisted(Exception):
    """A failure in neither the `ValueError` nor the `OSError` hierarchy."""


@pytest.mark.parametrize("route", ["save", "reset"])
def test_a_prompt_write_failure_of_a_kind_nobody_enumerated_is_still_a_page(
    tmp_path, monkeypatch, route
):
    """The breadth tripwire: narrow either clause to a type list and this dies.

    `except OSError` would carry every filesystem row in this file, so nothing
    else notices the narrowing. `_Unlisted` is outside both hierarchies the
    clause could plausibly be narrowed to.

    Patched on `webapp`, not on `verinote.prompts.library`: both names are
    module-level imports in `verinote/web/app.py` and the routes resolve them
    from module globals at call time.
    """

    def raiser(*args, **kwargs):
        raise _Unlisted("nobody enumerated this")

    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "extraction")

    if route == "save":
        monkeypatch.setattr(webapp, "save_prompt_override", raiser)
        r = c.post(
            "/prompts",
            data={"prompt_id": "extraction", "prompt_text": "Extract the facts."},
            follow_redirects=False,
        )
        expected = f"prompt extraction could not be saved to {override}:"
    else:
        monkeypatch.setattr(webapp, "delete_prompt_override", raiser)
        r = c.post(
            "/prompts/reset",
            data={"prompt_id": "extraction"},
            follow_redirects=False,
        )
        expected = f"prompt extraction override could not be deleted from {override}:"

    assert r.status_code == 500  # plan §3.1: the write did not happen
    assert expected in r.text
    assert "nobody enumerated this" in r.text


def test_an_unpatched_prompt_write_still_redirects(tmp_path):
    """The negative control for the tripwire above: the patch is what fails.

    Every assertion in the tripwire sits downstream of a monkeypatch, so a
    clause that answered 500 to a save that WORKED would satisfy it. This row
    drives the same two routes with nothing patched.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "extraction")

    saved = c.post(
        "/prompts",
        data={"prompt_id": "extraction", "prompt_text": "Extract the facts."},
        follow_redirects=False,
    )

    assert saved.status_code == 303
    assert override.is_file()

    reset = c.post(
        "/prompts/reset",
        data={"prompt_id": "extraction"},
        follow_redirects=False,
    )

    assert reset.status_code == 303
    assert not override.exists()


def test_a_failed_reset_over_an_invalid_override_offers_no_control(tmp_path):
    """The third shape a failed reset renders, and the one with nothing to click.

    An override that READS and fails validation makes `get_prompt` raise
    `PromptError`, so `_prompts_page` takes that branch: `prompt` is None, so no
    editor, and that branch passes `reset_only=False` outright, so no reset-only
    section. `_override_is_unreadable` is the `except Exception` branch's gate
    and is never consulted here — forcing it to return True leaves this response
    byte-identical, which is how that was checked rather than read off. What
    comes back is the banner and the prompt selector.

    Short of a repair affordance, deliberately: offering one for a
    readable-but-invalid override is #546, and the line it has to change is that
    hardcoded `reset_only=False`. It cannot simply become True, because
    `get_prompt` raises the same `PromptError` when the PACKAGED default is what
    fails validation — where a reset would delete the user's file and fix
    nothing.

    What #545 does deliver here is the diagnosis, and the contrast is measured,
    not assumed: on `2c96317` this same request answers a bare 500 with no page
    at all — no banner, no selector, nothing naming the file. Delete the broad
    clause in `reset_prompt_route` and this row returns to that.

    When #546 lands, the two `not in` assertions below are what tell you this
    docstring and `reset_prompt_route`'s comment need rewriting.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "query-translation")
    override.parent.mkdir(parents=True, exist_ok=True)
    # Valid UTF-8 and readable, so this is not the unreadable-override state; it
    # is missing the required `{qid}`, so `get_prompt` refuses it.
    override.write_text("Translate the question into Datalog.\n", encoding="utf-8")
    override.parent.chmod(0o500)
    try:
        probe = override.parent / ".write-probe"
        try:
            probe.write_text("x", encoding="utf-8")
        except PermissionError:
            pass
        else:
            probe.unlink()
            pytest.skip("this user writes straight through mode 0o500")

        r = c.post(
            "/prompts/reset",
            data={"prompt_id": "query-translation"},
            follow_redirects=False,
        )
    finally:
        override.parent.chmod(0o700)

    assert r.status_code == 500  # plan §3.1: the delete did not happen
    assert (
        f"prompt query-translation override could not be deleted from {override}:"
        in r.text
    )
    assert "{qid}" in r.text  # and why the page could not be loaded either
    assert 'value="ask-fallback"' in r.text  # the selector still routes away
    assert 'action="/prompts/reset"' not in r.text  # the #546 gap, pinned
    assert 'name="prompt_text"' not in r.text  # neither section rendered


def test_a_save_for_an_unknown_prompt_id_is_offered_no_save_form(tmp_path):
    """A control for an id `prompt_definition` rejected is a control that lies.

    The echo-back section renders for a save POST whose page could not load the
    prompt, and an unknown `prompt_id` is one of the ways a page fails to load —
    so without the membership test in its gate the section comes back with a
    Save button whose hidden `prompt_id` is the rejected id, and every click on
    it reproduces the same 400.
    `test_a_broken_override_does_not_blank_the_other_prompts` states that rule
    for the sibling reset control; this is the same rule for Save.

    The positive is in this row deliberately: an assertion that a section is
    absent proves nothing unless that section can render at all here, so a known
    id in the same could-not-load state is driven through the same client.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    submitted = "T44 is the text I typed."

    unknown = c.post(
        "/prompts",
        data={"prompt_id": "nope", "prompt_text": submitted},
        follow_redirects=False,
    )

    assert unknown.status_code == 400
    assert "unknown prompt: nope" in unknown.text  # still told why
    assert "Not saved" not in unknown.text
    assert 'name="prompt_text"' not in unknown.text
    assert submitted not in unknown.text
    assert 'value="nope"' not in unknown.text  # no hidden id either

    with _broken_override(prompt_override_path(tmp_path, "query-translation"), "chmod"):
        known = c.post(
            "/prompts",
            data={"prompt_id": "query-translation", "prompt_text": submitted},
            follow_redirects=False,
        )

    assert known.status_code == 400
    assert "Not saved" in known.text  # the section does render for a real id
    assert submitted in known.text


def test_a_save_that_cleared_the_textarea_gets_the_empty_textarea_back(tmp_path):
    """`is not none` and not truthiness, and the empty string is what buys it.

    Clearing the field and pressing Save sends `prompt_text=""`, and the save is
    refused for exactly that. On a truthy gate the echo-back section would not
    render, so over an override that cannot load the page would come back with
    no editor, no Save button, and — the override being readable, so
    `reset_only` is False — no reset control either: a 400 with nothing on the
    page to act on. `is not none` hands the empty field back, and the user can
    type into it and resubmit.

    `""` and not `"   "`: whitespace is truthy, so a whitespace submission
    survives that mutation and would pin nothing. The assertion the mutation
    kills is the textarea one; the status and both banner reasons are already
    covered by `test_a_refusal_over_an_invalid_stored_override_keeps_both_reasons`.
    """
    from verinote.prompts.library import prompt_override_path

    c = _prompts_client(tmp_path)
    override = prompt_override_path(tmp_path, "query-translation")
    override.parent.mkdir(parents=True, exist_ok=True)
    # Readable, so this is the validation-failure state, not the unreadable one.
    override.write_text("Translate the question into Datalog.\n", encoding="utf-8")

    r = c.post(
        "/prompts",
        data={"prompt_id": "query-translation", "prompt_text": ""},
        follow_redirects=False,
    )

    assert r.status_code == 400
    assert "prompt text is required" in r.text  # why the save was refused
    assert "{qid}" in r.text  # and why the page could not be loaded
    assert 'name="prompt_text"' in r.text  # the field comes back, empty
    # And it is the only control on this page: `reset_only` is False here, so
    # without that field the 400 would be a dead end.
    assert 'action="/prompts/reset"' not in r.text
